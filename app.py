# --- 1. 기본 라이브러리 및 설정 ---
import os
import html
import hashlib
import secrets
from urllib.parse import urlencode, urlsplit
import itertools
import random
import math
from dotenv import load_dotenv
from flask import (
    Flask, render_template, request, jsonify, session, redirect, url_for,
    flash, Response, send_file, stream_with_context, abort,
)
import json
from supabase import create_client, Client
from datetime import datetime, timedelta, timezone, date, time
from functools import wraps
import requests
import re
import pandas as pd
import numpy as np
from ortools.sat.python import cp_model
import uuid
import mwparserfromhell
import bleach
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
from group_history import (
    canonical_pair_key as _canonical_pair_key,
    meeting_details_from_history as _meeting_details_from_history,
    pair_keys_from_groups as _pair_keys_from_groups,
    matrix_rows_from_history as _matrix_rows_from_history,
    normalize_group_editor_payload,
)
from topic_preview import anonymous_topic_previews
from topic_document import number_topic_submissions, topic_submitter_identity
from seminar_absence import normalize_member_ids
from seminar_cycle import cycle_monday, is_member_signup_session, next_seminar_cycle
from group_restrictions import find_restriction_conflicts, restricted_pairs_from_rows
from recruitment_results import (
    RESULT_STATUS_LABELS,
    VALID_RESULT_STATUSES,
    normalize_applicant_name,
    normalize_student_id,
    parse_applicant_file,
    parse_applicant_rows,
)
from topic_lifecycle import topic_event_deadline, topic_event_is_expired
from security_utils import forwarded_client_address, safe_internal_next_url
from topic_credentials import (
    generate_topic_edit_token,
    legacy_pin_matches,
    topic_edit_token_digest,
    topic_edit_token_matches,
    topic_request_fingerprint,
)
from member_identity import (
    duplicate_student_id_values,
    member_student_id_conflicts,
    normalize_member_student_id,
    valid_member_student_id,
)

# .env 파일에서 환경 변수 로드
load_dotenv()

# Flask 앱 초기화 및 시크릿 키 설정
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY")
if not app.secret_key:
    raise ValueError("FLASK_SECRET_KEY가 .env 파일에 설정되지 않았습니다.")

# Supabase 클라이언트 초기화
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_SERVICE_KEY")
NOTION_PUBLIC_WIKI_URL = os.environ.get(
    "NOTION_PUBLIC_WIKI_URL",
    "https://quaint-sapphire-900.notion.site/bookbook",
).split('?', 1)[0]
PUBLIC_BASE_URL = os.environ.get(
    "PUBLIC_BASE_URL", "https://book-member-maker.onrender.com"
).rstrip("/")
_cookie_secure_setting = os.environ.get("SESSION_COOKIE_SECURE")
_cookie_secure = (
    _cookie_secure_setting.lower() in {"1", "true", "yes", "on"}
    if _cookie_secure_setting is not None
    else os.environ.get("PUBLIC_BASE_URL", "").startswith("https://")
)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SECURE=_cookie_secure,
    SESSION_COOKIE_SAMESITE="Lax",
    MAX_CONTENT_LENGTH=60 * 1024 * 1024,
)
EXTERNAL_HTTP_TIMEOUT = (5, 15)
if not SUPABASE_URL or not SUPABASE_KEY:
    raise ValueError("Supabase URL과 Key가 .env 파일에 설정되지 않았습니다.")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)


@app.context_processor
def public_navigation_context():
    return {"notion_public_wiki_url": NOTION_PUBLIC_WIKI_URL}


@app.after_request
def add_security_headers(response):
    """Apply low-risk browser protections without breaking existing inline UI scripts."""
    defaults = {
        "X-Content-Type-Options": "nosniff",
        "X-Frame-Options": "DENY",
        "Referrer-Policy": "strict-origin-when-cross-origin",
        "Permissions-Policy": "camera=(), microphone=(), geolocation=()",
    }
    if app.config.get("SESSION_COOKIE_SECURE"):
        defaults["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    for name, value in defaults.items():
        if name not in response.headers:
            response.headers[name] = value
    return response


@app.after_request
def add_default_social_preview(response):
    """모든 HTML에 첫 화면용 밝은 배경과 기본 공유 메타데이터를 넣는다."""
    content_type = response.headers.get("Content-Type", "")
    if response.status_code != 200 or not content_type.startswith("text/html"):
        return response
    page = response.get_data(as_text=True)
    if "</head>" not in page:
        return response
    # Public and legacy pages do not always include the shared header. Load
    # the current theme last so their page-local old palette cannot leak out.
    if 'data-wood-theme="1"' not in page and "/static/theme.css" not in page:
        theme_href = url_for("static", filename="theme.css")
        theme_link = f'\n  <link rel="stylesheet" href="{theme_href}" data-wood-theme="1">\n'
        page = page.replace("</head>", f"{theme_link}</head>", 1)
    if 'data-critical-theme="wood"' not in page:
        critical_theme = """
  <meta name="color-scheme" content="light">
  <meta name="theme-color" content="#FAF6EC">
  <style data-critical-theme="wood">
    html{color-scheme:light;background:#FAF6EC}
    body{background-color:#FAF6EC}
  </style>
"""
        page = page.replace("</head>", f"{critical_theme}</head>", 1)
    if 'property="og:title"' in page:
        response.set_data(page)
        return response
    title_match = re.search(r"<title[^>]*>(.*?)</title>", page, flags=re.I | re.S)
    title = re.sub(r"<[^>]+>", "", title_match.group(1)).strip() if title_match else "책 먹는 호반우"
    title = html.escape(title or "책 먹는 호반우", quote=True)
    description = html.escape(
        "책 먹는 호반우의 세미나, 발제문, 후기와 독서 활동에 참여하세요.",
        quote=True,
    )
    canonical = html.escape(f"{PUBLIC_BASE_URL}{request.full_path}".rstrip("?"), quote=True)
    image = f"{PUBLIC_BASE_URL}/static/og/club-preview.png"
    tags = f"""
  <meta name="description" content="{description}">
  <meta property="og:type" content="website">
  <meta property="og:site_name" content="책 먹는 호반우">
  <meta property="og:locale" content="ko_KR">
  <meta property="og:title" content="{title}">
  <meta property="og:description" content="{description}">
  <meta property="og:url" content="{canonical}">
  <meta property="og:image" content="{image}">
  <meta property="og:image:width" content="1200">
  <meta property="og:image:height" content="630">
  <meta name="twitter:card" content="summary_large_image">
  <meta name="twitter:title" content="{title}">
  <meta name="twitter:description" content="{description}">
  <meta name="twitter:image" content="{image}">
"""
    response.set_data(page.replace("</head>", f"{tags}</head>", 1))
    return response


@app.after_request
def protect_applicant_result_responses(response):
    """Applicant results contain personal decisions and must never be cached/indexed."""
    if request.endpoint == "applicant_result_portal" or session.get("applicant_portal_token"):
        response.headers["Cache-Control"] = "no-store, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["X-Robots-Tag"] = "noindex, nofollow, noarchive"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["Referrer-Policy"] = "no-referrer"
        response.headers["Content-Security-Policy"] = "frame-ancestors 'none'; base-uri 'none'; form-action 'self'"
    return response


@app.after_request
def protect_topic_edit_responses(response):
    """Do not cache pages or API responses that may contain a one-time edit token."""
    if request.endpoint in {'view_shared_topics', 'submit_topics', 'load_topics'}:
        response.headers['Cache-Control'] = 'no-store, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Referrer-Policy'] = 'no-referrer'
    return response


@app.before_request
def reject_cross_site_state_changes():
    """Reject explicit cross-site writes while preserving older same-site webviews."""
    if request.method not in {"POST", "PUT", "PATCH", "DELETE"} or not session:
        return None
    fetch_site = (request.headers.get("Sec-Fetch-Site") or "").lower()
    origin = request.headers.get("Origin")
    origin_mismatch = bool(origin and urlsplit(origin).netloc != request.host)
    if fetch_site == "cross-site" or origin_mismatch:
        if request.path.startswith("/api/"):
            return jsonify({"error": "허용되지 않은 요청 출처입니다."}), 403
        abort(403)
    return None


@app.before_request
def enforce_applicant_portal_boundary():
    """Keep a visitor who entered through a result link inside that portal.

    Static files are allowed, but login, signup, member pages, public tabs and
    unrelated APIs are blocked for the lifetime of the browser session.
    """
    portal_token = session.get("applicant_portal_token")
    if not portal_token:
        return None
    if request.endpoint == "static":
        return None
    # A newly issued result link must be able to replace an expired token in
    # the same Kakao in-app browser. The route clears and rebinds the session.
    if request.endpoint == "applicant_result_portal":
        return None
    if request.path.startswith("/api/"):
        return jsonify({"error": "지원자 결과 조회 링크에서는 다른 기능을 사용할 수 없습니다."}), 403
    return redirect(url_for("applicant_result_portal", token=portal_token), code=303)


# ==============================================================================
# --- 2. 헬퍼 함수 및 공용 시스템 ---
# ==============================================================================

# app.py 의 wiki_parser 함수를 아래 코드로 교체

def wiki_parser(wiki_text):
    """
    mwparserfromhell 라이브러리를 사용해 위키 텍스트를 안전한 HTML로 변환합니다.
    """
    if not wiki_text:
        return ""

    # 1. mwparserfromhell을 사용해 위키 텍스트를 파싱합니다.
    wikicode = mwparserfromhell.parse(wiki_text)

    # 2. 파싱된 코드를 HTML로 변환합니다.
    #    (strip_code는 태그 등을 제거하지만, 기본 HTML 변환에 사용될 수 있습니다.
    #     또는 to_html과 같은 커스텀 변환기를 만들 수도 있습니다.)
    #    여기서는 간단하게 문자열로 변환하여 기본 태그를 유지합니다.
    html = str(wikicode)

    # mwparserfromhell은 [[링크]] 등을 <wikilink> 같은 커스텀 태그로 만들 수 있으나,
    # 여기서는 간단한 문자열 치환으로 링크를 변환합니다.
    # (더 복잡한 변환은 mwparserfromhell의 node 탐색 기능을 사용해야 합니다)

    # 간단한 정규식으로 링크와 강조 등 기본 문법을 HTML 태그로 변환
    # (mwparserfromhell이 구조를 잡아주고, 세부 렌더링은 정규식으로 보완)
    html = re.sub(r"'''(.*?)'''", r'<strong>\1</strong>', html)
    html = re.sub(r"''(.*?)''", r'<em>\1</em>', html)
    html = re.sub(r'\[\[([^\]|]+?)\|([^\]]+?)\]\]', r'<a href="/docs/\1" class="wiki-link">\2</a>', html)
    html = re.sub(r'\[\[([^\]]+?)\]\]', r'<a href="/docs/\1" class="wiki-link">\1</a>', html)
    html = re.sub(r'\[(https?://\S+?)\s+([^\]]+?)\]', r'<a href="\1" target="_blank" rel="noopener noreferrer">\2</a>',
                  html)
    html = re.sub(r'==\s*(.*?)\s*==', r'<h2>\1</h2>', html, flags=re.MULTILINE)

    # 3. 보안을 위해 허용할 태그와 속성을 지정하여 Sanitize 처리
    allowed_tags = {
        'p', 'br', 'strong', 'em', 's', 'blockquote', 'hr',
        'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'ul', 'ol', 'li',
        'table', 'thead', 'tbody', 'tr', 'th', 'td',
        'pre', 'code', 'img', 'a', 'div'
    }
    allowed_attrs = {
        '*': ['class', 'align'],
        'img': ['src', 'alt', 'title'],
        'a': ['href', 'title', 'target', 'rel', 'class'],
    }

    safe_html = bleach.clean(html, tags=allowed_tags, attributes=allowed_attrs, strip=True)
    return safe_html


# 성별 표기 정규화 (DB 표준은 'M'/'W'이지만 과거 데이터에 'F', '남'/'여' 등이 섞일 수 있음)
def normalize_gender(g):
    if g is None:
        return None
    s = str(g).strip().lower()
    if not s or s in ('nan', 'none', 'null'):
        return None
    if s in ('m', 'male', '남', '남성', '남자'):
        return 'M'
    if s in ('w', 'f', 'female', '여', '여성', '여자'):
        return 'W'
    return None


# Jinja2 템플릿에서 날짜 형식을 예쁘게 보여주기 위한 필터
def format_datetime_filter(value, format_str="%Y년 %m월 %d일 %p %I:%M"):
    """
    [최종 수정] 시간대(timezone)를 올바르게 '변환'하여 KST로 표시하는 함수.
    astimezone()을 사용하여 모든 시간 관련 오류를 해결합니다.
    """
    if not value: return ""
    try:
        # 한국 시간대(KST, UTC+9)를 명확하게 정의합니다.
        KST = timezone(timedelta(hours=9))

        # 데이터베이스의 UTC 시간 문자열을 UTC-aware datetime 객체로 변환합니다.
        utc_dt = datetime.fromisoformat(value.replace('Z', '+00:00'))

        # astimezone()을 사용하여 UTC 시간을 KST 시간으로 정확하게 변환합니다.
        kst_dt = utc_dt.astimezone(KST)

        # 변환된 KST 시간을 원하는 형식의 문자열로 만듭니다.
        return kst_dt.strftime(format_str).replace("AM", "오전").replace("PM", "오후")
    except (ValueError, TypeError):
        # 혹시 모를 오류 발생 시, 원래 값을 그대로 보여줍니다.
        return value

# 2. 생성한 함수를 Jinja2 필터로 등록
app.jinja_env.filters['wiki_to_html'] = wiki_parser
app.jinja_env.filters['datetime'] = format_datetime_filter


# 로그인 여부 및 역할 확인을 위한 데코레이터 (문지기 함수)
def _access_denied(message, status=403):
    if request.path.startswith("/api/"):
        return jsonify({"error": message}), status
    flash(message, "warning" if status == 401 else "danger")
    target = "login" if status == 401 else "main_index"
    return redirect(url_for(target))


def login_required(role="ANY"):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if "user_role" not in session:
                return _access_denied("로그인이 필요합니다.", 401)

            user_role = session["user_role"]
            if role == "admin":
                user_id = session.get("user_id")
                try:
                    rows = supabase.table("members").select(
                        "role, is_active, member_status, account_status"
                    ).eq("id", user_id).limit(1).execute().data or []
                except Exception as exc:
                    app.logger.error("Admin authorization refresh failed: %s", exc)
                    return _access_denied("권한을 확인하지 못했습니다. 잠시 후 다시 시도해주세요.", 503)

                if not rows:
                    session.clear()
                    return _access_denied("로그인이 필요합니다.", 401)

                member = rows[0]
                user_role = member.get("role") or "member"
                session["user_role"] = user_role
                if member.get("account_status") != "active" or member.get("member_status") == "inactive":
                    session.clear()
                    return _access_denied("비활성화된 계정입니다.", 403)
                if member.get("is_active") is False or user_role not in ("admin", "officer"):
                    return _access_denied("이 페이지에 접근할 권한이 없습니다.", 403)
            elif role != "ANY" and user_role != role:
                return _access_denied("이 페이지에 접근할 권한이 없습니다.", 403)

            return f(*args, **kwargs)

        return decorated_function

    return decorator


def _current_user_is_primary_admin():
    """Return True only for a currently active member whose DB role is admin.

    The general admin decorator intentionally also permits officers. Confidential
    grouping restrictions require the narrower role and are checked against the
    database so a stale session cannot retain access after a role change.
    """
    user_id = session.get("user_id")
    if not user_id or session.get("user_role") != "admin":
        return False
    try:
        rows = supabase.table("members").select(
            "role, is_active, member_status, account_status"
        ).eq("id", user_id).limit(1).execute().data or []
        if not rows:
            return False
        member = rows[0]
        return bool(
            member.get("role") == "admin"
            and member.get("is_active")
            and member.get("member_status") != "inactive"
            and member.get("account_status") == "active"
        )
    except Exception as exc:
        app.logger.error("Primary admin role check failed: %s", exc)
        return False


def primary_admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user_role" not in session:
            if request.path.startswith("/api/"):
                return jsonify({"error": "로그인이 필요합니다."}), 401
            flash("로그인이 필요합니다.", "warning")
            return redirect(url_for("login"))
        if not _current_user_is_primary_admin():
            if request.path.startswith("/api/"):
                return jsonify({"error": "이 기능을 관리할 권한이 없습니다."}), 403
            flash("이 기능을 관리할 권한이 없습니다.", "danger")
            return redirect(url_for("main_index"))
        return f(*args, **kwargs)

    return decorated_function


def send_telegram_notification(message):
    """[수정] 여러 관리자에게 텔레그램 메시지를 발송합니다."""
    bot_token = os.environ.get("TELEGRAM_BOT_TOKEN")
    # [수정] 쉼표로 구분된 여러 채팅 ID를 불러옵니다.
    chat_ids_str = os.environ.get("TELEGRAM_CHAT_IDS")

    if not bot_token or not chat_ids_str:
        app.logger.error("텔레그램 봇 토큰 또는 관리자 채팅 ID가 설정되지 않았습니다.")
        return

    # 쉼표로 구분된 문자열을 개별 ID 리스트로 변환합니다.
    chat_ids = [chat_id.strip() for chat_id in chat_ids_str.split(',')]

    url = f"https://api.telegram.org/bot{bot_token}/sendMessage"

    # 각 채팅 ID에 대해 메시지를 발송합니다.
    for chat_id in chat_ids:
        params = {
            'chat_id': chat_id,
            'text': message,
            'parse_mode': 'Markdown'
        }
        try:
            response = requests.get(url, params=params, timeout=EXTERNAL_HTTP_TIMEOUT)
            response.raise_for_status()  # 오류 발생 시 예외 발생
            app.logger.info(f"{chat_id}로 텔레그램 알림이 성공적으로 발송되었습니다.")
        except Exception as e:
            app.logger.error(f"{chat_id}로 텔레그램 알림 발송 실패: {e}")

def get_next_monday():
    """오늘을 기준으로 다음 돌아오는 월요일의 날짜를 계산합니다."""
    today = datetime.now(timezone(timedelta(hours=9))).date()
    days_until_monday = (0 - today.weekday() + 7) % 7
    if days_until_monday == 0:
        days_until_monday = 7
    return today + timedelta(days=days_until_monday)


def get_next_seminar_dates():
    """가장 가까운 목요일 본 세미나와 다음 월요일 추가 세미나를 반환한다."""
    today = datetime.now(timezone(timedelta(hours=9))).date()
    return next_seminar_cycle(today)


# ==============================================================================
# --- 3. 로그인, 로그아웃, 메인 페이지 라우트 ---
# ==============================================================================

# [신규] 가장 기본이 되는 메인 페이지 라우트를 추가합니다.
@app.route('/keep-alive')
def keep_alive_endpoint():
    """Render liveness stays healthy even when the database is temporarily degraded."""
    checked_at = datetime.now(timezone.utc).isoformat()
    try:
        supabase.table('members').select('id').limit(1).execute()
        return jsonify({
            "status": "ok",
            "checked_at": checked_at,
            "supabase": "alive",
        }), 200
    except Exception as exc:
        app.logger.error("keep-alive dependency check failed: %s", exc)
        return jsonify({
            "status": "degraded",
            "checked_at": checked_at,
            "supabase": "unavailable",
        }), 200


@app.route('/')
def main_index():
    # --- [수정] D-데이 계산 로직 추가 ---
    try:
        # 모집 마감일을 설정합니다. (년, 월, 일)
        end_date_str = "2025-08-31"
        # templates/main_index.html 파일의 모집 기간 마지막 날짜와 일치시킵니다.
        end_date = datetime.strptime(end_date_str, "%Y-%m-%d")

        # 오늘 날짜를 가져옵니다.
        today = datetime.now()

        # 남은 날짜(D-데이)를 계산합니다.
        # 날짜만 비교하기 위해 .date()를 사용합니다.
        delta = end_date.date() - today.date()
        d_day = delta.days
    except Exception as e:
        app.logger.error(f"D-day calculation error: {e}")
        d_day = -1  # 오류 발생 시 기간이 지난 것으로 처리

    # 계산된 d_day 값을 템플릿으로 전달합니다.
    return render_template('main_index.html', d_day=d_day)


class KakaoOauth:
    def __init__(self):
        self.client_id = os.environ.get("KAKAO_OAUTH_CLIENT_ID")
        self.redirect_uri = os.environ.get("KAKAO_REDIRECT_URI")
        self.token_url = "https://kauth.kakao.com/oauth/token"
        self.user_info_url = "https://kapi.kakao.com/v2/user/me"

    def get_token(self, code):
        """인가 코드로 Access Token을 요청합니다."""
        data = {
            "grant_type": "authorization_code",
            "client_id": self.client_id,
            "redirect_uri": self.redirect_uri,
            "code": code,
        }
        # [수정] Client Secret을 활성화했으므로, 이 줄의 주석을 반드시 해제하고 값을 추가합니다.
        if os.environ.get("KAKAO_OAUTH_CLIENT_SECRET"):
            data["client_secret"] = os.environ.get("KAKAO_OAUTH_CLIENT_SECRET")

        response = requests.post(self.token_url, data=data, timeout=EXTERNAL_HTTP_TIMEOUT)
        response.raise_for_status()  # 오류 발생 시 예외 발생
        return response.json()

    def get_user_info(self, access_token):
        """Access Token으로 사용자 정보를 요청합니다."""
        headers = {"Authorization": f"Bearer {access_token}"}
        response = requests.post(self.user_info_url, headers=headers, timeout=EXTERNAL_HTTP_TIMEOUT)
        response.raise_for_status()
        return response.json()


# --- 기존 로그인 라우트들을 아래 코드로 교체합니다 ---

@app.route('/login')
def login():
    # 이 페이지는 이제 카카오 로그인 버튼만 보여줍니다.
    return render_template('login.html')


def _kakao_authorize_url(kakao_oauth, prompt=None):
    state = secrets.token_urlsafe(32)
    session["oauth_state"] = state
    params = {
        "client_id": kakao_oauth.client_id,
        "redirect_uri": kakao_oauth.redirect_uri,
        "response_type": "code",
        "state": state,
    }
    if prompt:
        params["prompt"] = prompt
    return "https://kauth.kakao.com/oauth/authorize?" + urlencode(params)


@app.route('/login/kakao')
def kakao_login():
    kakao_oauth = KakaoOauth()
    session['next_url'] = safe_internal_next_url(request.args.get('next'))
    mode = request.args.get('mode', 'login')
    session['auth_mode'] = mode if mode in ('login', 'signup') else 'login'
    return redirect(_kakao_authorize_url(kakao_oauth))


@app.route('/login/kakao/re-consent')
@login_required(role="ANY")
def kakao_reconsent_login():
    kakao_oauth = KakaoOauth()
    session['next_url'] = url_for('my_page')
    return redirect(_kakao_authorize_url(kakao_oauth, prompt="consent"))


# 2. 로그인 후 콜백을 처리할 라우트
@app.route('/login/kakao/callback')
def kakao_callback():
    try:
        expected_state = session.pop("oauth_state", None)
        received_state = request.args.get("state")
        if not expected_state or not received_state or not secrets.compare_digest(expected_state, received_state):
            session.pop("next_url", None)
            session.pop("auth_mode", None)
            flash("로그인 요청을 확인할 수 없습니다. 다시 시도해주세요.", "danger")
            return redirect(url_for("login"))

        code = request.args.get("code")
        if not code:
            flash("인증 코드를 받는데 실패했습니다.", "danger")
            return redirect(url_for("login"))

        kakao_oauth = KakaoOauth()
        token_info = kakao_oauth.get_token(code)
        access_token = token_info.get("access_token")
        user_info = kakao_oauth.get_user_info(access_token)

        social_id = str(user_info["id"])
        kakao_account = user_info.get("kakao_account", {})
        profile = kakao_account.get("profile", {})
        email = kakao_account.get("email")

        member_res = supabase.table("members").select("*").eq("social_id", social_id).execute()
        member = member_res.data[0] if member_res.data else None

        if member:
            # [수정] 세션 업데이트 로직 추가
            update_data = {}
            new_name = member['name']  # 기본값은 기존 이름

            if profile.get("profile_image_url"):
                update_data['profile_pic'] = profile.get("profile_image_url")

            if profile.get("nickname"):
                new_name = profile.get("nickname")
                update_data['name'] = new_name  # 닉네임도 함께 업데이트

            if update_data:
                supabase.table("members").update(update_data).eq("id", member['id']).execute()
                flash("카카오 프로필 정보가 업데이트되었습니다.", "success")
                # [핵심] DB 업데이트 후, 세션에 저장된 이름도 새로운 이름으로 갱신
                session['user_name'] = new_name

            # ... (계정 상태 및 활성 상태 체크 로직은 기존과 동일) ...
            if member.get('account_status') != 'active':
                flash("승인 대기 중입니다. 관리자가 가입/연동 요청을 확인 중입니다.", "warning")
                return redirect(url_for('login'))
            if member.get('member_status', 'active') == 'inactive':
                flash("비활성화된 계정입니다. 관리자에게 문의하세요.", "danger")
                return redirect(url_for('login'))
        else:
            # ... (신규 사용자 '계정 연결' 로직은 기존과 동일) ...
            social_data = {
                "social_id": social_id, "email": email,
                "social_name": profile.get("nickname"), "profile_pic": profile.get("profile_image_url")
            }
            session['temp_social_data'] = social_data
            return redirect(url_for('link_account_page'))

        # 세션 설정 (DB에서 읽은 최신 값으로 매번 갱신)
        session["user_id"] = member["id"]
        session["user_role"] = member["role"]
        session["user_name"] = member["name"]  # 항상 DB 최신값으로 갱신
        session.pop("member_preview", None)

        next_url = safe_internal_next_url(session.pop('next_url', None))
        if next_url:
            return redirect(next_url)
        if member["role"] in ("admin", "officer"):
            return redirect(url_for("admin_dashboard"))
        return redirect(url_for("my_page"))

    except Exception as e:
        flash("카카오 로그인 중 오류가 발생했습니다.", "danger")
        app.logger.error(f"Kakao callback error: {e}", exc_info=True)
        return redirect(url_for("login"))


@app.route('/link_account')
def link_account_page():
    # kakao_callback에서 임시 저장한 소셜 데이터가 없으면 로그인 페이지로
    if 'temp_social_data' not in session:
        return redirect(url_for('login'))

    social_data = session['temp_social_data']
    auth_mode = session.get('auth_mode', 'login')

    # [신규] 카카오 닉네임으로 미연결 멤버 자동 매칭
    # 동명이인 가능성 고려해 결과 0건 / 1건 / 2건 이상으로 분기
    matched_member = None
    multiple_matches = []
    nickname = (social_data.get('social_name') or '').strip()
    if nickname:
        try:
            res = supabase.table("members").select("id,name,student_id,profile_pic")\
                .eq("name", nickname)\
                .or_("social_id.is.null,social_id.eq.").execute()
            rows = res.data or []
            if len(rows) == 1:
                matched_member = rows[0]
            elif len(rows) > 1:
                multiple_matches = rows
        except Exception as e:
            app.logger.warning(f"link_account auto-match failed: {e}")

    return render_template('link_account.html',
                           **social_data,
                           auth_mode=auth_mode,
                           matched_member=matched_member,
                           multiple_matches=multiple_matches)


@app.route('/api/link/lookup', methods=['POST'])
def link_lookup():
    """계정 연동 중 이름과 학번으로 미연동 멤버 프로필을 조회합니다."""
    if 'temp_social_data' not in session:
        return jsonify({'error': '인증 세션이 만료되었습니다. 카카오 로그인부터 다시 진행해주세요.'}), 403
    data = request.json or {}
    name = (data.get('name') or '').strip()
    sid = (data.get('student_id') or '').strip()
    if not name:
        return jsonify({'error': '동아리 명부에 등록된 이름을 입력해주세요.'}), 400
    if not sid.isdigit() or len(sid) < 4:
        return jsonify({'error': '학번을 숫자로 정확히 입력해주세요.'}), 400
    try:
        res = supabase.table('members').select('id, name, department, recruiting_class, social_id') \
            .eq('student_id', sid).eq('name', name).execute()
        rows = res.data or []
        if not rows:
            return jsonify({'found': False})
        if all(r.get('social_id') for r in rows):
            return jsonify({'found': False, 'already_linked': True})
        m = next(r for r in rows if not r.get('social_id'))
        name = (m.get('name') or '').strip()
        if len(name) >= 3:
            masked = name[0] + '*' * (len(name) - 2) + name[-1]
        elif len(name) == 2:
            masked = name[0] + '*'
        else:
            masked = name
        return jsonify({
            'found': True,
            'masked_name': masked,
            'department': m.get('department') or '',
            'recruiting_class': m.get('recruiting_class'),
        })
    except Exception as e:
        app.logger.error(f"link_lookup error: {e}")
        return jsonify({'error': '조회 중 오류가 발생했습니다.'}), 500


@app.route('/link_account', methods=['POST'])
def link_account_submit():
    if 'temp_social_data' not in session:
        return redirect(url_for('login'))

    form = request.form
    action = form.get('action')
    social_info = session['temp_social_data']

    member = None
    if action == 'link':
        existing_name = form.get('existing_name', '').strip()
        student_id = form.get('student_id', '').strip()

        if existing_name and student_id:
            # 자동 승인 경로는 이름과 학번을 서버에서 다시 함께 검증합니다.
            member_res = supabase.table("members").select("*").eq("name", existing_name)\
                .eq("student_id", student_id)\
                .or_("social_id.is.null,social_id.eq.").execute()
        elif existing_name:
            member_res = supabase.table("members").select("*").eq("name", existing_name)\
                .or_("social_id.is.null,social_id.eq.").execute()
        else:
            flash("기존 활동명을 입력해주세요.", "danger")
            return redirect(url_for('link_account_page'))
        member_to_link = member_res.data[0] if member_res.data else None

        if member_to_link:
            # 이름과 학번이 모두 DB 값과 정확히 일치할 때만 자동 승인합니다.
            db_student_id = str(member_to_link.get('student_id') or '').strip()
            db_name = str(member_to_link.get('name') or '').strip()
            auto_approve = bool(
                existing_name and student_id
                and existing_name == db_name
                and student_id == db_student_id
            )

            update_data = {
                "social_id": social_info['social_id'],
                "profile_pic": social_info['profile_pic'],
                "account_status": 'active' if auto_approve else 'pending',
                "is_active": True if auto_approve else member_to_link.get('is_active', False)
            }
            # 이메일이 있고, 현재 멤버가 사용 중인 이메일이 아닌 경우에만 업데이트
            # (다른 멤버가 이미 같은 이메일을 사용 중이면 UNIQUE 제약 위반 방지)
            kakao_email = social_info.get('email')
            if kakao_email:
                email_conflict = supabase.table("members").select("id").eq("email", kakao_email).neq("id", member_to_link['id']).execute()
                if not email_conflict.data:
                    update_data["email"] = kakao_email
                else:
                    app.logger.warning(f"이메일 {kakao_email}이 이미 다른 멤버에게 사용 중 — 이메일 업데이트 생략")
            updated_member_response = supabase.table("members").update(update_data).eq("id", member_to_link['id']).execute()
            member = updated_member_response.data[0]

            if auto_approve:
                # 세션 설정 및 자동 로그인
                session.pop('temp_social_data', None)
                session['user_id'] = member['id']
                session['user_name'] = member['name']
                session['user_role'] = member.get('role', 'member')
                session.pop('member_preview', None)
                session['profile_pic'] = member.get('profile_pic', '')
                flash(f"학번 확인이 완료되었습니다. {member['name']}님, 환영합니다!", "success")
                return redirect(url_for('my_page'))
            else:
                # notifications 테이블에 알림 생성
                supabase.table('notifications').insert({
                    'type': 'account_link_request',
                    'related_member_id': member['id'],
                    'details': {
                        'original_name': member['name'],
                        'social_name': social_info['social_name'],
                        'social_email': social_info['email']
                    }
                }).execute()
                flash("기존 계정 연결 요청이 완료되었습니다. 관리자 승인 후 로그인 가능합니다.", "success")
                return redirect(url_for('login'))
        else:
            flash("해당 이름의 기존 계정을 찾을 수 없습니다. 신규 회원으로 가입해주세요.", "danger")
            return redirect(url_for('link_account_page'))

    elif action == 'create':
        new_member_data = {
            "name": social_info['social_name'],
            "email": social_info['email'],
            "social_id": social_info['social_id'],
            "profile_pic": social_info['profile_pic'],
            "role": "member",
            "account_status": 'pending'
        }
        # [수정] .insert() 뒤에 .select()를 제거하고, 실행 결과에서 바로 .data를 사용합니다.
        new_member_response = supabase.table("members").insert(new_member_data).execute()
        member = new_member_response.data[0]

        # notifications 테이블에 알림 생성
        supabase.table('notifications').insert({
            'type': 'new_user_request',
            'related_member_id': member['id'],
            'details': {
                'name': member['name'],
                'email': member['email']
            }
        }).execute()

        flash("회원가입 요청이 완료되었습니다. 관리자 승인 후 활동 가능합니다.", "success")
        return redirect(url_for('login'))






@app.route('/logout')
def logout():
    session.clear()
    flash('성공적으로 로그아웃되었습니다.', 'info')
    return redirect(url_for('main_index'))


@app.route('/toggle-member-preview', methods=['POST'])
@login_required(role="ANY")
def toggle_member_preview():
    user_id = session.get('user_id')
    members = supabase.table('members').select('role').eq('id', user_id).limit(1).execute().data or []
    if not members or members[0].get('role') != 'admin':
        abort(403)

    if session.get('member_preview'):
        session.pop('member_preview', None)
        session['user_role'] = 'admin'
        flash('관리자 화면으로 돌아왔습니다.', 'success')
        return redirect(url_for('admin_seminars'))

    session['member_preview'] = True
    session['user_role'] = 'member'
    flash('회원 화면으로 전환했습니다.', 'success')
    return redirect(url_for('my_page'))





@app.route('/api/attendance', methods=['POST'])
@login_required(role="ANY")
def update_attendance():
    return jsonify({'status': 'error', 'message': '세미나 참석·불참은 카카오톡 투표에서 변경해주세요.'}), 410


@app.route('/api/questions', methods=['POST'])
@login_required(role="ANY")
def create_question():
    """JavaScript(fetch) 요청을 처리하는 API 라우트"""
    data = request.json
    user_id = session.get('user_id')
    try:
        new_question = supabase.table('questions').insert({
            'user_id': user_id,
            'meeting_date': get_next_monday().isoformat(),
            'content': data.get('content')
        }).execute().data[0]

        member_res = supabase.table('members').select('name').eq('id', user_id).execute().data[0]
        new_question['members'] = member_res

        return jsonify({"status": "success", "question": new_question})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/questions/<int:question_id>', methods=['PUT', 'DELETE'])
@login_required(role="ANY")
def manage_question(question_id):
    """JavaScript(fetch) 요청을 처리하는 API 라우트"""
    user_id = session.get('user_id')
    try:
        question_res = supabase.table('questions').select('user_id').eq('id', question_id).single().execute()
        if not question_res.data or question_res.data['user_id'] != user_id:
            return jsonify({"status": "error", "message": "권한이 없습니다."}), 403

        if request.method == 'PUT':
            content = request.json.get('content')
            supabase.table('questions').update({'content': content}).eq('id', question_id).execute()
            return jsonify({"status": "success", "message": "질문이 수정되었습니다."})
        elif request.method == 'DELETE':
            supabase.table('questions').delete().eq('id', question_id).execute()
            return jsonify({"status": "success", "message": "질문이 삭제되었습니다."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/profiles')
@login_required(role="ANY")
def profiles_page():
    """모든 멤버의 프로필 목록을 보여주는 페이지"""
    try:
        all_members = supabase.table('members').select('*').eq('is_active', True).order('name').execute().data
    except Exception as e:
        flash(f"프로필 로딩 중 오류: {e}", "danger")
        all_members = []
    return render_template('profiles.html', all_members=all_members)


@app.route('/api/profiles/update', methods=['POST'])
@login_required(role="ANY")
def update_profile():
    """프로필의 '한 줄 소개'와 '상세 프로필(BBCode)'을 업데이트합니다."""
    user_id = session.get('user_id')
    data = request.json

    try:
        # 1. 클라이언트로부터 필요한 데이터만 추출합니다.
        update_data = {
            'profile_intro': data.get('intro'),
            'profile_content': data.get('content')  # BBCode 원문을 그대로 저장
        }

        # 2. 값이 없는 필드는 업데이트 대상에서 제외합니다.
        update_data = {k: v for k, v in update_data.items() if v is not None}

        # 3. 업데이트할 데이터가 없으면 오류를 반환합니다.
        if not update_data:
            return jsonify({"status": "error", "message": "전송된 데이터가 없습니다."}), 400

        # 4. 데이터베이스를 업데이트합니다.
        supabase.table('members').update(update_data).eq('id', user_id).execute()

        # 5. 프론트엔드가 페이지를 새로고침하므로 간단한 성공 메시지만 반환합니다.
        return jsonify({"status": "success", "message": "프로필이 저장되었습니다."})

    except Exception as e:
        app.logger.error(f"Profile update error: {e}", exc_info=True)
        return jsonify({"status": "error", "message": "프로필 저장 중 서버 오류가 발생했습니다."}), 500


@app.route('/profile/<int:member_id>')
@login_required(role="ANY")
def profile_detail_page(member_id):
    """특정 멤버의 상세 프로필 페이지를 보여줍니다."""
    try:
        # DB에서 요청된 ID의 멤버 정보를 조회합니다.
        member_res = supabase.table('members').select('*').eq('id', member_id).single().execute()
        member_data = member_res.data
        if not member_data:
            flash("존재하지 않는 회원입니다.", "danger")
            return redirect(url_for('profiles_page'))

    except Exception as e:
        flash(f"프로필 로딩 중 오류: {e}", "danger")
        return redirect(url_for('profiles_page'))

    # 새로 만들 profile_detail.html 템플릿으로 데이터를 전달합니다.
    return render_template('profile_detail.html', member=member_data)

# ==============================================================================
# <editor-fold desc="4. 관리자 (Admin) 전용 기능">
# --- 4.1. 관리자 대시보드 및 이벤트 관리 ---
@app.route('/admin/dashboard')
@login_required(role="admin")
def admin_dashboard():
    try:
        next_monday = get_next_monday()
    except Exception:
        next_monday = None

    try:
        all_members_res = supabase.table('members').select(
            'id, name, is_active, member_status, role, email, department, gender, student_id, recruiting_class'
        ).order('name').execute().data

        # 생성된 발제문 취합 이벤트 목록 불러오기
        topic_events_res = supabase.table('topic_events').select('*').order('created_at', desc=True).execute()
        topic_events = _auto_close_topic_events(topic_events_res.data or [])

        # 최근 조 편성 기록 (바로 수정하러 갈 수 있게 대시보드에 노출)
        latest_history_res = supabase.table('history').select('id, date, book_title') \
            .order('date', desc=True).limit(1).execute()
        latest_history = latest_history_res.data[0] if latest_history_res.data else None

        # 세미나 출석 투표 학기 목록
        seminar_terms_res = supabase.table('seminar_terms').select('*') \
            .order('start_date', desc=True).execute()
        seminar_terms = seminar_terms_res.data or []
        # 회차 수 카운트
        if seminar_terms:
            term_ids = [t['id'] for t in seminar_terms]
            sess_counts_res = supabase.table('seminar_sessions').select('term_id') \
                .in_('term_id', term_ids).execute()
            counts = {}
            for row in (sess_counts_res.data or []):
                counts[row['term_id']] = counts.get(row['term_id'], 0) + 1
            for t in seminar_terms:
                t['session_count'] = counts.get(t['id'], 0)
                t['share_url'] = f"{request.host_url}seminar_vote?token={t['share_token']}"

    except Exception as e:
        flash(f"대시보드 로딩 중 오류 발생: {e}", "danger")
        all_members_res = []
        topic_events = []
        seminar_terms = []
        latest_history = None

    return render_template(
        'admin_overview.html',
        all_members=all_members_res,
        meeting_date=next_monday,
        topic_events=topic_events,
        seminar_terms=seminar_terms,
        latest_history=latest_history,
    )


def _recruitment_campaign(campaign_id):
    rows = supabase.table("recruitment_campaigns").select("*") \
        .eq("id", str(campaign_id)).limit(1).execute().data or []
    return rows[0] if rows else None


def _applicant_request_hash():
    address = forwarded_client_address(
        request.headers.get("X-Forwarded-For"), request.remote_addr
    )
    secret = str(app.secret_key or "")
    return hashlib.sha256(f"{secret}|{address}".encode("utf-8")).hexdigest()


@app.route('/applicant-result/<uuid:token>', methods=['GET', 'POST'])
def applicant_result_portal(token):
    """Isolated, no-login result lookup by exact name and student ID."""
    token_text = str(token)
    is_staff_preview = session.get("user_role") in ("admin", "officer")
    if not is_staff_preview and session.get("applicant_portal_token") != token_text:
        session.clear()
        session["applicant_portal_token"] = token_text

    campaigns = supabase.table("recruitment_campaigns").select("*") \
        .eq("share_token", token_text).limit(1).execute().data or []
    campaign = campaigns[0] if campaigns else None
    if not campaign:
        return render_template(
            "applicant_result_portal.html",
            campaign=None,
            portal_state="unavailable",
            result=None,
            error_message=None,
        ), 404

    if not campaign.get("is_active"):
        portal_state = "closed"
    elif not campaign.get("is_published"):
        portal_state = "pending"
    else:
        portal_state = "lookup"

    result = None
    error_message = None
    response_status = 200
    if request.method == 'POST' and portal_state == "lookup":
        name = (request.form.get("name") or "").strip()
        student_id = normalize_student_id(request.form.get("student_id"))
        name_key = normalize_applicant_name(name)
        if not name_key or not re.fullmatch(r"[0-9]{4,20}", student_id):
            error_message = "이름과 학번을 정확히 입력해주세요."
        else:
            ip_hash = _applicant_request_hash()
            rate_start = (datetime.now(timezone.utc) - timedelta(minutes=15)).isoformat()
            attempts = supabase.table("recruitment_lookup_attempts").select("id", count="exact") \
                .eq("campaign_id", campaign["id"]).eq("ip_hash", ip_hash) \
                .gte("created_at", rate_start).execute()
            if (attempts.count or 0) >= 10:
                error_message = "조회가 잠시 제한되었습니다. 15분 후 다시 시도해주세요."
                response_status = 429
            else:
                matches = supabase.table("recruitment_applicants").select(
                    "id, name, result_status, personal_message"
                ).eq("campaign_id", campaign["id"]).eq("student_id", student_id) \
                    .eq("name_key", name_key).limit(1).execute().data or []
                matched = matches[0] if matches else None
                try:
                    supabase.table("recruitment_lookup_attempts").insert({
                        "campaign_id": campaign["id"],
                        "ip_hash": ip_hash,
                        "succeeded": bool(matched),
                    }).execute()
                except Exception as exc:
                    app.logger.warning("recruitment lookup audit failed: %s", exc)
                if not matched:
                    error_message = "입력한 정보와 일치하는 결과를 찾지 못했습니다. 이름과 학번을 다시 확인해주세요."
                else:
                    status = matched.get("result_status") or "pending"
                    result = {
                        "name": matched.get("name") or name,
                        "status": status,
                        "status_label": RESULT_STATUS_LABELS.get(status, "발표 전"),
                        "common_message": campaign.get(f"{status}_message") or campaign.get("pending_message") or "",
                        "personal_message": matched.get("personal_message") or "",
                    }

    return render_template(
        "applicant_result_portal.html",
        campaign=campaign,
        portal_state=portal_state,
        result=result,
        error_message=error_message,
    ), response_status


@app.route('/admin/recruitment-results')
@login_required(role="admin")
def admin_recruitment_results():
    campaigns = supabase.table("recruitment_campaigns").select("*") \
        .order("created_at", desc=True).execute().data or []
    applicant_rows = supabase.table("recruitment_applicants").select("campaign_id, result_status").execute().data or []
    campaign_counts = defaultdict(lambda: defaultdict(int))
    for row in applicant_rows:
        campaign_counts[row.get("campaign_id")][row.get("result_status") or "pending"] += 1
    for campaign in campaigns:
        counts = campaign_counts[campaign["id"]]
        campaign["applicant_count"] = sum(counts.values())
        campaign["status_counts"] = dict(counts)
        campaign["share_url"] = f"{PUBLIC_BASE_URL}{url_for('applicant_result_portal', token=campaign['share_token'])}"
    return render_template("admin_recruitment_results.html", campaigns=campaigns)


@app.route('/admin/recruitment-results/create', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_results_create():
    title = (request.form.get("title") or "").strip()
    if not title:
        flash("모집 차수 이름을 입력해주세요.", "danger")
        return redirect(url_for("admin_recruitment_results"))
    try:
        created = supabase.table("recruitment_campaigns").insert({
            "title": title,
            "created_by": session.get("user_id"),
        }).execute().data or []
        flash("지원자 결과 페이지를 만들었습니다. 명단을 넣고 발표 상태를 확인해주세요.", "success")
        return redirect(url_for("admin_recruitment_result_detail", campaign_id=created[0]["id"]))
    except Exception as exc:
        app.logger.error("recruitment campaign create failed: %s", exc, exc_info=True)
        flash("모집 차수를 만들지 못했습니다.", "danger")
        return redirect(url_for("admin_recruitment_results"))


@app.route('/admin/recruitment-results/<uuid:campaign_id>')
@login_required(role="admin")
def admin_recruitment_result_detail(campaign_id):
    campaign = _recruitment_campaign(campaign_id)
    if not campaign:
        flash("모집 차수를 찾을 수 없습니다.", "danger")
        return redirect(url_for("admin_recruitment_results"))
    applicants = supabase.table("recruitment_applicants").select("*") \
        .eq("campaign_id", str(campaign_id)).order("name").execute().data or []
    counts = {status: 0 for status in VALID_RESULT_STATUSES}
    for applicant in applicants:
        counts[applicant.get("result_status") or "pending"] += 1
    campaign["share_url"] = f"{PUBLIC_BASE_URL}{url_for('applicant_result_portal', token=campaign['share_token'])}"
    return render_template(
        "admin_recruitment_result_detail.html",
        campaign=campaign,
        applicants=applicants,
        status_counts=counts,
        status_labels=RESULT_STATUS_LABELS,
    )


@app.route('/admin/recruitment-results/<uuid:campaign_id>/settings', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_result_settings(campaign_id):
    if not _recruitment_campaign(campaign_id):
        return redirect(url_for("admin_recruitment_results"))
    title = (request.form.get("title") or "").strip()
    if not title:
        flash("모집 차수 이름은 비워둘 수 없습니다.", "danger")
        return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))
    update = {
        "title": title,
        "intro_text": (request.form.get("intro_text") or "").strip(),
        "pending_message": (request.form.get("pending_message") or "").strip(),
        "accepted_message": (request.form.get("accepted_message") or "").strip(),
        "waitlisted_message": (request.form.get("waitlisted_message") or "").strip(),
        "rejected_message": (request.form.get("rejected_message") or "").strip(),
        "contact_text": (request.form.get("contact_text") or "").strip() or None,
        "is_active": request.form.get("is_active") == "on",
        "is_published": request.form.get("is_published") == "on",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if any(not update[key] for key in ("intro_text", "pending_message", "accepted_message", "waitlisted_message", "rejected_message")):
        flash("공통 안내문은 비워둘 수 없습니다.", "danger")
        return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))
    try:
        supabase.table("recruitment_campaigns").update(update).eq("id", str(campaign_id)).execute()
        flash("공개 설정과 안내문을 저장했습니다.", "success")
    except Exception as exc:
        app.logger.error("recruitment settings update failed: %s", exc, exc_info=True)
        flash("설정을 저장하지 못했습니다. 글자 수를 확인해주세요.", "danger")
    return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))


@app.route('/admin/recruitment-results/<uuid:campaign_id>/rotate-link', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_result_rotate_link(campaign_id):
    if _recruitment_campaign(campaign_id):
        supabase.table("recruitment_campaigns").update({
            "share_token": str(uuid.uuid4()),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }).eq("id", str(campaign_id)).execute()
        flash("공유 링크를 새로 발급했습니다. 이전 링크는 즉시 사용할 수 없습니다.", "success")
    return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))


@app.route('/admin/recruitment-results/<uuid:campaign_id>/bulk-import', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_result_bulk_import(campaign_id):
    if not _recruitment_campaign(campaign_id):
        return redirect(url_for("admin_recruitment_results"))
    raw_rows = request.form.get("applicant_rows") or ""
    uploaded = request.files.get("applicant_file")
    has_file = bool(uploaded and uploaded.filename)
    if has_file and raw_rows.strip():
        flash("Excel 파일 업로드와 명단 붙여넣기 중 한 가지만 사용해주세요.", "danger")
        return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))
    if has_file:
        parsed, errors = parse_applicant_file(uploaded.filename, uploaded.read(2 * 1024 * 1024 + 1))
    else:
        parsed, errors = parse_applicant_rows(raw_rows)
    if errors:
        flash(" / ".join(errors[:8]), "danger")
        return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))
    if not parsed:
        flash("추가할 지원자 명단을 붙여넣어 주세요.", "danger")
        return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))
    student_ids = [row["student_id"] for row in parsed]
    existing_rows = supabase.table("recruitment_applicants").select(
        "student_id, result_status, personal_message"
    ).eq("campaign_id", str(campaign_id)).in_("student_id", student_ids).execute().data or []
    existing_by_id = {row["student_id"]: row for row in existing_rows}
    payload = []
    for row in parsed:
        existing = existing_by_id.get(row["student_id"], {})
        status_provided = row.pop("_status_provided")
        message_provided = row.pop("_message_provided")
        if existing and not status_provided:
            row["result_status"] = existing.get("result_status") or "pending"
        if existing and not message_provided:
            row["personal_message"] = existing.get("personal_message")
        row["campaign_id"] = str(campaign_id)
        row["updated_at"] = datetime.now(timezone.utc).isoformat()
        payload.append(row)
    try:
        supabase.table("recruitment_applicants").upsert(
            payload, on_conflict="campaign_id,student_id"
        ).execute()
        flash(f"지원자 {len(payload)}명의 명단을 반영했습니다.", "success")
    except Exception as exc:
        app.logger.error("recruitment bulk import failed: %s", exc, exc_info=True)
        flash("명단을 반영하지 못했습니다. 중복 학번과 입력 형식을 확인해주세요.", "danger")
    return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))


@app.route('/admin/recruitment-results/<uuid:campaign_id>/applicants/create', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_result_create_applicant(campaign_id):
    if not _recruitment_campaign(campaign_id):
        return redirect(url_for("admin_recruitment_results"))
    name = (request.form.get("name") or "").strip()
    student_id = normalize_student_id(request.form.get("student_id"))
    result_status = (request.form.get("result_status") or "pending").strip()
    personal_message = (request.form.get("personal_message") or "").strip() or None
    if not normalize_applicant_name(name) or not re.fullmatch(r"[0-9]{4,20}", student_id):
        flash("지원자 이름과 4~20자리 숫자 학번을 확인해주세요.", "danger")
    elif result_status not in VALID_RESULT_STATUSES:
        flash("지원자 결과를 확인해주세요.", "danger")
    elif personal_message and len(personal_message) > 3000:
        flash("개인 안내는 3,000자 이하로 입력해주세요.", "danger")
    else:
        try:
            supabase.table("recruitment_applicants").insert({
                "campaign_id": str(campaign_id),
                "name": name,
                "name_key": normalize_applicant_name(name),
                "student_id": student_id,
                "result_status": result_status,
                "personal_message": personal_message,
            }).execute()
            flash(f"{name} 지원자를 명단에 추가했습니다.", "success")
        except Exception as exc:
            app.logger.error("recruitment applicant create failed: %s", exc, exc_info=True)
            flash("지원자를 추가하지 못했습니다. 같은 학번이 이미 있는지 확인해주세요.", "danger")
    return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))


@app.route('/admin/recruitment-results/<uuid:campaign_id>/applicants/<uuid:applicant_id>', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_result_update_applicant(campaign_id, applicant_id):
    name = (request.form.get("name") or "").strip()
    student_id = normalize_student_id(request.form.get("student_id"))
    result_status = (request.form.get("result_status") or "").strip()
    personal_message = (request.form.get("personal_message") or "").strip() or None
    if not normalize_applicant_name(name) or not re.fullmatch(r"[0-9]{4,20}", student_id) or result_status not in VALID_RESULT_STATUSES:
        flash("지원자 이름·학번·결과를 확인해주세요.", "danger")
    else:
        try:
            supabase.table("recruitment_applicants").update({
                "name": name,
                "name_key": normalize_applicant_name(name),
                "student_id": student_id,
                "result_status": result_status,
                "personal_message": personal_message,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }).eq("id", str(applicant_id)).eq("campaign_id", str(campaign_id)).execute()
            flash(f"{name} 지원자의 결과를 저장했습니다.", "success")
        except Exception as exc:
            app.logger.error("recruitment applicant update failed: %s", exc, exc_info=True)
            flash("지원자 정보를 저장하지 못했습니다. 학번 중복을 확인해주세요.", "danger")
    return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))


@app.route('/admin/recruitment-results/<uuid:campaign_id>/applicants/<uuid:applicant_id>/delete', methods=['POST'])
@login_required(role="admin")
def admin_recruitment_result_delete_applicant(campaign_id, applicant_id):
    supabase.table("recruitment_applicants").delete().eq("id", str(applicant_id)) \
        .eq("campaign_id", str(campaign_id)).execute()
    flash("지원자를 명단에서 삭제했습니다.", "success")
    return redirect(url_for("admin_recruitment_result_detail", campaign_id=campaign_id))





@app.route('/api/admin/members/<int:member_id>/set_status', methods=['POST'])
@login_required(role="admin")
def set_member_status(member_id):
    """관리자가 특정 멤버의 상태를 active / dormant / inactive 중 하나로 설정합니다."""
    data = request.json
    new_status = data.get('member_status')
    if new_status not in ('active', 'dormant', 'inactive'):
        return jsonify({"status": "error", "message": "유효하지 않은 상태입니다."}), 400
    try:
        if member_id == session.get('user_id') and new_status == 'inactive':
            return jsonify({"status": "error", "message": "자기 자신을 비활성화할 수 없습니다."}), 403
        supabase.table('members').update({
            'member_status': new_status,
            'is_active': new_status == 'active'  # 하위 호환성 유지
        }).eq('id', member_id).execute()
        return jsonify({"status": "success", "message": "멤버 상태가 변경되었습니다."})
    except Exception as e:
        app.logger.error(f"Error setting member status for {member_id}: {e}")
        return jsonify({"status": "error", "message": "상태 변경 중 오류 발생"}), 500


def _member_student_id_conflict(student_id, exclude_member_id=None):
    rows = supabase.table('members').select('id, student_id').execute().data or []
    return member_student_id_conflicts(rows, student_id, exclude_member_id)


@app.route('/api/admin/members/create', methods=['POST'])
@login_required(role="admin")
def create_member():
    """관리자가 새 멤버를 등록합니다."""
    data = request.json
    try:
        name = data.get('name', '').strip()
        if not name:
            return jsonify({"status": "error", "message": "이름은 필수입니다."}), 400
        requested_role = data.get('role', 'member')
        if requested_role not in ('member', 'officer', 'admin'):
            return jsonify({"status": "error", "message": "유효하지 않은 권한입니다."}), 400
        if requested_role == 'admin' and not _current_user_is_primary_admin():
            return jsonify({"status": "error", "message": "관리자 권한은 관리자만 지정할 수 있습니다."}), 403
        student_id = normalize_member_student_id(data.get('student_id'))
        if not valid_member_student_id(student_id):
            return jsonify({"status": "error", "message": "학번은 숫자 4~20자리로 입력해주세요."}), 400
        if student_id and _member_student_id_conflict(student_id):
            return jsonify({"status": "error", "message": "이미 다른 회원에게 등록된 학번입니다. 중복 학번 표시를 확인해주세요."}), 409
        insert_fields = {
            'name': name,
            'role': requested_role,
            'member_status': 'active',
            'is_active': True,
            'account_status': 'active'
        }
        # 빈 문자열은 NULL로 (email은 UNIQUE 제약이 있어 ''가 중복되면 insert가 실패함)
        for field in ('email', 'gender', 'department', 'student_id', 'recruiting_class'):
            val = data.get(field)
            if field == 'student_id':
                val = student_id
            if isinstance(val, str) and val.strip() == '':
                val = None
            insert_fields[field] = val
        supabase.table('members').insert(insert_fields).execute()
        return jsonify({"status": "success", "message": f"{name} 멤버가 추가되었습니다."})
    except Exception as e:
        app.logger.error(f"Error creating member: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/admin/members/<int:member_id>/edit', methods=['POST'])
@login_required(role="admin")
def edit_member(member_id):
    """관리자가 멤버 정보를 편집합니다."""
    data = request.json
    try:
        current_rows = supabase.table('members').select('role, student_id').eq('id', member_id).limit(1).execute().data or []
        if not current_rows:
            return jsonify({"status": "error", "message": "회원을 찾을 수 없습니다."}), 404
        current_role = current_rows[0].get('role') or 'member'
        if 'role' in data:
            requested_role = data.get('role')
            if requested_role not in ('member', 'officer', 'admin'):
                return jsonify({"status": "error", "message": "유효하지 않은 권한입니다."}), 400
            if requested_role != current_role and not _current_user_is_primary_admin():
                return jsonify({"status": "error", "message": "회원 권한은 관리자만 변경할 수 있습니다."}), 403
            if member_id == session.get('user_id') and current_role == 'admin' and requested_role != 'admin':
                return jsonify({"status": "error", "message": "자신의 관리자 권한은 직접 해제할 수 없습니다."}), 403
            if current_role == 'admin' and requested_role != 'admin':
                active_admins = supabase.table('members').select('id', count='exact') \
                    .eq('role', 'admin').eq('is_active', True).execute()
                if (active_admins.count or 0) <= 1:
                    return jsonify({"status": "error", "message": "최소 한 명의 활성 관리자가 필요합니다."}), 409
        update_fields = {}
        for field in ('name', 'email', 'gender', 'role', 'department', 'student_id', 'recruiting_class'):
            if field in data:
                val = data[field]
                if isinstance(val, str) and val.strip() == '':
                    update_fields[field] = None  # 빈 문자열은 NULL로
                else:
                    update_fields[field] = val
        if 'student_id' in update_fields:
            student_id = normalize_member_student_id(update_fields.get('student_id'))
            if not valid_member_student_id(student_id):
                return jsonify({"status": "error", "message": "학번은 숫자 4~20자리로 입력해주세요."}), 400
            current_student_id = normalize_member_student_id(current_rows[0].get('student_id'))
            if student_id and student_id != current_student_id \
                    and _member_student_id_conflict(student_id, exclude_member_id=member_id):
                return jsonify({"status": "error", "message": "이미 다른 회원에게 등록된 학번입니다. 중복 학번 표시를 확인해주세요."}), 409
            update_fields['student_id'] = student_id or None
        if not update_fields:
            return jsonify({"status": "error", "message": "변경할 내용이 없습니다."}), 400
        supabase.table('members').update(update_fields).eq('id', member_id).execute()
        return jsonify({"status": "success", "message": "멤버 정보가 수정되었습니다."})
    except Exception as e:
        app.logger.error(f"Error editing member {member_id}: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/admin/members/<int:member_id>/delete', methods=['POST'])
@login_required(role="admin")
def delete_member(member_id):
    """관리자가 멤버를 삭제합니다."""
    try:
        if member_id == session.get('user_id'):
            return jsonify({"status": "error", "message": "자기 자신을 삭제할 수 없습니다."}), 403
        supabase.table('members').delete().eq('id', member_id).execute()
        return jsonify({"status": "success", "message": "멤버가 삭제되었습니다."})
    except Exception as e:
        app.logger.error(f"Error deleting member {member_id}: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/admin/members/merge', methods=['POST'])
@login_required(role="admin")
def merge_members():
    """두 멤버를 하나로 합칩니다.
    예: '민수'(성 빠짐) → '김민수' 로 모든 활동 이력 이전 후 source 삭제.
    body: {source_id: int, target_id: int}
    """
    try:
        data = request.json or {}
        source_id = int(data.get('source_id'))
        target_id = int(data.get('target_id'))
        if source_id == target_id:
            return jsonify({"status": "error", "message": "같은 멤버입니다."}), 400
        if source_id == session.get('user_id'):
            return jsonify({"status": "error", "message": "자기 자신을 source 로 지정할 수 없습니다."}), 403

        # 두 멤버 모두 존재해야 함
        both = supabase.table('members').select('id, name, social_id, email').in_('id', [source_id, target_id]).execute().data or []
        if len(both) != 2:
            return jsonify({"status": "error", "message": "멤버를 찾을 수 없습니다."}), 404
        source_row = next(m for m in both if m['id'] == source_id)
        target_row = next(m for m in both if m['id'] == target_id)

        # 카카오 로그인 정보(social_id/email)를 target 으로 이전.
        # target 에 이미 값이 있으면 덮어쓰지 않음. source 의 값은 unique 충돌 방지를 위해 먼저 NULL 처리.
        login_transfer = {}
        if (source_row.get('social_id') or '').strip() and not (target_row.get('social_id') or '').strip():
            login_transfer['social_id'] = source_row['social_id']
        if (source_row.get('email') or '').strip() and not (target_row.get('email') or '').strip():
            login_transfer['email'] = source_row['email']
        if login_transfer:
            # source 에서 먼저 비워서 unique 충돌 방지
            supabase.table('members').update({
                k: None for k in login_transfer.keys()
            }).eq('id', source_id).execute()
            # target 으로 이전
            supabase.table('members').update(login_transfer).eq('id', target_id).execute()

        # (table, column, conflict_columns) — conflict_columns 가 있으면 unique 충돌 시 source 행 삭제
        moves = [
            ('attendance', 'user_id', ['user_id', 'meeting_date']),
            ('seminar_votes', 'member_id', ['session_id', 'member_id']),
            ('brick_session_members', 'member_id', ['session_id', 'member_id']),
            ('study_session_members', 'member_id', ['session_id', 'member_id']),
            ('special_event_attendees', 'member_id', ['event_id', 'member_id']),
        ]
        moved_summary = {}
        for table, col, conflict_cols in moves:
            try:
                source_rows = supabase.table(table).select(','.join(['id'] + conflict_cols)) \
                    .eq(col, source_id).execute().data or []
                if not source_rows:
                    moved_summary[table] = 0
                    continue
                # target 이 이미 갖고 있는 conflict_cols 조합 조회
                other_col = [c for c in conflict_cols if c != col][0]
                other_vals = list({r[other_col] for r in source_rows})
                target_existing_res = supabase.table(table).select(other_col) \
                    .eq(col, target_id).in_(other_col, other_vals).execute().data or []
                target_has = {r[other_col] for r in target_existing_res}

                to_delete_ids = [r['id'] for r in source_rows if r[other_col] in target_has]
                to_update_ids = [r['id'] for r in source_rows if r[other_col] not in target_has]

                if to_delete_ids:
                    supabase.table(table).delete().in_('id', to_delete_ids).execute()
                if to_update_ids:
                    supabase.table(table).update({col: target_id}).in_('id', to_update_ids).execute()
                moved_summary[table] = {'updated': len(to_update_ids), 'deleted_dup': len(to_delete_ids)}
            except Exception as e:
                app.logger.warning(f"merge_members move {table} 실패: {e}")
                moved_summary[table] = f"error: {e}"

        # set null 계열: special_events.created_by — source 가 만든 이벤트는 target 으로 이전
        try:
            supabase.table('special_events').update({'created_by': target_id}) \
                .eq('created_by', source_id).execute()
        except Exception as e:
            app.logger.warning(f"merge_members special_events.created_by 실패: {e}")

        # 마지막으로 source 삭제
        supabase.table('members').delete().eq('id', source_id).execute()

        return jsonify({
            "status": "success",
            "message": "멤버가 합쳐졌습니다.",
            "moved": moved_summary,
            "login_transferred": login_transfer,
        })
    except Exception as e:
        app.logger.error(f"merge_members error: {e}", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500



@app.route('/api/notifications')
@login_required(role="admin")
def get_notifications():
    """관리자에게 보여줄 승인 대기 중인 알림 목록을 반환합니다."""
    try:
        notifications = supabase.table('notifications') \
            .select('*') \
            .eq('status', 'pending') \
            .order('created_at', desc=True) \
            .execute().data
        return jsonify(notifications)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# app.py의 handle_notification 함수를 아래 코드로 교체

@app.route('/api/notifications/<int:notif_id>/handle', methods=['POST'])
@login_required(role="admin")
def handle_notification(notif_id):
    data = request.json
    action = data.get('action')  # 'approve' or 'deny'

    if action not in ['approve', 'deny']:
        return jsonify({"error": "Invalid action"}), 400

    try:
        # 1. 알림 정보를 먼저 조회
        notification_to_handle_res = supabase.table('notifications').select('related_member_id, type') \
            .eq('id', notif_id).single().execute()

        notification_to_handle = notification_to_handle_res.data
        if not notification_to_handle:
            return jsonify({"error": "Notification not found"}), 404

        member_id = notification_to_handle.get('related_member_id')
        notif_type = notification_to_handle.get('type')

        # 2. 알림의 상태를 'approved' 또는 'denied'로 업데이트
        supabase.table('notifications').update({
            'status': 'approved' if action == 'approve' else 'denied'
        }).eq('id', notif_id).execute()

        # 3. 알림 유형에 따라 후속 조치 실행
        if member_id:
            # 가입 또는 계정 연결 요청 처리
            if notif_type in ['new_user_request', 'account_link_request']:
                if action == 'approve':
                    supabase.table('members').update({'account_status': 'active'}).eq('id', member_id).execute()
                elif action == 'deny':
                    if notif_type == 'account_link_request':
                        # 연결 거절 시: 다시 시도할 수 있도록 social_id 해제 및 상태 원복
                        supabase.table('members').update({
                            'social_id': None, 
                            'account_status': 'active'
                        }).eq('id', member_id).execute()
                    elif notif_type == 'new_user_request':
                        # 신규 가입 거절 시: 임시로 생성된 pending 멤버 레코드 자체를 삭제하여 재가입 가능하도록 함
                        supabase.table('members').delete().eq('id', member_id).execute()

            # 불참 요청 처리
            elif notif_type == 'absence_request':
                next_monday = get_next_monday()

                # [핵심 수정] update_data를 정의하고, where 절을 명확하게 지정하여 업데이트
                if action == 'approve':
                    # 불참 승인: 불참 확정 및 상태 변경
                    update_data = {
                        'absence_request_status': 'approved',
                        'attending_seminar': False
                    }
                else:  # action == 'deny'
                    # 불참 반려: 요청 상태만 변경
                    update_data = {
                        'absence_request_status': 'denied'
                    }

                # 업데이트할 행을 명확하게 지정
                supabase.table('attendance').update(update_data) \
                    .eq('user_id', member_id) \
                    .eq('meeting_date', next_monday.isoformat()) \
                    .eq('absence_request_status', 'pending') \
                    .execute()

        return jsonify({"status": "success", "message": f"요청이 {action}되었습니다."})

    except Exception as e:
        app.logger.error(f"Error handling notification {notif_id}: {e}")
        return jsonify({"error": str(e)}), 500


# --- 4.2. 독서 모임 조 편성 (관리자 전용) ---


def _topic_facilitators_for_session(seminar_session, all_active_members):
    """Match a session's topic submitters to active members."""
    if not seminar_session:
        return None, set(), []

    topic_query = supabase.table('topic_events').select('*')
    if seminar_session.get('seminar_week_id'):
        topic_query = topic_query.eq('seminar_week_id', seminar_session['seminar_week_id'])
    else:
        topic_query = topic_query.eq('seminar_session_id', seminar_session['id'])
    topic_rows = topic_query.limit(1).execute().data or []
    if not topic_rows:
        return None, set(), []

    topic_event = topic_rows[0]
    submissions = supabase.table('topic_submissions').select('author_name, student_id') \
        .eq('event_id', topic_event['id']).execute().data or []
    by_student_id = {
        str(member.get('student_id')).strip(): member
        for member in all_active_members if member.get('student_id')
    }
    by_name = defaultdict(list)
    for member in all_active_members:
        by_name[(member.get('name') or '').strip()].append(member)

    matched_names = set()
    unmatched = []
    for submission in submissions:
        member = None
        student_id = str(submission.get('student_id') or '').strip()
        author_name = (submission.get('author_name') or '').strip()
        if student_id:
            member = by_student_id.get(student_id)
        if member is None and len(by_name.get(author_name, [])) == 1:
            member = by_name[author_name][0]
        if member:
            matched_names.add(member['name'])
        else:
            unmatched.append({'author_name': author_name, 'student_id': student_id})
    return topic_event, matched_names, unmatched


def _load_group_pair_restriction_rows():
    return supabase.table('group_pair_restrictions') \
        .select('id, member_a_id, member_b_id, note, created_at, updated_at') \
        .order('created_at').execute().data or []


def _restricted_name_pairs(member_rows=None):
    restrictions = _load_group_pair_restriction_rows()
    if member_rows is None:
        member_rows = supabase.table('members').select('id, name').execute().data or []
    return restricted_pairs_from_rows(restrictions, member_rows)


@app.route('/api/admin/group-pair-restrictions', methods=['POST'])
@primary_admin_required
def create_group_pair_restriction():
    data = request.get_json(silent=True) or {}
    try:
        member_ids = sorted((int(data.get('member_a_id')), int(data.get('member_b_id'))))
    except (TypeError, ValueError):
        return jsonify({'error': '두 회원을 모두 선택해주세요.'}), 400
    if member_ids[0] == member_ids[1]:
        return jsonify({'error': '서로 다른 두 회원을 선택해주세요.'}), 400

    active_rows = supabase.table('members').select('id').in_('id', member_ids) \
        .eq('is_active', True).execute().data or []
    if {row['id'] for row in active_rows} != set(member_ids):
        return jsonify({'error': '현재 활성 상태인 회원만 선택할 수 있습니다.'}), 400

    note = (data.get('note') or '').strip()
    if len(note) > 200:
        return jsonify({'error': '관리 메모는 200자 이내로 입력해주세요.'}), 400
    now_iso = datetime.now(timezone.utc).isoformat()
    payload = {
        'member_a_id': member_ids[0],
        'member_b_id': member_ids[1],
        'note': note or None,
        'created_by_member_id': session.get('user_id'),
        'updated_at': now_iso,
    }
    supabase.table('group_pair_restrictions').upsert(
        payload, on_conflict='member_a_id,member_b_id'
    ).execute()
    return jsonify({'status': 'ok'})


@app.route('/api/admin/group-pair-restrictions/<int:restriction_id>', methods=['DELETE'])
@primary_admin_required
def delete_group_pair_restriction(restriction_id):
    supabase.table('group_pair_restrictions').delete().eq('id', restriction_id).execute()
    return jsonify({'status': 'ok'})


@app.route('/making_team', methods=['GET'])
@login_required(role="admin")
def bookclub_index():
    app.logger.info("---/making_team 경로 함수 실행 시작---")
    seminar_session_id = (request.args.get('session_id') or '').strip() or None
    day_choice = (request.args.get('day') or 'mon').lower()
    if day_choice not in ('mon', 'thu', 'all'):
        day_choice = 'mon'

    mon_attendee_ids: set = set()
    thu_attendee_ids: set = set()
    mon_date_iso = ''
    thu_date_iso = ''
    selected_session = None
    selected_topic_event = None
    pre_checked_facilitator_names = set()
    unmatched_facilitators = []
    can_manage_pair_restrictions = False
    pair_restrictions = []

    try:
        all_active_members_res = supabase.table("members").select("id, name, student_id, department, gender") \
            .eq('is_active', True).order("name").execute()
        all_active_members = all_active_members_res.data or []
        for member in all_active_members:
            member['gender_code'] = normalize_gender(member.get('gender'))
        active_member_ids = {member['id'] for member in all_active_members}
        app.logger.info(f"[2] DB에서 가져온 전체 활성 멤버 수: {len(all_active_members)}명")

        if seminar_session_id:
            selected_session = supabase.table('seminar_sessions').select('*') \
                .eq('id', seminar_session_id).single().execute().data
            if not selected_session:
                raise ValueError('선택한 세미나 회차를 찾을 수 없습니다.')
            meeting_date = selected_session['meeting_date']
            day_choice = selected_session.get('day_type') or day_choice
            if day_choice == 'mon':
                mon_date_iso = meeting_date
                target_ids = mon_attendee_ids
            else:
                thu_date_iso = meeting_date
                target_ids = thu_attendee_ids

            participation_mode = selected_session.get('participation_mode') or 'legacy_explicit'
            if selected_session.get('planned_member_ids') is not None:
                target_ids |= {int(member_id) for member_id in selected_session['planned_member_ids']}
            elif participation_mode == 'absence_only':
                absence_rows = supabase.table('seminar_absences').select('member_id') \
                    .eq('session_id', seminar_session_id).is_('cancelled_at', 'null').execute().data or []
                absent_ids = {row['member_id'] for row in absence_rows}
                target_ids |= active_member_ids - absent_ids
            elif participation_mode == 'opt_in':
                vote_rows = supabase.table('seminar_votes').select('member_id') \
                    .eq('session_id', seminar_session_id).eq('attending', True).execute().data or []
                target_ids |= {row['member_id'] for row in vote_rows}
            else:
                attendance_rows = supabase.table('attendance').select('user_id') \
                    .eq('meeting_date', meeting_date).eq('attending_seminar', True).execute().data or []
                target_ids |= {row['user_id'] for row in attendance_rows}
                vote_rows = supabase.table('seminar_votes').select('member_id') \
                    .eq('session_id', seminar_session_id).eq('attending', True).execute().data or []
                target_ids |= {row['member_id'] for row in vote_rows}
            pre_checked_attendee_ids = set(target_ids)

            selected_topic_event, pre_checked_facilitator_names, unmatched_facilitators = \
                _topic_facilitators_for_session(selected_session, all_active_members)
        else:
            seminar_dates = get_next_seminar_dates()
            app.logger.info(f"[1] 다음 세미나 날짜: {[d.isoformat() for d in seminar_dates]}")
            for d in seminar_dates:
                if d.weekday() == 0:
                    mon_date_iso = d.isoformat()
                elif d.weekday() == 3:
                    thu_date_iso = d.isoformat()
            date_strs = [d.isoformat() for d in seminar_dates]
            attendance_res = supabase.table('attendance').select('user_id, meeting_date') \
                .in_('meeting_date', date_strs).eq('attending_seminar', True).execute()
            for row in (attendance_res.data or []):
                if row.get('meeting_date') == mon_date_iso:
                    mon_attendee_ids.add(row['user_id'])
                elif row.get('meeting_date') == thu_date_iso:
                    thu_attendee_ids.add(row['user_id'])
            sess_rows = supabase.table('seminar_sessions').select(
                'id, meeting_date, day_type, participation_mode, seminar_week_id, book_title, planned_member_ids'
            ) \
                .in_('meeting_date', date_strs).eq('is_active', True).execute().data or []
            for s in sess_rows:
                target_ids = mon_attendee_ids if s.get('day_type') == 'mon' else thu_attendee_ids
                mode = s.get('participation_mode') or 'legacy_explicit'
                if s.get('planned_member_ids') is not None:
                    target_ids.clear()
                    target_ids |= {int(member_id) for member_id in s['planned_member_ids']}
                elif mode == 'absence_only':
                    target_ids.clear()
                    absent = supabase.table('seminar_absences').select('member_id') \
                        .eq('session_id', s['id']).is_('cancelled_at', 'null').execute().data or []
                    target_ids |= active_member_ids - {row['member_id'] for row in absent}
                else:
                    if mode == 'opt_in':
                        target_ids.clear()
                    votes = supabase.table('seminar_votes').select('member_id') \
                        .eq('session_id', s['id']).eq('attending', True).execute().data or []
                    target_ids |= {row['member_id'] for row in votes}
            if day_choice == 'mon':
                pre_checked_attendee_ids = mon_attendee_ids
            elif day_choice == 'thu':
                pre_checked_attendee_ids = thu_attendee_ids
            else:
                pre_checked_attendee_ids = mon_attendee_ids | thu_attendee_ids

            # 조 편성 메뉴를 직접 열어도 선택한 요일의 회차와 발제문을 연결한다.
            # 월·목 합산은 단일 회차로 저장할 수 없으므로 회차 자동 연결을 하지 않는다.
            if day_choice in ('mon', 'thu'):
                linked_session = next(
                    (row for row in sess_rows if row.get('day_type') == day_choice), None
                )
                if linked_session:
                    seminar_session_id = linked_session['id']
                    selected_topic_event, pre_checked_facilitator_names, unmatched_facilitators = \
                        _topic_facilitators_for_session(linked_session, all_active_members)

        # Keep snapshotted attendees visible even if their member status changed.
        snapshot_missing_ids = set(pre_checked_attendee_ids) - {member['id'] for member in all_active_members}
        if snapshot_missing_ids:
            snapshot_members = supabase.table('members').select('id, name, student_id, department, gender') \
                .in_('id', list(snapshot_missing_ids)).execute().data or []
            for member in snapshot_members:
                member['gender_code'] = normalize_gender(member.get('gender'))
            all_active_members.extend(snapshot_members)
            all_active_members.sort(key=lambda member: member.get('name') or '')

        can_manage_pair_restrictions = _current_user_is_primary_admin()
        if can_manage_pair_restrictions:
            member_by_id = {member['id']: member for member in all_active_members}
            for row in _load_group_pair_restriction_rows():
                member_a = member_by_id.get(row.get('member_a_id'))
                member_b = member_by_id.get(row.get('member_b_id'))
                if member_a and member_b:
                    pair_restrictions.append({
                        **row,
                        'member_a': member_a,
                        'member_b': member_b,
                    })
        app.logger.info(f"[5] day={day_choice}, 미리 체크될 참석자 수: {len(pre_checked_attendee_ids)}명 (월 {len(mon_attendee_ids)} / 목 {len(thu_attendee_ids)})")

    except Exception as e:
        app.logger.error(f"!!! /making_team 경로에서 오류 발생: {e}", exc_info=True)
        flash(f"데이터를 불러오는 중 오류가 발생했습니다: {e}", "danger")
        all_active_members = []
        pre_checked_attendee_ids = set()

    return render_template(
        'bookclub_index.html',
        all_members=all_active_members,
        pre_checked_attendee_ids=pre_checked_attendee_ids,
        mon_attendee_ids=mon_attendee_ids,
        thu_attendee_ids=thu_attendee_ids,
        mon_date=mon_date_iso,
        thu_date=thu_date_iso,
        day_choice=day_choice,
        selected_session=selected_session,
        selected_topic_event=selected_topic_event,
        seminar_session_id=seminar_session_id,
        pre_checked_facilitator_names=pre_checked_facilitator_names,
        unmatched_facilitators=unmatched_facilitators,
        can_manage_pair_restrictions=can_manage_pair_restrictions,
        pair_restrictions=pair_restrictions,
    )


@app.route('/start_group_generation')
@login_required(role="admin")
def start_group_generation():
    # --- [수정] 로그 추가 ---
    app.logger.info("\n---/start_group_generation 경로 함수 실행 시작---")
    present_names = request.args.getlist('present')
    facilitator_names = request.args.getlist('facilitators')
    group_count_str = request.args.get('group_count')
    group_names_str = request.args.get('group_names', '')
    seminar_session_id = (request.args.get('seminar_session_id') or '').strip() or None
    meeting_date = (request.args.get('meeting_date') or '').strip() or None
    session_book_title = None
    if seminar_session_id:
        selected_session = supabase.table('seminar_sessions').select('meeting_date, book_title') \
            .eq('id', seminar_session_id).single().execute().data
        if not selected_session:
            return jsonify({'error': '선택한 세미나 회차를 찾을 수 없습니다.'}), 404
        meeting_date = selected_session['meeting_date']
        session_book_title = selected_session.get('book_title')

    # 방어 필터: 참석 명단에 없는 발제자는 조 편성에서 제외한다.
    # (프런트 검증을 우회하거나 잘못된 파라미터가 들어와도 안전하도록.)
    present_set = set(present_names)
    dropped_facs = [n for n in facilitator_names if n not in present_set]
    if dropped_facs:
        app.logger.info(f"[1.5] 참석 명단에 없어 제외된 발제자: {dropped_facs}")
        facilitator_names = [n for n in facilitator_names if n in present_set]

    app.logger.info(f"[1] 전달받은 참석자 명단 (총 {len(present_names)}명): {present_names}")
    app.logger.info(f"[2] 전달받은 발제자 명단: {facilitator_names}")
    app.logger.info(f"[3] 전달받은 그룹 수: '{group_count_str}'")
    # --- [수정 끝] ---

    manual_entry_url = url_for('manual_entry')

    def generate_events(manual_url):
        cancel_event = None
        try:
            app.logger.info("[4] DB에서 전체 회원 및 히스토리 데이터 로드 시작")
            members_res = supabase.table("members").select("*").order("name").execute().data
            if len(set(present_names)) != len(present_names):
                raise ValueError('참석자 이름이 중복되어 있습니다. 동명이인 또는 중복 선택을 확인해주세요.')
            _validate_group_member_names([present_names])
            members_df = pd.DataFrame(members_res)
            effective_history_rows = _effective_group_history_rows()
            app.logger.info(f"[5] 데이터 로드 완료: 회원 {len(members_df)}명, 히스토리 {len(effective_history_rows)}건")
            restricted_pairs = _restricted_name_pairs(members_res)
            app.logger.info("[5.5] 비공개 편성 제한 %d건 적용", len(restricted_pairs))

            group_count_override = None
            if group_count_str and group_count_str.isdigit():
                group_count_override = int(group_count_str)

            # 저장된 계획표가 아니라 실제 참석 기준으로 즉시 계산한다.
            # 회차에 등록된 미연락 불참자는 해당 날짜의 만남에서 자동 제외된다.
            calculated_matrix = _matrix_rows_from_history(effective_history_rows)
            co_matrix = {key: row['count'] for key, row in calculated_matrix.items()}
            app.logger.info(f"[6] co_matrix {len(co_matrix)}개 항목 로드 완료")

            yield f"event: progress\ndata: {json.dumps({'progress': 10})}\n\n"

            app.logger.info("[7] '종합 최적화(단일)' CP-SAT 알고리즘 실행 시작")

            # 솔버를 별도 스레드에서 돌리고, 진행률을 큐를 통해 실시간 스트리밍
            import threading, queue
            progress_queue = queue.Queue()
            cancel_event = threading.Event()

            def progress_callback(pct):
                try:
                    progress_queue.put(('progress', pct))
                except Exception:
                    pass

            solver_result = {'solutions': None, 'error': None}

            def run_solver():
                try:
                    solver_result['solutions'] = run_cp_grouping(
                        members_df, co_matrix, present_names, facilitator_names,
                        optimize_for='combined', top_n=12,
                        group_count_override=group_count_override,
                        progress_callback=progress_callback,
                        restricted_pairs=restricted_pairs,
                        cancel_event=cancel_event,
                    )
                except Exception as ex:
                    solver_result['error'] = ex
                finally:
                    progress_queue.put(('done', None))

            t = threading.Thread(target=run_solver, daemon=True)
            t.start()

            # 큐를 폴링하면서 진행률을 실시간 yield
            last_sent_pct = 10
            while True:
                try:
                    kind, payload = progress_queue.get(timeout=5)
                except queue.Empty:
                    # heartbeat (SSE 연결 유지)
                    yield ": keep-alive\n\n"
                    continue
                if kind == 'progress':
                    if payload > last_sent_pct:
                        last_sent_pct = payload
                        yield f"event: progress\ndata: {json.dumps({'progress': payload})}\n\n"
                elif kind == 'done':
                    break

            t.join()
            if solver_result['error']:
                raise solver_result['error']
            combined_solutions = solver_result['solutions'] or []
            app.logger.info(f"[8] '종합 최적화' 완료, {len(combined_solutions)}개")

            yield f"event: progress\ndata: {json.dumps({'progress': 90})}\n\n"

            # 결과 페이지에서 'M'/'W'로 비교하므로 정규화된 값으로 전달
            member_genders = {
                row['name']: normalize_gender(row.get('gender'))
                for row in members_df.to_dict(orient='records')
            }

            meeting_history = _meeting_details_from_history(effective_history_rows, before_date=meeting_date)

            app.logger.info("[10] 최종 결과 페이지(HTML) 렌더링 시작")
            with app.app_context():
                group_names = [name.strip() for name in group_names_str.split(',') if name.strip()]
                final_html = render_template(
                    'bookclub_ga_results.html',
                    combined_solutions=combined_solutions,
                    present=present_names,
                    facilitators=facilitator_names,
                    group_names=group_names,
                    meeting_history=meeting_history,
                    member_genders=member_genders,
                    manual_entry_url=manual_url,
                    seminar_session_id=seminar_session_id,
                    meeting_date=meeting_date,
                    book_title=session_book_title,
                )
                complete_data = json.dumps({'html': final_html})
                yield f"event: complete\ndata: {complete_data}\n\n"
            app.logger.info("[11] 성공적으로 최종 HTML을 브라우저로 전송 완료")

        except GeneratorExit:
            if cancel_event is not None:
                cancel_event.set()
            app.logger.info("Group generation stream closed by the client")
            raise
        except Exception as e:
            app.logger.error(f"!!! 조 편성 중 심각한 오류 발생: {e}", exc_info=True)
            error_data = json.dumps({'error': str(e)})
            yield f"event: error\ndata: {error_data}\n\n"
        finally:
            if cancel_event is not None:
                cancel_event.set()

    # Response가 generator를 순회하는 시점에는 원래 view 함수가 이미 반환된 뒤다.
    # 요청 컨텍스트를 스트림 수명 동안 유지해야 결과 템플릿의 header/sidebar가
    # request와 session을 안전하게 참조할 수 있다.
    return Response(
        stream_with_context(generate_events(manual_entry_url)),
        mimetype='text/event-stream',
    )



def run_cp_grouping(members_df, co_matrix, attendee_names, presenter_names,
                    optimize_for='gender', top_n=10, group_count_override=None,
                    progress_callback=None, restricted_pairs=None, cancel_event=None):
    """
    OR-Tools CP-SAT 기반 조 편성 알고리즘.
    optimize_for: 'gender' (성비우선) or 'new_face' (새만남우선)
    top_n: 반환할 다양한 조합 수
    """
    app.logger.info(f"[CP-SAT] 시작: optimize_for={optimize_for}, top_n={top_n}, attendees={len(attendee_names)}")

    names = list(attendee_names)
    n = len(names)

    if n < 3:
        app.logger.warning("[CP-SAT] 참석자 3명 미만, 중단")
        return []

    # 멤버 정보 딕셔너리
    name_to_info = {}
    for _, row in members_df.iterrows():
        name_to_info[row['name']] = row.to_dict()

    # 그룹 수 결정
    MIN_GROUP_SIZE = 4
    if group_count_override and group_count_override > 0:
        num_groups = group_count_override
    else:
        q, r = divmod(n, 4)
        num_groups = q if r == 0 else q + 1
    # 모든 그룹이 최소 MIN_GROUP_SIZE명 이상 가질 수 있도록 그룹 수 상한선 적용
    num_groups = max(1, min(num_groups, n // MIN_GROUP_SIZE))

    min_size = MIN_GROUP_SIZE
    max_size = math.ceil(n / num_groups)  # +1 제거: 일부 조에 여유 생갈 경우 3인조 발생 방지

    app.logger.info(f"[CP-SAT] 그룹={num_groups}, 크기={min_size}~{max_size}")

    # 성별 정규화: 미상은 별도 카테고리로 두어 자동으로 여성에 합산되지 않도록 함.
    is_male_arr = []   # 1 if 남성 else 0
    is_female_arr = [] # 1 if 여성 else 0
    for name in names:
        info = name_to_info.get(name, {})
        norm = normalize_gender(info.get('gender'))
        is_male_arr.append(1 if norm == 'M' else 0)
        is_female_arr.append(1 if norm == 'W' else 0)
    # 하위호환을 위해 genders도 유지 (기존 변수 참조 자리)
    genders = is_male_arr

    # 발제자 인덱스
    presenter_set = set(presenter_names)
    presenter_indices = [i for i, nm in enumerate(names) if nm in presenter_set]

    name_to_index = {name: index for index, name in enumerate(names)}
    restricted_index_pairs = []
    for first_name, second_name in (restricted_pairs or set()):
        if first_name in name_to_index and second_name in name_to_index:
            restricted_index_pairs.append((name_to_index[first_name], name_to_index[second_name]))

    # 쌍별 만남 횟수
    def get_pair_count(a, b):
        key = _canonical_pair_key(a, b)
        return co_matrix.get(key, 0)

    # objective 계산을 위한 쌍 사전
    pair_counts = {}
    for i in range(n):
        for j in range(i + 1, n):
            pair_counts[(i, j)] = get_pair_count(names[i], names[j])

    SCALE = 1000
    results = []
    # 金지 조합: 이미 찾은 조합들 (각각 frozenset of frozensets)
    found_groupings = []

    for attempt in range(top_n * 2): # 최대 탐색 횟수 조정
        if len(results) >= top_n:
            break

        if cancel_event is not None and cancel_event.is_set():
            app.logger.info("[CP-SAT] client cancellation requested")
            break

        model = cp_model.CpModel()

        # x[i][g]: 멤버 i가 그룹 g에 소속
        x = [[model.NewBoolVar(f'x_{i}_{g}') for g in range(num_groups)] for i in range(n)]

        # --- [대칭성 파괴(Symmetry Breaking)] ---
        # 첫 번째 사람을 무조건 그룹 0에 넣음으로써 불필요한 자리바꿈 탐색 공간을 기하급수적으로 줄임
        if n > 0 and num_groups > 0:
            model.Add(x[0][0] == 1)

        # 각 멤버는 정확히 하나의 그룹
        for i in range(n):
            model.AddExactlyOne(x[i][g] for g in range(num_groups))

        # 그룹 크기 제약
        for g in range(num_groups):
            sz = sum(x[i][g] for i in range(n))
            model.Add(sz >= min_size)
            model.Add(sz <= max_size)

        # 비공개 관리자 제한: 지정된 두 사람은 어떤 추천안에서도 같은 조가 될 수 없다.
        for left_index, right_index in restricted_index_pairs:
            for g in range(num_groups):
                model.Add(x[left_index][g] + x[right_index][g] <= 1)

        # 발제자 분산: 그룹 수 >= 발제자 수일 때 각 그룹에 1명씩
        if len(presenter_indices) <= num_groups:
            for pi in range(len(presenter_indices)):
                for pj in range(pi + 1, len(presenter_indices)):
                    for g in range(num_groups):
                        model.Add(x[presenter_indices[pi]][g] + x[presenter_indices[pj]][g] <= 1)

        # 이전에 찾은 조합 금지: 동일한 pair grouping을 피하기 위해
        for attempt_idx, found_pairs in enumerate(found_groupings):
            if not found_pairs:
                continue
            same_g_vars = []
            for (i, j) in found_pairs:
                for g in range(num_groups):
                    b = model.NewBoolVar(f'f_{attempt_idx}_{i}_{j}_{g}')
                    # x[i][g] == 1 and x[j][g] == 1 이면 b >= 1 이어야 함
                    model.Add(x[i][g] + x[j][g] - 1 <= b)
                    same_g_vars.append(b)
            # 이전에 같은 그룹이었던 쌍들 중 무조건 하나 이상은 이번엔 다른 그룹이어야 함
            if same_g_vars:
                model.Add(sum(same_g_vars) <= len(found_pairs) - 1)

        # --- 목적함수 ---
        obj = []

        # 성비 점수: 각 그룹 내 성별 불균형(|남-여|)을 최소화.
        # 성별 미상자는 어느 쪽으로도 카운트하지 않아 결과 왜곡을 막음.
        for g in range(num_groups):
            males = model.NewIntVar(0, n, f'm_{g}')
            model.Add(males == sum(is_male_arr[i] * x[i][g] for i in range(n)))
            females = model.NewIntVar(0, n, f'f_{g}')
            model.Add(females == sum(is_female_arr[i] * x[i][g] for i in range(n)))
            diff = model.NewIntVar(-n, n, f'd_{g}')
            model.Add(diff == males - females)
            abs_diff = model.NewIntVar(0, n, f'ad_{g}')
            model.AddAbsEquality(abs_diff, diff)
            gender_w = 40 if optimize_for in ['gender', 'combined'] else 6
            obj.append(abs_diff * (-SCALE * gender_w))

        # 새만남 점수 (최적화 변환): 
        # 기존: 모든 쌍에 대해 보너스를 주는 방식 (O(N^2) 변수 생성)
        # 변경: 이미 만난 1~2단계 쌍에 대해서만 '페널티(loss)'를 부과 (O(E) 90% 이상 변수 축소)
        new_face_w = 10 if optimize_for in ['new_face', 'combined'] else 6
        max_nf_val = SCALE * new_face_w
        for (i, j), cnt in pair_counts.items():
            if cnt > 0:
                nf_val = max_nf_val // (cnt + 1)
                loss = max_nf_val - nf_val
                if loss > 0:
                    for g in range(num_groups):
                        bv = model.NewBoolVar(f's_{attempt}_{i}_{j}_{g}')
                        model.Add(x[i][g] + x[j][g] - 1 <= bv)
                        obj.append(bv * (-loss))

        model.Maximize(sum(obj))

        solver = cp_model.CpSolver()
        # 탐색 구조를 대폭 줄였으므로 1.5초만에 최적해/실현가능 영역에 도달
        solver.parameters.max_time_in_seconds = 1.5
        solver.parameters.num_workers = 4
        solver.parameters.random_seed = attempt * 13 + 7

        status = solver.Solve(model)

        if cancel_event is not None and cancel_event.is_set():
            app.logger.info("[CP-SAT] stopping after active solve due to cancellation")
            break
        if status not in (cp_model.OPTIMAL, cp_model.FEASIBLE):
            app.logger.warning(f"[CP-SAT] attempt {attempt}: 실패 status={status}")
            continue

        # 결과 추출
        assignment = {}
        for i in range(n):
            for g in range(num_groups):
                if solver.Value(x[i][g]) == 1:
                    assignment[names[i]] = g
                    break

        groups_dict = {g: [] for g in range(num_groups)}
        for name, g in assignment.items():
            groups_dict[g].append(name)
        formatted_groups = [groups_dict[g] for g in sorted(groups_dict) if groups_dict[g]]

        # 중복 체크
        frozen = frozenset(frozenset(grp) for grp in formatted_groups)
        if frozen in [frozenset(frozenset(grp) for grp in r['groups']) for r in results]:
            app.logger.info(f"[CP-SAT] attempt {attempt}: 중복, 스킵")
            continue

        # 점수 계산 (표시용 — combined 가중치 40:10 적용).
        # 성별 미상자는 분모/분자에서 모두 제외해 점수 왜곡 방지.
        gender_score = 0.0
        for grp in formatted_groups:
            m = sum(1 for nm in grp if normalize_gender(name_to_info.get(nm, {}).get('gender')) == 'M')
            f = sum(1 for nm in grp if normalize_gender(name_to_info.get(nm, {}).get('gender')) == 'W')
            if m > 0 and f > 0:
                gender_score += min(m, f) / max(m, f)

        new_face_score = 0.0
        for (i, j), cnt in pair_counts.items():
            if assignment.get(names[i]) == assignment.get(names[j]):
                new_face_score += 1.0 / (cnt + 1)

        total_score = (gender_score * 40 + new_face_score * 10)

        results.append({
            'score': f"{total_score:.2f}",
            'details': [f"{gender_score:.2f}", f"{new_face_score:.2f}", "0.00", "0.00"],
            'groups': formatted_groups
        })
        app.logger.info(f"[CP-SAT] attempt {attempt}: 성공 {len(results)}/{top_n}, score={total_score:.2f}")

        # 현재 진행률을 progress_callback으로 통보 (10~85% 범위를 results 수에 청해 비례 배분)
        if progress_callback:
            pct = min(85, 12 + int((len(results) / top_n) * 73))
            progress_callback(pct)

        # 이 조합의 동일-그룹 쌍을 금지 목록에 추가
        same_pairs = set()
        for i in range(n):
            for j in range(i + 1, n):
                if assignment.get(names[i]) == assignment.get(names[j]):
                    same_pairs.add((i, j))
        found_groupings.append(same_pairs)

    app.logger.info(f"[CP-SAT] 완료: {len(results)}개 반환")
    return results



#todo 미리보기에서도 * 거르기


def _validate_groups_against_restrictions(groups):
    conflicts = find_restriction_conflicts(groups, _restricted_name_pairs())
    if conflicts:
        app.logger.warning("조 편성 제한 조건 충돌 %d건 차단", len(conflicts))
        return False
    return True


def _validate_group_member_names(groups):
    """Name-keyed legacy records cannot safely identify duplicate member names."""
    names = {name for group in groups for name in group}
    if not names:
        return
    rows = supabase.table('members').select('id, name').in_('name', list(names)).execute().data or []
    counts = defaultdict(int)
    for row in rows:
        counts[row.get('name')] += 1
    if any(count > 1 for count in counts.values()):
        raise ValueError('동명이인 회원이 포함되어 이름만으로 구분할 수 없습니다. 회원 정보를 확인한 뒤 편성해주세요.')


def _actual_member_ids_from_names(names):
    names = list(dict.fromkeys(name.strip() for name in names if isinstance(name, str) and name.strip()))
    if not names:
        return []
    rows = supabase.table('members').select('id, name').in_('name', names).execute().data or []
    by_name = defaultdict(list)
    for member in rows:
        by_name[member.get('name')].append(member['id'])
    if any(len(by_name.get(name, [])) != 1 for name in names):
        raise ValueError('기존 참석 명단에 회원 정보가 없거나 동명이인이 있어 실제 출석을 확정할 수 없습니다. 회원 정보를 먼저 확인해주세요.')
    return [by_name[name][0] for name in names]


@app.route('/api/bookclub/validate-groups', methods=['POST'])
@login_required(role="admin")
def validate_bookclub_groups():
    data = request.get_json(silent=True) or {}
    groups = data.get('groups') or []
    try:
        groups, _ = normalize_group_editor_payload(groups)
        _validate_group_member_names(groups)
    except ValueError as exc:
        return jsonify({'valid': False, 'message': str(exc)}), 400
    if not _validate_groups_against_restrictions(groups):
        return jsonify({
            'valid': False,
            'message': '관리자가 지정한 비공개 편성 제한 조건과 충돌합니다. 명단을 다시 편성해주세요.',
        }), 409
    return jsonify({'valid': True})


@app.route('/api/bookclub/save', methods=['POST'])
@login_required(role="admin")
def bookclub_save():
    data = request.get_json(silent=True) or {}
    try:
        groups, editor_state = normalize_group_editor_payload(data.get('groups'), data.get('group_editor_state'))
    except ValueError as exc:
        return jsonify({'status': 'error', 'message': str(exc)}), 400
    # 헬퍼 함수를 호출하여 저장 로직 실행
    result = save_group_record_to_db(
        data.get("date"), [name for group in groups for name in group], data.get('facilitators') or [], groups,
        book_title=data.get('book_title'), genre=data.get('genre'),
        seminar_session_id=data.get('seminar_session_id'), group_editor_state=editor_state
    )

    if result["status"] == "ok":
        return jsonify(result)
    else:
        return jsonify(result), 500


def _chunks(items, size=100):
    seq = list(items)
    for idx in range(0, len(seq), size):
        yield seq[idx:idx + size]


def _effective_group_history_rows():
    """Only confirmed actual attendees contribute to past-meeting statistics."""
    history_rows = supabase.table('history').select(
        'id, date, groups, present, seminar_session_id, attendance_confirmed_at, actual_member_ids'
    ).execute().data or []
    today = datetime.now(KST).date().isoformat()
    history_rows = [row for row in history_rows if str(row.get('date') or '') <= today]
    session_ids = {row['seminar_session_id'] for row in history_rows if row.get('seminar_session_id')}
    sessions, no_show_rows = [], []
    if session_ids:
        sessions = supabase.table('seminar_sessions').select('id, actual_member_ids, attendance_confirmed_at') \
            .in_('id', list(session_ids)).execute().data or []
        no_show_rows = supabase.table('seminar_no_shows').select('session_id, member_id') \
            .in_('session_id', list(session_ids)).is_('cancelled_at', 'null').execute().data or []
    session_by_id = {row['id']: row for row in sessions}
    member_rows = supabase.table('members').select('id, name').execute().data or []
    names_to_ids = defaultdict(list)
    member_name_by_id = {}
    for member in member_rows:
        names_to_ids[member.get('name')].append(member['id'])
        member_name_by_id[member['id']] = member.get('name')
    no_shows_by_session = defaultdict(set)
    for absence in no_show_rows:
        name = member_name_by_id.get(absence['member_id'])
        if name:
            no_shows_by_session[absence['session_id']].add(name)
    effective = []
    for row in history_rows:
        linked = session_by_id.get(row.get('seminar_session_id')) or {}
        # Attendance may be confirmed before the seating chart is saved/replaced.
        if linked.get('attendance_confirmed_at') and linked.get('actual_member_ids') is not None:
            actual_ids = set(linked['actual_member_ids'])
        elif row.get('attendance_confirmed_at'):
            actual_ids = set(row['actual_member_ids']) if row.get('actual_member_ids') is not None else None
        else:
            continue
        excluded = set(no_shows_by_session.get(row.get('seminar_session_id'), set()))
        for group in row.get('groups') or []:
            for name in group if isinstance(group, list) else []:
                ids = names_to_ids.get(name) or []
                if actual_ids is not None and (len(ids) != 1 or ids[0] not in actual_ids):
                    excluded.add(name)
                elif actual_ids is None and row.get('present') is not None and name not in row['present']:
                    excluded.add(name)
        effective.append({**row, 'excluded_names': sorted(excluded)})
    return effective


def _rebuild_matrix_for_session(session_id):
    linked_rows = supabase.table('history').select('groups') \
        .eq('seminar_session_id', session_id).execute().data or []
    affected_keys = set()
    for row in linked_rows:
        affected_keys |= _pair_keys_from_groups(row.get('groups') or [])
    if affected_keys:
        return rebuild_co_matrix(affected_keys)
    return {'history_count': 0, 'pair_count': 0, 'removed_count': 0, 'scope': 'affected'}


def rebuild_co_matrix(pair_keys=None):
    """history를 원본으로 만남 매트릭스를 재계산한다.

    pair_keys가 주어지면 해당 쌍만 갱신하고, None이면 전체를 복구한다.
    """
    target_keys = set(pair_keys) if pair_keys is not None else None
    history_rows = _effective_group_history_rows()
    calculated = _matrix_rows_from_history(history_rows, target_keys)

    if target_keys is None:
        existing_rows = supabase.table('bookclub_co_matrix').select('pair_key').execute().data or []
        existing_keys = {row['pair_key'] for row in existing_rows}
        target_keys = existing_keys | set(calculated)

    stale_keys = target_keys - set(calculated)
    for batch in _chunks(stale_keys):
        supabase.table('bookclub_co_matrix').delete().in_('pair_key', batch).execute()
    for batch in _chunks(calculated.values()):
        supabase.table('bookclub_co_matrix').upsert(batch, on_conflict='pair_key').execute()

    return {
        'history_count': len(history_rows),
        'pair_count': len(calculated),
        'removed_count': len(stale_keys),
        'scope': 'full' if pair_keys is None else 'affected',
    }


# [신규] 조 편성 기록을 DB에 저장하는 헬퍼 함수
def save_group_record_to_db(date, present, facilitators, groups, book_title=None, genre=None,
                            seminar_session_id=None, group_editor_state=None):
    """주어진 데이터로 조 편성 기록과 만남 횟수 매트릭스를 DB에 저장/업데이트합니다."""
    try:
        groups, group_editor_state = normalize_group_editor_payload(groups, group_editor_state)
        _validate_group_member_names(groups)
        present = [name for group in groups for name in group]
        if not isinstance(facilitators, list):
            raise ValueError('발제자 명단 형식이 올바르지 않습니다.')
        facilitators = list(dict.fromkeys(name for name in facilitators if name in present))
        if not _validate_groups_against_restrictions(groups):
            raise ValueError('관리자가 지정한 비공개 편성 제한 조건과 충돌합니다. 명단을 다시 편성해주세요.')

        # 1. history 테이블에 기록 저장
        if seminar_session_id:
            sess_res = supabase.table('seminar_sessions').select('id, meeting_date, book_title') \
                .eq('id', seminar_session_id).single().execute().data
            if not sess_res:
                raise ValueError('연결할 세미나 회차를 찾을 수 없습니다.')
            date = sess_res['meeting_date']
            if not book_title:
                book_title = sess_res.get('book_title')

        record = {"date": date, "present": present, "facilitators": facilitators, "groups": groups,
                  "group_editor_state": group_editor_state}
        if seminar_session_id:
            record['seminar_session_id'] = seminar_session_id
        if book_title:
            record["book_title"] = book_title.strip()
        if genre:
            record["genre"] = genre.strip()
        insert_res = supabase.table("history").insert(record).execute()
        history_id = insert_res.data[0]['id'] if insert_res.data else None

        # 2. history를 원본으로 해당 사람 쌍을 다시 계산한다.
        affected_keys = _pair_keys_from_groups(groups)
        if affected_keys:
            rebuild_co_matrix(affected_keys)

        return {"status": "ok", "history_id": history_id}
    except Exception as e:
        app.logger.error(f"Error saving group record: {e}")
        return {"status": "error", "message": str(e)}

@app.route('/manual_entry')
@login_required(role="admin")
def manual_entry():
    """수동으로 조 편성을 입력하는 페이지를 렌더링합니다."""
    try:
        members_res = supabase.table("members").select("name").eq('is_active', True).order("name").execute().data
        all_members = [m['name'] for m in members_res]
    except Exception as e:
        flash(f"회원 정보를 불러오는 중 오류 발생: {e}", "danger")
        all_members = []
    # 회차 관리 페이지에서 "+ 수동 추가" 클릭 시 prefill (?date=, ?book_title=)
    prefill = {
        'date': (request.args.get('date') or '').strip(),
        'book_title': (request.args.get('book_title') or '').strip(),
    }
    try:
        genres = _load_genres()
    except Exception:
        genres = []
    return render_template('manual_entry.html', all_members=all_members,
                           prefill=prefill, genres=genres)


@app.route('/save_manual_groups', methods=['POST'])
@login_required(role="admin")
def save_manual_groups():
    try:
        form_data = request.form
        meeting_date = form_data.get('meeting_date')
        book_title = (form_data.get('book_title') or '').strip()
        genre = (form_data.get('genre') or '').strip()

        groups = []
        present_members_set = set()
        facilitator_members_set = set()  # [수정] 발제자 목록을 저장할 Set

        for i in range(1, 16):
            group_text = form_data.get(f'group_{i}')
            if group_text:
                member_names_raw = re.split(r'[,;\s\n]+', group_text)

                cleaned_group = []
                for name_raw in member_names_raw:
                    name = name_raw.strip()
                    if not name:
                        continue

                    # [수정] 이름 뒤에 '*'가 있는지 확인
                    if name.endswith('*'):
                        clean_name = name[:-1]  # '*'를 제거한 순수 이름
                        facilitator_members_set.add(clean_name)
                        cleaned_group.append(clean_name)
                    else:
                        cleaned_group.append(name)

                if cleaned_group:
                    groups.append(cleaned_group)
                    present_members_set.update(cleaned_group)

        present_members = sorted(list(present_members_set))
        facilitator_members = sorted(list(facilitator_members_set))  # [수정] 발제자 목록 정렬

        if not all([meeting_date, present_members, groups]):
            flash("날짜와 최소 1명 이상의 그룹 멤버를 모두 입력해야 합니다.", "danger")
            return redirect(url_for('manual_entry'))

        # [수정] 발제자/도서/장르 정보도 함께 DB에 저장
        result = save_group_record_to_db(meeting_date, present_members, facilitator_members, groups,
                                         book_title=book_title, genre=genre)

        if result["status"] == "ok":
            flash("수동 조 편성 기록이 성공적으로 저장되었습니다. 이 페이지에서 바로 수정할 수 있습니다.", "success")
            if result.get("history_id"):
                return redirect(url_for('records_seminar_detail', history_id=result["history_id"]))
            return redirect(url_for('records_seminars'))
        else:
            flash(f"저장 중 오류 발생: {result['message']}", "danger")
            return redirect(url_for('manual_entry'))

    except Exception as e:
        flash(f"처리 중 예외 발생: {e}", "danger")
        return redirect(url_for('manual_entry'))

@app.route('/api/bookclub/history', methods=['GET'])
@login_required(role="admin")
def bookclub_api_get_history():
    response = supabase.table("history").select("*").order("date", desc=True).execute()
    return jsonify(response.data)


@app.route('/api/bookclub/history/delete', methods=['POST'])
@login_required(role="admin")
def bookclub_api_delete_history():
    """기록을 삭제한 뒤 영향받은 사람 쌍을 남은 전체 이력으로 재계산한다."""
    record_id = request.json.get("id")
    if not record_id:
        return jsonify({"status": "error", "message": "record id required"}), 400
    try:
        # 1. 삭제할 기록 조회
        del_res = supabase.table("history").select("groups, date").eq("id", record_id).execute()
        if not del_res.data:
            return jsonify({"status": "error", "message": "Record not found"}), 404
        deleted_record = del_res.data[0]

        # 2. 실제 삭제
        supabase.table("history").delete().eq("id", record_id).execute()

        deleted_groups = deleted_record.get("groups", []) or []
        affected_keys = _pair_keys_from_groups(deleted_groups)
        if affected_keys:
            rebuild_co_matrix(affected_keys)

        return jsonify({"status": "ok"})
    except Exception as e:
        app.logger.error(f"Error deleting history: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500
# </editor-fold>
# ==============================================================================



#=== 위키 관련 모음

@app.route('/docs/<doc_title>')
def view_document(doc_title):
    return redirect(NOTION_PUBLIC_WIKI_URL, code=302)
    """
    데이터베이스에서 문서를 찾아 제목과 내용을 보여주는 페이지.
    """
    try:
        doc_res = supabase.table('documents').select('*').eq('title', doc_title).single().execute()
        document = doc_res.data

        if not document:
            flash(f"'{doc_title}' 문서를 찾을 수 없습니다.", "warning")
            return render_template('doc_not_found.html', title=doc_title), 404

        rendered_content = wiki_parser(document.get('content', ''))

        return render_template('doc_view.html', doc=document, content=rendered_content)

    except Exception as e:
        app.logger.error(f"Error viewing document '{doc_title}': {e}")
        flash("문서를 불러오는 중 오류가 발생했습니다.", "danger")
        return redirect(url_for('main_index'))


# 1. 문서 생성 페이지를 보여주는 라우트
@app.route('/docs/create')
@login_required(role="ANY")
def create_document_page():
    return redirect(NOTION_PUBLIC_WIKI_URL, code=302)


# 2. 문서 생성 요청을 처리하는 API 라우트
@app.route('/api/docs/create', methods=['POST'])
@login_required(role="ANY")
def handle_create_document():
    return jsonify({
        "status": "error", "message": "앱 위키 편집은 종료되었습니다.",
        "wiki_url": NOTION_PUBLIC_WIKI_URL,
    }), 410
    data = request.json
    title = data.get('title')
    content = data.get('content')
    author_id = session.get('user_id')  # 세션에서 현재 로그인한 사용자의 ID를 가져옴

    if not title or not content:
        return jsonify({"status": "error", "message": "제목과 내용을 모두 입력해야 합니다."}), 400

    try:
        # DB에 새로운 문서 삽입
        supabase.table('documents').insert({
            'title': title,
            'content': content,
            'author_id': author_id
        }).execute()

        # 성공 시, 새로 만들어진 문서 페이지로 바로 이동할 수 있도록 URL 반환
        return jsonify({"status": "success", "message": "문서가 성공적으로 생성되었습니다.", "doc_title": title})

    except Exception as e:
        # Supabase에서 title UNIQUE 제약 조건 위반 시 특정 에러 코드를 반환합니다.
        if '23505' in str(e):  # UNIQUE VIOLATION
            return jsonify({"status": "error", "message": f"이미 '{title}' 제목의 문서가 존재합니다."}), 409
        app.logger.error(f"Error creating document: {e}")
        return jsonify({"status": "error", "message": "문서 생성 중 오류 발생"}), 500


# 3. 전체 문서 목록을 보여주는 라우트
@app.route('/docs')
def view_all_documents():
    return redirect(NOTION_PUBLIC_WIKI_URL, code=302)
    """
    지금까지 생성된 모든 문서의 목록을 보여주는 페이지.
    """
    try:
        # 문서의 제목, 마지막 수정일, 그리고 작성자(members 테이블과 join)의 이름을 가져옵니다.
        # 최근 수정된 문서가 위로 오도록 정렬합니다.
        docs_res = supabase.table('documents').select('title, updated_at, members(name)') \
            .order('updated_at', desc=True).execute()

        documents = docs_res.data

        return render_template('doc_list.html', documents=documents)

    except Exception as e:
        app.logger.error(f"Error fetching document list: {e}")
        flash("문서 목록을 불러오는 중 오류가 발생했습니다.", "danger")
        return redirect(url_for('main_index'))


# 4. 문서 수정 페이지를 보여주는 라우트
@app.route('/docs/edit/<doc_title>')
@login_required(role="ANY")
def edit_document_page(doc_title):
    return redirect(NOTION_PUBLIC_WIKI_URL, code=302)
    try:
        doc_res = supabase.table('documents').select('*').eq('title', doc_title).single().execute()
        document = doc_res.data

        if not document:
            flash(f"'{doc_title}' 문서를 찾을 수 없습니다.", "warning")
            return redirect(url_for('view_all_documents'))

        # 권한 확인: 작성자 본인이거나 관리자인지 확인
        # if document['author_id'] != session.get('user_id') and session.get('user_role') != 'admin':
        #     flash("이 문서를 수정할 권한이 없습니다.", "danger")
        #     return redirect(url_for('view_document', doc_title=doc_title))

        # 생성 시 사용했던 doc_edit.html 템플릿을 재사용하되,
        # 기존 문서 데이터를 함께 전달하여 폼을 채워넣음
        return render_template('doc_edit.html', doc=document)

    except Exception as e:
        app.logger.error(f"Error loading edit page for document '{doc_title}': {e}")
        flash("편집 페이지를 불러오는 중 오류가 발생했습니다.", "danger")
        return redirect(url_for('view_all_documents'))


# 5. 문서 수정 요청을 처리하는 API 라우트
@app.route('/api/docs/edit/<doc_id>', methods=['POST'])
@login_required(role="ANY")
def handle_edit_document(doc_id):
    return jsonify({
        "status": "error", "message": "앱 위키 편집은 종료되었습니다.",
        "wiki_url": NOTION_PUBLIC_WIKI_URL,
    }), 410
    data = request.json
    content = data.get('content')
    editor_id = session.get('user_id')

    if not content:
        return jsonify({"status": "error", "message": "내용이 없습니다."}), 400

    try:
        # [수정] 권한 및 문서 존재 여부 확인을 위해 먼저 title을 가져옵니다.
        doc_res = supabase.table('documents').select('author_id, title').eq('id', doc_id).single().execute()
        document = doc_res.data

        if not document:
            return jsonify({"status": "error", "message": "수정할 문서를 찾을 수 없습니다."}), 404

        # 1. 'documents' 테이블의 내용을 업데이트합니다. (반환값은 사용하지 않음)
        supabase.table('documents').update({
            'content': content,
            'updated_at': 'now()'
        }).eq('id', doc_id).execute()

        # 2. 'document_logs' 테이블에 변경 이력을 삽입합니다.
        supabase.table('document_logs').insert({
            'document_id': doc_id,
            'editor_id': editor_id,
            'content': content
        }).execute()

        # 3. 바로 화면에 반영할 수 있도록 렌더링된 HTML을 생성합니다.
        rendered_html = wiki_parser(content)

        return jsonify({
            "status": "success",
            "message": "문서가 성공적으로 수정되었습니다.",
            "doc_title": document['title'],  # [수정] 기존에 조회한 문서의 title을 사용합니다.
            "rendered_html": rendered_html
        })

    except Exception as e:
        app.logger.error(f"Error updating document id {doc_id}: {e}")
        return jsonify({"status": "error", "message": "문서 수정 중 오류 발생"}), 500


# 2. [신규] 문서 수정 로그를 가져오는 API 라우트
@app.route('/api/docs/<doc_id>/history')
@login_required(role="ANY")
def get_document_history(doc_id):
    try:
        logs_res = supabase.table('document_logs').select('created_at, content, members(name)') \
            .eq('document_id', doc_id).order('created_at', desc=True).execute()

        return jsonify(logs_res.data)
    except Exception as e:
        app.logger.error(f"Error fetching history for doc id {doc_id}: {e}")
        return jsonify({"error": "로그를 불러오는 중 오류 발생"}), 500


# 6. [신규] 문서 삭제 요청을 처리하는 API 라우트
@app.route('/api/docs/delete/<doc_id>', methods=['POST'])
@login_required(role="ANY")
def handle_delete_document(doc_id):
    return jsonify({
        "status": "error", "message": "기존 위키 자료는 보존 기간이라 삭제할 수 없습니다.",
        "wiki_url": NOTION_PUBLIC_WIKI_URL,
    }), 410
    try:
        # 삭제 권한 확인을 위해 먼저 문서의 작성자 정보를 가져옵니다.
        doc_res = supabase.table('documents').select('author_id, title').eq('id', doc_id).single().execute()
        document = doc_res.data

        if not document:
            return jsonify({"status": "error", "message": "삭제할 문서를 찾을 수 없습니다."}), 404

        # 권한 확인: 작성자 본인이거나 관리자가 아니면 삭제 불가
        if document['author_id'] != session.get('user_id') and session.get('user_role') != 'admin':
            return jsonify({"status": "error", "message": "이 문서를 삭제할 권한이 없습니다."}), 403

        # 문서 삭제 실행
        # 'document_logs' 테이블에 ON DELETE CASCADE를 설정했기 때문에,
        # 원본 문서만 삭제해도 관련 로그가 모두 자동으로 삭제됩니다.
        supabase.table('documents').delete().eq('id', doc_id).execute()

        flash(f"'{document['title']}' 문서가 성공적으로 삭제되었습니다.", "success")
        return jsonify({"status": "success", "message": "문서가 삭제되었습니다."})

    except Exception as e:
        app.logger.error(f"Error deleting document id {doc_id}: {e}")
        return jsonify({"status": "error", "message": "문서 삭제 중 오류 발생"}), 500


# 수동 추가 시 정보 확인용
@app.route('/api/bookclub/preview_manual_groups', methods=['POST'])
@login_required(role="admin")
def preview_manual_groups():
    try:
        data = request.get_json()
        groups = data.get('groups')

        if not groups:
            return jsonify({"error": "그룹 정보가 없습니다."}), 400

        # 1. DB에서 전체 회원 정보와 만남 기록을 가져옵니다.
        all_members_res = supabase.table("members").select("name, gender").execute()
        name_to_gender = {m['name']: normalize_gender(m.get('gender')) for m in all_members_res.data}

        # [수정] last_met 컬럼도 함께 가져옵니다.
        co_matrix_res = supabase.table("bookclub_co_matrix").select("pair_key, count, last_met").execute()
        co_matrix = {item['pair_key']: {'count': item['count'], 'last_met': item.get('last_met')} for item in
                     co_matrix_res.data}

        # 발제자 표시용 '*' 접미사를 제거하여 실제 회원 이름으로 정규화
        def strip_facilitator_mark(nm):
            return nm[:-1].strip() if isinstance(nm, str) and nm.endswith('*') else nm

        groups = [[strip_facilitator_mark(n) for n in g if n] for g in groups]

        # 2. 그룹별 분석 시작
        group_analysis = []
        for i, group in enumerate(groups):
            # 성비 계산
            gender_counts = {'M': 0, 'W': 0, 'Unknown': 0}
            for name in group:
                gender = name_to_gender.get(name)
                if gender in ['M', 'W']:
                    gender_counts[gender] += 1
                else:
                    gender_counts['Unknown'] += 1

            # 만남 기록 분석
            new_encounters = []
            past_encounters = []

            # itertools.combinations를 사용하여 그룹 내 모든 쌍을 생성
            for name1, name2 in itertools.combinations(group, 2):
                pair_key = _canonical_pair_key(name1, name2)

                if pair_key in co_matrix:
                    # 만난 기록이 있는 경우
                    record = co_matrix[pair_key]
                    past_encounters.append({
                        "pair": f"{name1} & {name2}",
                        "count": record['count'],
                        "last_met": record.get('last_met', 'N/A')  # last_met이 없을 경우 대비
                    })
                else:
                    # 처음 만나는 경우
                    new_encounters.append(f"{name1} & {name2}")

            # 결과 구조화
            group_analysis.append({
                "group_index": i + 1,
                "gender_balance": gender_counts,
                "encounters": {
                    "new": new_encounters,
                    "past": sorted(past_encounters, key=lambda x: x['count'], reverse=True)  # 횟수 내림차순 정렬
                }
            })

        return jsonify({"group_analysis": group_analysis})

    except Exception as e:
        app.logger.error(f"Error in preview_manual_groups: {e}", exc_info=True)
        return jsonify({"error": "데이터 분석 중 오류가 발생했습니다."}), 500


#=== 마이페이지
@app.route('/mypage')
@login_required(role="ANY")
def my_page():
    user_id = session.get('user_id')
    user_name = session.get('user_name')

    try:
        # 1. 내 정보 조회
        user_res = supabase.table('members').select('*').eq('id', user_id).single().execute()
        user_data = user_res.data
        if not user_data:
            flash("사용자 정보를 찾을 수 없습니다.", "danger")
            return redirect(url_for('main_index'))

        attendance_records_res = supabase.table('attendance').select('meeting_date') \
            .eq('user_id', user_id).eq('attending_seminar', True).order('meeting_date', desc=True).execute()
        attendance_records = attendance_records_res.data


        # 다음 세미나 날짜 계산 (월/목)
        seminar_date_objs = get_next_seminar_dates()
        weekday_labels = {0: '월', 1: '화', 2: '수', 3: '목', 4: '금', 5: '토', 6: '일'}
        seminar_dates = [
            {
                'date': d.isoformat(),
                'label': f"{d.strftime('%m/%d')} ({weekday_labels[d.weekday()]})"
            }
            for d in seminar_date_objs
        ]
        date_strs = [s['date'] for s in seminar_dates]

        # 나의 참석 확인 날짜 목록
        confirmed_res = supabase.table('attendance').select('meeting_date') \
            .eq('user_id', user_id) \
            .in_('meeting_date', date_strs) \
            .eq('attending_seminar', True).execute()
        my_confirmed_dates = {r['meeting_date'] for r in (confirmed_res.data or [])}

        # 현재 진행 중인 발제문 제출 이벤트들 (다중 지원)
        active_topic_events = []
        try:
            topic_res = supabase.table('topic_events').select('*').eq('is_active', True).order('meeting_date', desc=True).execute()
            active_topic_events = [
                event for event in _auto_close_topic_events(topic_res.data or [])
                if event.get('is_active')
            ]
        except Exception:
            pass

        # 내 활동 요약 (세미나/발제/벽돌책/소모임)
        try:
            activity = _aggregate_member_activity(user_id, user_data.get('name', ''))
        except Exception as e:
            app.logger.warning(f"my_page activity error: {e}")
            activity = {'seminar_count': 0, 'facilitator_count': 0, 'brick_sessions': [], 'study_sessions': []}

        # 내가 참여한 스페셜 이벤트 (최근 5개)
        my_special_events = _member_special_events(user_id)[:5]

        # 링크형 참여 기록(후기·추천·벽돌책 지원)은 로그인 계정과 같은 회원 ID로 묶인다.
        engagement_activity = {'seminar_reviews': 0, 'book_suggestions': 0, 'brick_applications': 0, 'brick_reviews': 0}
        try:
            engagement_activity = {
                'seminar_reviews': supabase.table('seminar_reviews').select('id', count='exact')
                    .eq('member_id', user_id).is_('deleted_at', 'null').execute().count or 0,
                'book_suggestions': supabase.table('book_suggestions').select('id', count='exact')
                    .eq('created_by', user_id).neq('status', 'archived').execute().count or 0,
                'brick_applications': supabase.table('brick_applications').select('id', count='exact')
                    .eq('member_id', user_id).execute().count or 0,
                'brick_reviews': supabase.table('brick_reviews').select('id', count='exact')
                    .eq('member_id', user_id).is_('deleted_at', 'null').execute().count or 0,
            }
        except Exception as e:
            app.logger.warning(f"my_page engagement activity error: {e}")

        # 현재 투표 가능한 세미나 회차 (학기별)
        my_open_votes = []
        active_seminar_terms = []
        try:
            now_kst = datetime.now(KST)
            terms = supabase.table('seminar_terms').select('id, name, share_token, max_capacity') \
                .eq('is_active', True).execute().data or []
            active_seminar_terms = [
                {'name': t['name'], 'share_token': t['share_token']} for t in terms
            ]
            for term in terms:
                t_sess = supabase.table('seminar_sessions') \
                    .select('id, meeting_date, day_type, participation_mode, capacity, vote_open_at, vote_close_at') \
                    .eq('term_id', term['id']).eq('is_active', True) \
                    .eq('day_type', 'mon').eq('participation_mode', 'opt_in') \
                    .order('meeting_date').execute().data or []
                # 내 기존 투표
                sids = [s['id'] for s in t_sess]
                my_votes = {}
                if sids:
                    mv = supabase.table('seminar_votes').select('session_id, attending') \
                        .in_('session_id', sids).eq('member_id', user_id).execute().data or []
                    my_votes = {v['session_id']: v['attending'] for v in mv}
                for s in t_sess:
                    open_at, close_at = _voting_window_for(s)
                    if open_at <= now_kst <= close_at:
                        s['my_vote'] = my_votes.get(s['id'])  # True/False/None
                        s['term_token'] = term['share_token']
                        s['term_name'] = term['name']
                        s['vote_close_label'] = close_at.strftime('%m/%d %H:%M')
                        my_open_votes.append(s)
        except Exception as e:
            app.logger.warning(f"my_page open votes error: {e}")

        return render_template(
            'my_page_member.html',
            user=user_data,
            attendance_records=attendance_records,
            seminar_dates=seminar_dates,
            my_confirmed_dates=my_confirmed_dates,
            active_topic_events=active_topic_events,
            activity=activity,
            my_open_votes=my_open_votes,
            active_seminar_terms=active_seminar_terms,
            my_special_events=my_special_events,
            engagement_activity=engagement_activity,
        )

    except Exception as e:
        app.logger.error(f"Error loading my page for user {user_id}: {e}")
        flash("마이페이지를 불러오는 중 오류가 발생했습니다.", "danger")
        return redirect(url_for('main_index'))


@app.route('/api/request_absence', methods=['POST'])
@login_required(role="ANY")
def request_absence():
    user_id = session.get('user_id')
    reason = request.json.get('reason')
    next_monday = get_next_monday()

    if datetime.now(timezone(timedelta(hours=9))).date().weekday() == 0:
        return jsonify({"error": "당일에는 불참 요청을 할 수 없습니다. 관리자에게 직접 문의하세요."}), 403

    if not reason or not reason.strip():
        return jsonify({"error": "불참 사유를 입력해야 합니다."}), 400

    try:
        # 1. [수정] attendance 테이블의 상태를 'pending'으로 설정
        supabase.table('attendance').upsert({
            'user_id': user_id,
            'meeting_date': next_monday.isoformat(),
            'attending_seminar': False,  # 우선 불참으로 설정
            'attending_afterparty': False,
            'absence_reason': reason,
            'absence_request_status': 'pending'  # 승인 대기 상태
        }, on_conflict='user_id, meeting_date').execute()

        # 2. [수정] 관리자에게 '승인/반려'가 필요한 알림 생성
        supabase.table('notifications').insert({
            'type': 'absence_request',  # 승인/반려가 필요한 타입으로 변경
            'related_member_id': user_id,
            'details': {'name': session.get('user_name'), 'reason': reason}
        }).execute()

        return jsonify({"status": "success", "message": "불참 요청이 관리자에게 전달되었습니다."})
    except Exception as e:
        app.logger.error(f"Error on absence request for user {user_id}: {e}")
        return jsonify({"error": "처리 중 오류가 발생했습니다."}), 500



#=== 벌점 추가
# ==============================================================================
# --- [신규] 주간 발제문 수집 및 문서화 시스템 ---
# ==============================================================================


def _auto_close_topic_events(events):
    """공유 발제문은 마지막 연결 회차가 끝난 다음 날 마감한다."""
    rows = list(events or [])
    week_ids = list({row['seminar_week_id'] for row in rows if row.get('seminar_week_id')})
    session_ids = list({row['seminar_session_id'] for row in rows
                        if row.get('seminar_session_id') and not row.get('seminar_week_id')})
    related_sessions = []
    if week_ids:
        related_sessions.extend(supabase.table('seminar_sessions')
            .select('id, seminar_week_id, meeting_date').in_('seminar_week_id', week_ids).execute().data or [])
    if session_ids:
        related_sessions.extend(supabase.table('seminar_sessions')
            .select('id, seminar_week_id, meeting_date').in_('id', session_ids).execute().data or [])
    for row in rows:
        row['session_dates'] = sorted({str(item['meeting_date'])[:10] for item in related_sessions
            if item.get('meeting_date') and (
                (row.get('seminar_week_id') and item.get('seminar_week_id') == row['seminar_week_id'])
                or (not row.get('seminar_week_id') and item.get('id') == row.get('seminar_session_id'))
            )})
        deadline = topic_event_deadline(row)
        row['submission_deadline'] = deadline.isoformat() if deadline else None
    today_kst = datetime.now(timezone(timedelta(hours=9))).date()
    expired = [row for row in rows if row.get('is_active') and topic_event_is_expired(row, today=today_kst)]
    expired_ids = [row.get('id') for row in expired if row.get('id')]
    if expired_ids:
        try:
            supabase.table('topic_events').update({'is_active': False}).in_('id', expired_ids).execute()
        except Exception as exc:
            app.logger.warning("topic event auto-close sync failed: %s", exc)
        for row in expired:
            row['is_active'] = False
            row['auto_closed'] = True
    return rows


def _topic_event_for_submission(event_rows, session_id=None):
    """회차 링크의 마감은 공유 이벤트 전체를 닫지 않고 별도로 검사한다."""
    events = _auto_close_topic_events(event_rows)
    if not events:
        return None
    event = dict(events[0])
    if session_id:
        seminar_session = next((item for item in _topic_event_sessions(event)
                                if str(item['id']) == str(session_id)), None)
        if not seminar_session:
            raise ValueError('이 발제문에 연결되지 않은 세미나 회차입니다.')
        event['submission_session_id'] = seminar_session['id']
        event['submission_meeting_date'] = seminar_session['meeting_date']
        deadline = topic_event_deadline(event)
        event['submission_deadline'] = deadline.isoformat() if deadline else None
        if topic_event_is_expired(event, today=datetime.now(timezone(timedelta(hours=9))).date()):
            event['is_active'] = False
    return event


def _topic_event_sessions(event):
    """다운로드 날짜·사회자를 선택할 때 이벤트 소속 회차만 허용한다."""
    query = supabase.table('seminar_sessions').select('*')
    if event.get('seminar_week_id'):
        query = query.eq('seminar_week_id', event['seminar_week_id'])
    elif event.get('seminar_session_id'):
        query = query.eq('id', event['seminar_session_id'])
    else:
        return []
    return query.order('meeting_date').execute().data or []


# 1. 관리자: 발제문 이벤트 생성 API
def _open_topic_event_for_session(session_id):
    seminar_session = supabase.table('seminar_sessions').select(
        'id, meeting_date, book_title, book_author, seminar_week_id'
    ).eq('id', session_id).single().execute().data
    if not seminar_session:
        raise ValueError('세미나 회차를 찾을 수 없습니다.')

    week_id = seminar_session.get('seminar_week_id')
    existing_query = supabase.table('topic_events').select('*')
    existing_query = (existing_query.eq('seminar_week_id', week_id) if week_id
                      else existing_query.eq('seminar_session_id', session_id))
    existing = existing_query.execute().data or []
    if existing:
        return _auto_close_topic_events(existing)[0], False

    week = None
    if week_id:
        rows = supabase.table('seminar_weeks').select('*').eq('id', week_id).execute().data or []
        week = rows[0] if rows else None
    book_title = ((week or {}).get('book_title') or seminar_session.get('book_title') or '').strip()
    book_author = ((week or {}).get('book_author') or seminar_session.get('book_author') or '').strip()
    if not book_title:
        raise ValueError('먼저 이번 회차 도서 제목을 입력해주세요.')

    payload = {
        'meeting_date': seminar_session['meeting_date'],
        'book_title': book_title,
        'book_author': book_author or None,
        'seminar_session_id': session_id,
        'seminar_week_id': week_id,
        'share_token': str(uuid.uuid4()),
        'is_active': True,
    }
    try:
        created = supabase.table('topic_events').insert(payload).execute().data or []
        return created[0], True
    except Exception:
        # 동시에 두 번 눌린 경우 unique index가 중복을 막는다.
        raced_query = supabase.table('topic_events').select('*')
        raced_query = (raced_query.eq('seminar_week_id', week_id) if week_id
                       else raced_query.eq('seminar_session_id', session_id))
        raced = raced_query.execute().data or []
        if raced:
            return raced[0], False
        raise


@app.route('/api/admin/topic_events/create', methods=['POST'])
@login_required(role="admin")
def create_topic_event():
    try:
        data = request.json or {}
        if data.get('seminar_session_id'):
            event, created = _open_topic_event_for_session(data['seminar_session_id'])
            return jsonify({
                "status": "success", "created": created, "event": event,
                "share_url": f"{request.host_url}shared_topics?token={event['share_token']}"
            })
        token = str(uuid.uuid4())
        supabase.table('topic_events').insert({
            'meeting_date': data.get('meeting_date'),
            'book_title': data.get('book_title'),
            'book_author': data.get('book_author'),
            'share_token': token
        }).execute()
        return jsonify({"status": "success", "message": "발제문 수집 링크가 생성되었습니다."})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/open_topics', methods=['POST'])
@login_required(role="admin")
def seminar_session_open_topics(session_id):
    try:
        event, created = _open_topic_event_for_session(session_id)
        return jsonify({
            'status': 'success', 'created': created, 'event': event,
            'share_url': f"{request.host_url}shared_topics?token={event['share_token']}"
        })
    except ValueError as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400
    except Exception as e:
        app.logger.error(f"seminar_session_open_topics error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_weeks/<week_id>/open_topics', methods=['POST'])
@login_required(role="admin")
def seminar_week_open_topics(week_id):
    try:
        sessions = supabase.table('seminar_sessions').select('id') \
            .eq('seminar_week_id', week_id).order('meeting_date').execute().data or []
        if not sessions:
            return jsonify({'status': 'error', 'message': '이 주차의 세미나 회차가 없습니다.'}), 404
        event, created = _open_topic_event_for_session(sessions[0]['id'])
        return jsonify({
            'status': 'success', 'created': created, 'event': event,
            'share_url': f"{request.host_url}shared_topics?token={event['share_token']}"
        })
    except ValueError as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400
    except Exception as e:
        app.logger.error(f"seminar_week_open_topics error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500

# 1.5. 관리자: 발제문 이벤트 영구 삭제 API (제출 내역까지 포함)
@app.route('/api/admin/topic_events/<event_id>/delete', methods=['POST'])
@login_required(role="admin")
def delete_topic_event(event_id):
    try:
        supabase.table('topic_submissions').delete().eq('event_id', event_id).execute()
        supabase.table('topic_events').delete().eq('id', event_id).execute()
        return jsonify({"status": "success", "message": "발제문 이벤트가 영구 삭제되었습니다."})
    except Exception as e:
        app.logger.error(f"Error deleting topic event: {e}")
        return jsonify({"error": "이벤트 삭제 중 서버 오류가 발생했습니다."}), 500


# 1.6. 관리자: 발제문 이벤트 활성/숨김 토글 (soft delete)
@app.route('/api/admin/topic_events/<event_id>/toggle_active', methods=['POST'])
@login_required(role="admin")
def toggle_topic_event(event_id):
    try:
        cur = supabase.table('topic_events').select('*').eq('id', event_id).single().execute().data
        if not cur:
            return jsonify({"error": "이벤트를 찾을 수 없습니다."}), 404
        cur = _auto_close_topic_events([cur])[0]
        new_state = not bool(cur.get('is_active'))
        today_kst = datetime.now(timezone(timedelta(hours=9))).date()
        if new_state and topic_event_is_expired(cur, today=today_kst):
            return jsonify({"error": "마지막 세미나 날짜가 지난 발제문은 자동 마감되어 다시 열 수 없습니다."}), 400
        supabase.table('topic_events').update({'is_active': new_state}).eq('id', event_id).execute()
        return jsonify({"status": "success", "is_active": new_state})
    except Exception as e:
        app.logger.error(f"Error toggling topic event: {e}")
        return jsonify({"error": "상태 변경 중 오류가 발생했습니다."}), 500


def _authenticated_topic_member():
    """Return the active member tied to the authenticated session.

    Topic submission links are public, so client-provided names and student IDs
    must never be treated as proof of membership.  Refresh the member from the
    database using the server-side session ID and fail closed when the account is
    missing or no longer active.
    """
    user_id = session.get('user_id')
    if not user_id:
        return None

    try:
        rows = supabase.table('members').select(
            'id, name, department, student_id, account_status, member_status, is_active'
        ).eq('id', user_id).limit(1).execute().data or []
    except Exception as exc:
        app.logger.error("Topic member authorization refresh failed: %s", exc)
        return None

    if not rows:
        return None

    member = rows[0]
    if (
        member.get('account_status') != 'active'
        or member.get('member_status') == 'inactive'
        or member.get('is_active') is False
        or not (member.get('name') or '').strip()
    ):
        return None
    return member


TOPIC_EDIT_FAILURE_LIMIT = 10
TOPIC_EDIT_FAILURE_WINDOW_MINUTES = 15


def _topic_edit_request_hash(event_id):
    address = forwarded_client_address(
        request.headers.get('X-Forwarded-For'), request.remote_addr
    )
    return topic_request_fingerprint(address, event_id, app.secret_key)


def _topic_edit_is_rate_limited(event_id):
    request_hash = _topic_edit_request_hash(event_id)
    rate_start = (
        datetime.now(timezone.utc) - timedelta(minutes=TOPIC_EDIT_FAILURE_WINDOW_MINUTES)
    ).isoformat()
    try:
        attempts = supabase.table('topic_edit_attempts').select('id', count='exact') \
            .eq('event_id', event_id).eq('request_hash', request_hash) \
            .eq('succeeded', False).gte('created_at', rate_start).execute()
        return (attempts.count or 0) >= TOPIC_EDIT_FAILURE_LIMIT
    except Exception as exc:
        app.logger.warning("topic edit rate-limit lookup failed: %s", exc)
        return False


def _record_topic_edit_attempt(event_id, succeeded):
    try:
        supabase.table('topic_edit_attempts').insert({
            'event_id': event_id,
            'request_hash': _topic_edit_request_hash(event_id),
            'succeeded': bool(succeeded),
        }).execute()
    except Exception as exc:
        app.logger.warning("topic edit attempt audit failed: %s", exc)


def _guest_topic_credential_matches(existing_record, credential):
    identity_kind = existing_record.get('identity_kind')
    if not identity_kind:
        identity_kind = (
            'legacy_member'
            if existing_record.get('pin_code') == 'MEMBER'
            else 'legacy_pin'
        )
    if identity_kind == 'guest':
        return topic_edit_token_matches(
            existing_record.get('edit_token_hash'), credential, app.secret_key
        )
    if identity_kind == 'legacy_pin':
        return legacy_pin_matches(existing_record.get('pin_code'), credential)
    return False


def _registered_topic_student_id(student_id):
    sid = str(student_id or '').strip()
    if not sid:
        return False
    try:
        rows = supabase.table('members').select('id').eq('student_id', sid).limit(1).execute().data or []
        return bool(rows)
    except Exception as exc:
        app.logger.error("topic member student-id check failed: %s", exc)
        return None


# 2. 사용자: 공유 링크를 통한 발제문 작성 페이지
@app.route('/shared_topics')
def view_shared_topics():
    token = request.args.get('token')
    if not token:
        flash("잘못된 접근입니다.", "danger")
        return redirect(url_for('main_index'))

    try:
        event_res = supabase.table('topic_events').select('*').eq('share_token', token).single().execute()
        event_data = event_res.data
        if event_data:
            event_data = _topic_event_for_submission([event_data], request.args.get('session_id'))
        if not event_data or not event_data.get('is_active'):
            flash("마감되었거나 유효하지 않은 링크입니다.", "warning")
            return redirect(url_for('main_index'))

        # 중복 방지용 공개 영역에는 발제 JSON만 조회한다. 작성자 정보는 서버에서부터 가져오지 않는다.
        preview_rows = supabase.table('topic_submissions').select('topics') \
            .eq('event_id', event_data['id']).order('created_at').execute().data or []
        existing_topic_previews = anonymous_topic_previews(preview_rows)

        member = _authenticated_topic_member()
        user_name = member.get('name') if member else None
        user_department = member.get('department') if member else None
        user_student_id = member.get('student_id') if member else None
        return render_template('topic_submit.html',
                               event=event_data,
                               user_name=user_name,
                               user_department=user_department,
                               user_student_id=user_student_id,
                               existing_topic_previews=existing_topic_previews)
    except Exception as e:
        app.logger.error(f"Error loading topic event: {e}")
        return "유효하지 않은 링크입니다.", 404


# 3. 사용자: 발제문 제출/수정 API (PIN 검증 포함)
@app.route('/api/topics/submit', methods=['POST'])
def submit_topics():
    data = request.get_json(silent=True) or {}
    event_id = data.get('event_id')
    author_name = (data.get('author_name') or '').strip()
    department = (data.get('department') or '').strip()
    edit_credential = str(data.get('edit_credential') or data.get('pin_code') or '').strip()
    student_id = (data.get('student_id') or '').strip()
    topics = data.get('topics')  # JSON Array

    # 회원 권한은 입력한 이름/학번이 아니라 서버 세션의 member id로만 확인한다.
    member = _authenticated_topic_member()
    is_logged_in_member = member is not None
    if member:
        author_name = (member.get('name') or '').strip()
        department = (member.get('department') or department).strip()
        student_id = str(member.get('student_id') or student_id).strip()

    if not all([event_id, author_name, department, topics]):
        return jsonify({"error": "필수 정보를 모두 입력해주세요."}), 400

    try:
        event_rows = supabase.table('topic_events').select('*') \
            .eq('id', event_id).limit(1).execute().data or []
        event_data = _topic_event_for_submission(event_rows, data.get('session_id'))
        if not event_data or not event_data.get('is_active'):
            return jsonify({"error": "해당 세미나 날짜가 지나 발제문 제출이 마감되었습니다."}), 400

        # 기존 제출 내역 확인 (학번 우선 매칭 — 학과 표기만 바꿔 여러 건 제출하는 것 방지)
        existing_record = _find_topic_submission(
            event_id,
            author_name,
            department,
            student_id,
            member_id=member.get('id') if member else None,
        )

        # 학번에서 입학년도 2자리 추출 (예: "2022123456" → "22")
        # student_id 가 4자 이상이면 3-4번째 문자 사용
        admission_year = ''
        sid = (student_id or '').strip()
        if len(sid) >= 4 and sid[2:4].isdigit():
            admission_year = sid[2:4]

        # 발제문 개수 제한: 1개는 필수, 기본 한도는 선택 제출 1개를 더한 2개.
        # 회장이 개별적으로 바꾼 기존 한도는 그대로 따른다.
        topic_limit = (existing_record or {}).get('topic_limit') or 2
        if len(topics) > topic_limit:
            return jsonify({"error": f"발제문은 {topic_limit}개까지 작성할 수 있습니다. 더 쓰고 싶다면 회장에게 문의해주세요."}), 400

        if existing_record:
            issued_edit_token = None
            if is_logged_in_member:
                owner_id = existing_record.get('member_id')
                if owner_id is not None and str(owner_id) != str(member.get('id')):
                    return jsonify({"error": "수정 권한을 확인할 수 없습니다."}), 403
            else:
                if _topic_edit_is_rate_limited(event_id):
                    return jsonify({"error": "수정 확인이 잠시 제한되었습니다. 15분 후 다시 시도해주세요."}), 429
                credential_ok = _guest_topic_credential_matches(existing_record, edit_credential)
                _record_topic_edit_attempt(event_id, credential_ok)
                if not credential_ok:
                    return jsonify({"error": "수정 정보를 확인할 수 없습니다."}), 403

            # 업데이트 실행
            update_payload = {'topics': topics, 'updated_at': 'now()', 'department': department}
            if admission_year:
                update_payload['admission_year'] = admission_year
            if sid:
                update_payload['student_id'] = sid
            if is_logged_in_member:
                update_payload.update({
                    'member_id': member.get('id'),
                    'identity_kind': 'member',
                    'credential_version': 2,
                    'pin_code': 'MEMBER',
                    'edit_token_hash': None,
                })
            elif existing_record.get('identity_kind') in (None, 'legacy_pin'):
                issued_edit_token = generate_topic_edit_token()
                update_payload.update({
                    'identity_kind': 'guest',
                    'credential_version': 2,
                    'pin_code': 'TOKEN',
                    'edit_token_hash': topic_edit_token_digest(
                        issued_edit_token, app.secret_key
                    ),
                })
            supabase.table('topic_submissions').update(update_payload) \
                .eq('id', existing_record['id']).execute()
            response = {"status": "success", "message": "발제문이 성공적으로 수정되었습니다."}
            if issued_edit_token:
                response['edit_token'] = issued_edit_token
                response['message'] = "발제문이 수정되었고, 안전한 새 수정 코드가 발급되었습니다."
            return jsonify(response)
        else:
            # 신규 생성 모드
            if not is_logged_in_member:
                registered_student_id = _registered_topic_student_id(sid)
                if registered_student_id is None:
                    return jsonify({"error": "회원 정보를 확인하지 못했습니다. 잠시 후 다시 시도해주세요."}), 503
                if registered_student_id:
                    return jsonify({"error": "등록된 회원 학번은 로그인 후 제출해주세요."}), 403
            issued_edit_token = None
            insert_payload = {
                'event_id': event_id,
                'author_name': author_name,
                'department': department,
                'admission_year': admission_year or None,
                'student_id': sid or None,
                'topics': topics
            }
            if is_logged_in_member:
                insert_payload.update({
                    'member_id': member.get('id'),
                    'identity_kind': 'member',
                    'credential_version': 2,
                    'pin_code': 'MEMBER',
                    'edit_token_hash': None,
                })
            else:
                issued_edit_token = generate_topic_edit_token()
                insert_payload.update({
                    'member_id': None,
                    'identity_kind': 'guest',
                    'credential_version': 2,
                    'pin_code': 'TOKEN',
                    'edit_token_hash': topic_edit_token_digest(
                        issued_edit_token, app.secret_key
                    ),
                })
            supabase.table('topic_submissions').insert(insert_payload).execute()
            response = {"status": "success", "message": "발제문이 성공적으로 제출되었습니다."}
            if issued_edit_token:
                response['edit_token'] = issued_edit_token
                response['message'] = "발제문이 제출되었습니다. 아래 수정 코드를 꼭 보관해주세요."
            return jsonify(response)

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        app.logger.error(f"Error submitting topics: {e}")
        return jsonify({"error": "제출 중 서버 오류가 발생했습니다."}), 500


def _find_topic_submission(event_id, author_name, department, student_id, member_id=None):
    """같은 이벤트에서 같은 사람의 기존 제출물을 찾는다.
    학번이 있으면 (event_id, student_id)로 우선 매칭하고,
    없거나 매칭 실패 시 과거 데이터 호환을 위해 (event_id, author_name, department)로 폴백한다."""
    if member_id is not None:
        member_res = supabase.table('topic_submissions').select('*') \
            .eq('event_id', event_id).eq('member_id', member_id).execute()
        if member_res.data:
            return member_res.data[0]

    sid = (student_id or '').strip()
    if sid:
        res = supabase.table('topic_submissions').select('*') \
            .eq('event_id', event_id).eq('student_id', sid).execute()
        if res.data:
            candidate = res.data[0]
            if member_id is None:
                return candidate
            if (
                candidate.get('member_id') is None
                and candidate.get('pin_code') == 'MEMBER'
            ):
                return candidate
    res = supabase.table('topic_submissions').select('*') \
        .eq('event_id', event_id).eq('author_name', author_name).eq('department', department).execute()
    if not res.data:
        return None
    candidate = res.data[0]
    if member_id is None:
        return candidate
    if (
        candidate.get('member_id') is None
        and candidate.get('pin_code') == 'MEMBER'
    ):
        return candidate
    return None


# 3.5. 사용자: 발제문 불러오기 API
@app.route('/api/topics/load', methods=['POST'])
def load_topics():
    data = request.get_json(silent=True) or {}
    event_id = data.get('event_id')
    author_name = (data.get('author_name') or '').strip()
    department = (data.get('department') or '').strip()
    edit_credential = str(data.get('edit_credential') or data.get('pin_code') or '').strip()
    student_id = (data.get('student_id') or '').strip()

    # 회원 권한은 입력한 이름/학번이 아니라 서버 세션의 member id로만 확인한다.
    member = _authenticated_topic_member()
    is_logged_in_member = member is not None
    if member:
        author_name = (member.get('name') or '').strip()
        department = (member.get('department') or department).strip()
        student_id = str(member.get('student_id') or student_id).strip()

    if not all([event_id, author_name, department]):
        return jsonify({"error": "이름과 소속을 모두 입력해주세요."}), 400

    try:
        event_rows = supabase.table('topic_events').select('*') \
            .eq('id', event_id).limit(1).execute().data or []
        event_data = _topic_event_for_submission(event_rows, data.get('session_id'))
        if not event_data or not event_data.get('is_active'):
            return jsonify({"error": "해당 세미나 날짜가 지나 발제문 제출이 마감되었습니다."}), 400

        existing_record = _find_topic_submission(
            event_id,
            author_name,
            department,
            student_id,
            member_id=member.get('id') if member else None,
        )

        if existing_record:
            if is_logged_in_member:
                owner_id = existing_record.get('member_id')
                if owner_id is not None and str(owner_id) != str(member.get('id')):
                    return jsonify({"error": "수정 권한을 확인할 수 없습니다."}), 403
            else:
                if _topic_edit_is_rate_limited(event_id):
                    return jsonify({"error": "수정 확인이 잠시 제한되었습니다. 15분 후 다시 시도해주세요."}), 429
                credential_ok = _guest_topic_credential_matches(existing_record, edit_credential)
                _record_topic_edit_attempt(event_id, credential_ok)
                if not credential_ok:
                    return jsonify({"error": "수정 정보를 확인할 수 없습니다."}), 403

            return jsonify({
                "status": "success",
                "topics": existing_record['topics'],
                "topic_limit": existing_record.get('topic_limit') or 2,
            })
        else:
            return jsonify({"error": "작성된 발제문 내역이 없습니다. 처음 작성하는 것이 맞나요?"}), 404

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        app.logger.error(f"Error loading topics: {e}")
        return jsonify({"error": "불러오기 중 서버 오류가 발생했습니다."}), 500


# 3.7 관리자: 발제문 제출 내역 삭제/수정 API (다른 사람의 것도 관리 가능)
@app.route('/api/admin/topic_submissions/<submission_id>/delete', methods=['POST'])
@login_required(role="admin")
def admin_delete_topic_submission(submission_id):
    try:
        supabase.table('topic_submissions').delete().eq('id', submission_id).execute()
        return jsonify({"status": "success", "message": "발제문이 삭제되었습니다."})
    except Exception as e:
        app.logger.error(f"admin_delete_topic_submission error: {e}")
        return jsonify({"error": "삭제 중 오류가 발생했습니다."}), 500


@app.route('/api/admin/topic_submissions/<submission_id>/update', methods=['POST'])
@login_required(role="admin")
def admin_update_topic_submission(submission_id):
    try:
        data = request.json or {}
        topics = data.get('topics')
        if not topics or not isinstance(topics, list):
            return jsonify({"error": "발제문 내용이 비어있습니다."}), 400
        update_fields = {'topics': topics, 'updated_at': 'now()'}
        # 작성자/소속도 함께 수정할 수 있게 허용
        if data.get('author_name'):
            update_fields['author_name'] = data['author_name']
        if data.get('department'):
            update_fields['department'] = data['department']
        supabase.table('topic_submissions').update(update_fields).eq('id', submission_id).execute()
        return jsonify({"status": "success", "message": "발제문이 수정되었습니다."})
    except Exception as e:
        app.logger.error(f"admin_update_topic_submission error: {e}")
        return jsonify({"error": "수정 중 오류가 발생했습니다."}), 500


@app.route('/api/admin/topic_submissions/<submission_id>/set_limit', methods=['POST'])
@login_required(role="admin")
def admin_set_topic_limit(submission_id):
    """회장이 특정 제출자의 발제문 한도를 개별적으로 조정합니다."""
    try:
        limit = int((request.json or {}).get('topic_limit', 2))
        if not (1 <= limit <= 10):
            return jsonify({"error": "한도는 1~10 사이여야 합니다."}), 400
        supabase.table('topic_submissions').update({'topic_limit': limit}).eq('id', submission_id).execute()
        return jsonify({"status": "success", "message": f"발제문 한도가 {limit}개로 변경되었습니다."})
    except (TypeError, ValueError):
        return jsonify({"error": "올바른 숫자를 입력해주세요."}), 400
    except Exception as e:
        app.logger.error(f"admin_set_topic_limit error: {e}")
        return jsonify({"error": "한도 변경 중 오류가 발생했습니다."}), 500


# 3.8 관리자: 발제문 상세 보기 및 취합 페이지
@app.route('/admin/topics/<event_id>/view')
@login_required(role="admin")
def view_admin_topics(event_id):
    try:
        event = supabase.table('topic_events').select('*').eq('id', event_id).single().execute().data
        if not event:
            abort(404)
        event = _auto_close_topic_events([event])[0]
        submissions = supabase.table('topic_submissions').select('*').eq('event_id', event_id).order(
            'created_at').execute().data
        submissions = number_topic_submissions(submissions)
        
        return render_template('admin_topic_view.html', event=event, submissions=submissions,
                               seminar_sessions=_topic_event_sessions(event))
    except Exception as e:
        flash(f"상세 정보를 불러오는 중 오류 발생: {str(e)}", "danger")
        return redirect(url_for('admin_dashboard'))


def _moderator_name(value):
    name = ' '.join(str(value or '').split())
    if not name or len(name) > 100:
        raise ValueError('사회자 이름을 1~100자로 입력해주세요.')
    return name


def _render_topic_document_setup(event, seminar_sessions, selected_session, download, error=None):
    members = supabase.table('members').select('id, name, department') \
        .eq('is_active', True).order('name').execute().data or []
    csrf_token = session.setdefault('topic_moderator_csrf', secrets.token_urlsafe(32))
    return render_template('topic_document_setup.html', event=event,
        seminar_sessions=seminar_sessions, selected_session=selected_session,
        moderator_name=(selected_session or event).get('moderator_name', ''),
        members=members, download=download, error=error, csrf_token=csrf_token)


def _valid_topic_moderator_csrf():
    expected = session.get('topic_moderator_csrf') or ''
    supplied = request.form.get('csrf_token') or ''
    return bool(expected and supplied and secrets.compare_digest(expected, supplied))


@app.route('/admin/seminar_sessions/<session_id>/moderator', methods=['GET', 'POST'])
@login_required(role='admin')
def seminar_session_moderator(session_id):
    rows = supabase.table('seminar_sessions').select('*').eq('id', session_id).limit(1).execute().data or []
    if not rows:
        abort(404)
    seminar_session = rows[0]
    if request.method == 'POST':
        if not _valid_topic_moderator_csrf():
            abort(403)
        try:
            name = _moderator_name(request.form.get('moderator_name'))
            updated = supabase.table('seminar_sessions').update({'moderator_name': name}) \
                .eq('id', session_id).execute().data or []
            if not updated:
                raise RuntimeError('회차가 변경되었거나 삭제되었습니다.')
        except ValueError as exc:
            return _render_topic_document_setup(seminar_session, [seminar_session],
                seminar_session, False, error=str(exc)), 400
        except Exception as exc:
            app.logger.error('seminar moderator save failed: %s', exc)
            return _render_topic_document_setup(seminar_session, [seminar_session],
                seminar_session, False, error='사회자를 저장하지 못했습니다. 잠시 후 다시 시도해주세요.'), 503
        flash('이 회차의 사회자를 저장했습니다.', 'success')
        return redirect(url_for('admin_seminars'))
    return _render_topic_document_setup(seminar_session, [seminar_session], seminar_session, False)


@app.route('/admin/topics/<event_id>/document_setup', methods=['GET', 'POST'])
@login_required(role='admin')
def topic_document_setup(event_id):
    rows = supabase.table('topic_events').select('*').eq('id', event_id).limit(1).execute().data or []
    if not rows:
        abort(404)
    event = rows[0]
    seminar_sessions = _topic_event_sessions(event)
    session_id = request.form.get('session_id') if request.method == 'POST' else request.args.get('session_id')
    selected = next((item for item in seminar_sessions if str(item['id']) == str(session_id)), None)
    if session_id and not selected:
        abort(400, description='이 발제문에 연결되지 않은 세미나 회차입니다.')
    if len(seminar_sessions) == 1 and not selected:
        selected = seminar_sessions[0]
    download = request.values.get('download', '1') != '0'
    if request.method == 'POST':
        if not _valid_topic_moderator_csrf():
            abort(403)
        try:
            if seminar_sessions and not selected:
                raise ValueError('Word에 표시할 세미나 날짜를 선택해주세요.')
            name = _moderator_name(request.form.get('moderator_name'))
            table = 'seminar_sessions' if selected else 'topic_events'
            target_id = selected['id'] if selected else event_id
            updated = supabase.table(table).update({'moderator_name': name}).eq('id', target_id).execute().data or []
            if not updated:
                raise RuntimeError('세미나 정보가 변경되었거나 삭제되었습니다.')
        except ValueError as exc:
            return _render_topic_document_setup(event, seminar_sessions, selected, download, error=str(exc)), 400
        except Exception as exc:
            app.logger.error('topic document metadata save failed: %s', exc)
            return _render_topic_document_setup(event, seminar_sessions, selected, download,
                error='사회자를 저장하지 못했습니다. 잠시 후 다시 시도해주세요.'), 503
        if download:
            return redirect(url_for('download_topics_word', event_id=event_id,
                                    session_id=selected['id'] if selected else None))
        flash('사회자를 저장했습니다.', 'success')
        return redirect(url_for('view_admin_topics', event_id=event_id))
    return _render_topic_document_setup(event, seminar_sessions, selected, download)


# 4. 관리자: Word 파일로 출력 (template.docx 디자인 유지)
# - docxtpl 라이브러리 버전 차이로 production에서 일부 iteration이 누락되는 이슈가 보고됨.
#   대응:
#   1) requirements.txt에 python-docx, docxtpl 버전 핀 (로컬과 동일하게 맞춤)
#   2) 렌더링 후 모든 제출자의 이름이 문서에 실제로 들어갔는지 검증
#   3) 누락 발견 시 누락된 제출만 python-docx로 문서 끝에 append (디자인은 일부 상이하나 누락 방지)
@app.route('/admin/topics/<event_id>/download_word')
@login_required(role="admin")
def download_topics_word(event_id):
    try:
        from docxtpl import DocxTemplate
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_BREAK

        event = supabase.table('topic_events').select('*').eq('id', event_id).single().execute().data
        if not event:
            return '발제문을 찾을 수 없습니다.', 404
        seminar_sessions = _topic_event_sessions(event)
        session_id = request.args.get('session_id')
        selected = next((item for item in seminar_sessions if str(item['id']) == str(session_id)), None)
        if session_id and not selected:
            return '이 발제문에 연결되지 않은 세미나 회차입니다.', 400
        if len(seminar_sessions) == 1 and not selected:
            selected = seminar_sessions[0]
        metadata = selected or event
        if (seminar_sessions and not selected) or not str(metadata.get('moderator_name') or '').strip():
            return redirect(url_for('topic_document_setup', event_id=event_id, session_id=session_id))
        # GET 다운로드에서는 저장하지 않는다. 사회자·날짜는 선택한 회차의 표시값만 사용한다.
        event = dict(event, meeting_date=metadata.get('meeting_date'), moderator_name=metadata['moderator_name'])
        submissions = supabase.table('topic_submissions').select('*') \
            .eq('event_id', event_id).order('created_at').execute().data or []
        submissions = number_topic_submissions(submissions)
        app.logger.info(f"download_topics_word: 제출 {len(submissions)}건")

        template_path = os.path.join(app.root_path, 'templates', 'template.docx')
        if not os.path.exists(template_path):
            flash("템플릿 워드 파일(template.docx)을 templates 폴더에서 찾을 수 없습니다.", "danger")
            return redirect(url_for('admin_dashboard'))

        # 1) docxtpl로 템플릿 렌더링 (원본 디자인 유지)
        # ※ 발제문 내용에 '<책제목>' 같이 꺾쇠괄호가 들어가면 Word XML 구조를 깨뜨려서
        #   해당 지점 이후 렌더링이 중단되고 파일이 손상됨.
        #   docxtpl의 자동 escape이 production 환경에서 일관되지 않게 동작하므로,
        #   사전에 풀-와이드 유니코드(〈, 〉)로 치환해서 시각적으로는 동일하지만 안전하게 처리.
        def _safe(s):
            if s is None:
                return ''
            return (str(s)
                    .replace('<', '〈')
                    .replace('>', '〉'))

        doc = DocxTemplate(template_path)
        date_str = (event.get('meeting_date') or '').replace('-', '.')
        template_submissions = []
        for sub in submissions:
            identity = topic_submitter_identity(sub)
            template_submissions.append({
                # 기존 템플릿의 department + author_name 배치를 유지하면서
                # '전자공학부 22 박민서' 형태로 표시한다.
                'department': _safe(identity['department_and_year']),
                'author_name': _safe(identity['author_name']),
                'topics': [
                    {
                        'topic': _safe(t.get('topic', '')),
                        'page': _safe(t.get('page', '')),
                        'reference': _safe(t.get('reference', '')),
                        'number': t.get('number'),
                    }
                    for t in (sub.get('topics') or [])
                ],
            })
        context = {
            'book_title': _safe(event.get('book_title', '')),
            'meeting_date': date_str,
            'book_author': _safe(event.get('book_author', '')),
            'moderator_name': _safe(event['moderator_name']),
            'submissions': template_submissions,
        }
        doc.render(context, autoescape=True)

        # 메모리에 저장
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        # 2) 렌더링 검증 — 모든 제출자가 본문에 들어갔는지 확인 (누락 방지)
        import zipfile, re
        buf.seek(0)
        with zipfile.ZipFile(BytesIO(buf.getvalue())) as zin:
            body_xml = zin.read('word/document.xml').decode('utf-8', errors='replace')
        # 본문 plaintext만 추출
        body_text = ''.join(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', body_xml))
        # 각 제출자가 '발제자' 섹션에 들어갔는지 검사: 이름 등장 횟수가 2 이상이면 OK
        # (1회: 참석자 명단, 2회 이상: 발제자 섹션 포함)
        missing = []
        for sub in submissions:
            nm = (sub.get('author_name') or '').strip()
            dept = (sub.get('department') or '').strip()
            if not nm:
                continue
            # 이름이 '발제자' 섹션에 등장하는지 검사: '발제자' 텍스트 이후에 이름이 등장하는지
            # 간단하게 본문 전체에서 이름 등장 횟수로 판별
            cnt = body_text.count(nm)
            if cnt < 2:  # 참석자 명단 1회만 있고 발제 섹션엔 없음
                missing.append(sub)
        if missing:
            app.logger.warning(f"download_topics_word: 누락 감지 {len(missing)}건 → python-docx로 append")
            # 3) 누락된 제출을 문서 끝에 직접 추가
            doc2 = Document(BytesIO(buf.getvalue()))
            for sub in missing:
                # 페이지 나눔
                br_p = doc2.add_paragraph()
                br_p.add_run().add_break(WD_BREAK.PAGE)
                # 발제자 헤더
                hp = doc2.add_paragraph()
                r1 = hp.add_run('발제자: ')
                r1.bold = True
                r1.font.size = Pt(13)
                r1.font.color.rgb = RGBColor(0, 102, 204)
                r2 = hp.add_run(topic_submitter_identity(sub)['full_label'])
                r2.bold = True
                r2.font.size = Pt(13)
                # 발제 내용
                for t in sub.get('topics') or []:
                    topic_text = (t.get('topic') or '').strip()
                    page = (t.get('page') or '').strip()
                    reference = (t.get('reference') or '').strip()
                    p = doc2.add_paragraph()
                    rn = p.add_run(f"{t.get('number')}. "); rn.bold = True; rn.font.size = Pt(11)
                    first = True
                    for line in topic_text.split('\n'):
                        if not first:
                            p.add_run().add_break()
                        rr = p.add_run(line); rr.font.size = Pt(11); first = False
                    if page or reference:
                        meta = doc2.add_paragraph()
                        parts = []
                        if page: parts.append(f"페이지: {page}")
                        if reference: parts.append(f"참조: {reference}")
                        rm = meta.add_run('   ' + ' | '.join(parts))
                        rm.italic = True; rm.font.size = Pt(10)
                        rm.font.color.rgb = RGBColor(102, 102, 102)
            buf = BytesIO()
            doc2.save(buf)
            buf.seek(0)

        # === 최종 폰트 일괄 적용: 함초롱바탕 ===
        # 한글은 w:eastAsia 속성까지 함께 지정해야 Word 에서 정확히 표시됨.
        try:
            from docx import Document as _Doc
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            TARGET_FONT = '함초롱바탕'

            def _apply_font(run, name):
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'),    name)
                rFonts.set(qn('w:hAnsi'),    name)
                rFonts.set(qn('w:eastAsia'), name)
                rFonts.set(qn('w:cs'),       name)

            _final = _Doc(BytesIO(buf.getvalue()))
            for p in _final.paragraphs:
                for r in p.runs:
                    _apply_font(r, TARGET_FONT)
            for table in _final.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                _apply_font(r, TARGET_FONT)
            buf = BytesIO()
            _final.save(buf)
            buf.seek(0)
        except Exception as fe:
            app.logger.warning(f"폰트 적용 실패(무시 가능): {fe}")

        filename = f"발제문_{event.get('book_title','')}_{event.get('meeting_date','')}.docx"
        app.logger.info(f"download_topics_word: 완료, 크기={len(buf.getvalue())} bytes")
        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-disposition": f"attachment; filename={filename.encode('utf-8').decode('latin1')}"}
        )
    except Exception as e:
        app.logger.error(f"download_topics_word error: {e}", exc_info=True)
        flash(f"문서 생성 중 오류 발생: {str(e)}", "danger")
        return redirect(url_for('admin_dashboard'))

# ==============================================================================
# --- 6.5 세미나 출석 투표 (학기 단위) ---
# ==============================================================================

KST = timezone(timedelta(hours=9))


def _parse_db_ts(val):
    """Supabase timestamptz 문자열 → KST datetime."""
    if not val:
        return None
    if isinstance(val, datetime):
        return val.astimezone(KST) if val.tzinfo else val.replace(tzinfo=KST)
    try:
        s = str(val).replace('Z', '+00:00')
        # Postgres가 "2026-05-10 23:59:59+00" 같은 형식을 줄 수도 있으므로 보정
        if 'T' not in s and ' ' in s:
            s = s.replace(' ', 'T', 1)
        dt = datetime.fromisoformat(s)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=KST)
        return dt.astimezone(KST)
    except Exception:
        return None


def _default_voting_window(meeting_date):
    """기본 규칙(전주 금 18:00 ~ 전주 일 23:59:59 KST)."""
    if isinstance(meeting_date, str):
        d = date.fromisoformat(meeting_date)
    else:
        d = meeting_date
    monday = d - timedelta(days=d.weekday())
    friday_before = monday - timedelta(days=3)
    sunday_before = monday - timedelta(days=1)
    open_at = datetime.combine(friday_before, time(18, 0), tzinfo=KST)
    close_at = datetime.combine(sunday_before, time(23, 59, 59), tzinfo=KST)
    return open_at, close_at


def _voting_window_for(session_or_date):
    """세미나 회차의 투표 오픈/마감 시각 반환.
    - session dict가 들어오면 vote_open_at/vote_close_at(관리자 지정값)이 있으면 그걸 우선.
    - 둘 다 없거나 date/str만 들어오면 기본 규칙(전주 금 18:00 ~ 전주 일 23:59:59 KST).
    """
    if isinstance(session_or_date, dict):
        meeting_date = session_or_date.get('meeting_date')
        custom_open = _parse_db_ts(session_or_date.get('vote_open_at'))
        custom_close = _parse_db_ts(session_or_date.get('vote_close_at'))
    else:
        meeting_date = session_or_date
        custom_open = None
        custom_close = None
    default_open, default_close = _default_voting_window(meeting_date)
    return (custom_open or default_open), (custom_close or default_close)


def _is_voting_open(session_or_date):
    open_at, close_at = _voting_window_for(session_or_date)
    now = datetime.now(KST)
    return open_at <= now <= close_at


def _enumerate_mon_thu(start_date, end_date):
    """start_date~end_date(둘 다 포함) 사이 모든 월/목 날짜를 (date, day_type) 리스트로 반환."""
    if isinstance(start_date, str):
        start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
    if isinstance(end_date, str):
        end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
    result = []
    d = start_date
    while d <= end_date:
        # weekday(): 월=0, 목=3
        if d.weekday() == 0:
            result.append((d, 'mon'))
        elif d.weekday() == 3:
            result.append((d, 'thu'))
        d += timedelta(days=1)
    return result


def _week_start(value):
    return cycle_monday(value)


def _ensure_term_weeks(term_id, dates):
    """주차를 먼저 보장하고 week_start -> week id 매핑을 반환한다."""
    starts = sorted({_week_start(d).isoformat() for d, _ in dates})
    existing = supabase.table('seminar_weeks').select('id, week_start') \
        .eq('term_id', term_id).execute().data or []
    by_start = {row['week_start']: row['id'] for row in existing}
    missing = [{'term_id': term_id, 'week_start': value} for value in starts if value not in by_start]
    if missing:
        supabase.table('seminar_weeks').insert(missing).execute()
        existing = supabase.table('seminar_weeks').select('id, week_start') \
            .eq('term_id', term_id).execute().data or []
        by_start = {row['week_start']: row['id'] for row in existing}
    return by_start


def _session_payload(term_id, meeting_date, day_type, weeks):
    week_id = weeks[_week_start(meeting_date).isoformat()]
    return {
        'term_id': term_id,
        'seminar_week_id': week_id,
        'meeting_date': meeting_date.isoformat(),
        'day_type': day_type,
        'participation_mode': 'opt_in' if day_type == 'mon' else 'absence_only',
        'capacity': None,  # 카카오톡 참석 투표: 정원·선착순 제한 없음
    }


@app.route('/api/admin/seminar_terms/create', methods=['POST'])
@login_required(role="admin")
def seminar_term_create():
    try:
        data = request.json or request.form
        name = (data.get('name') or '').strip()
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        max_capacity = int(data.get('max_capacity') or 32)
        if not (name and start_date and end_date):
            return jsonify({'status': 'error', 'message': '학기명/시작일/종료일은 필수입니다.'}), 400

        term_res = supabase.table('seminar_terms').insert({
            'name': name,
            'start_date': start_date,
            'end_date': end_date,
            'max_capacity': max_capacity,
            'is_active': True,
        }).execute()
        term = term_res.data[0]

        dates = _enumerate_mon_thu(start_date, end_date)
        weeks = _ensure_term_weeks(term['id'], dates)
        sessions_payload = [_session_payload(term['id'], d, dt, weeks) for d, dt in dates]
        if sessions_payload:
            supabase.table('seminar_sessions').insert(sessions_payload).execute()

        share_url = f"{request.host_url}seminar_vote?token={term['share_token']}"
        return jsonify({
            'status': 'success',
            'term': term,
            'session_count': len(sessions_payload),
            'share_url': share_url,
        })
    except Exception as e:
        app.logger.error(f"seminar_term_create error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_terms/<term_id>/update', methods=['POST'])
@login_required(role="admin")
def seminar_term_update(term_id):
    try:
        data = request.json or request.form
        update = {}
        for key in ('name', 'start_date', 'end_date'):
            if data.get(key) is not None:
                update[key] = data.get(key)
        if data.get('max_capacity') is not None:
            update['max_capacity'] = int(data.get('max_capacity'))
        if data.get('is_active') is not None:
            v = data.get('is_active')
            update['is_active'] = v if isinstance(v, bool) else str(v).lower() in ('true', '1', 'on', 'yes')
        if update:
            supabase.table('seminar_terms').update(update).eq('id', term_id).execute()

        # 기간이 변경된 경우, 누락된 회차만 추가 (기존 회차는 보존)
        if 'start_date' in update or 'end_date' in update:
            term_res = supabase.table('seminar_terms').select('start_date, end_date').eq('id', term_id).single().execute()
            t = term_res.data
            existing_res = supabase.table('seminar_sessions').select('meeting_date').eq('term_id', term_id).execute()
            existing_dates = {row['meeting_date'] for row in (existing_res.data or [])}
            dates = _enumerate_mon_thu(t['start_date'], t['end_date'])
            weeks = _ensure_term_weeks(term_id, dates)
            to_insert = []
            for d, dt in dates:
                if d.isoformat() not in existing_dates:
                    to_insert.append(_session_payload(term_id, d, dt, weeks))
            if to_insert:
                supabase.table('seminar_sessions').insert(to_insert).execute()

        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"seminar_term_update error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_terms/<term_id>/delete', methods=['POST'])
@login_required(role="admin")
def seminar_term_delete(term_id):
    try:
        supabase.table('seminar_terms').update({'is_active': False}).eq('id', term_id).execute()
        supabase.table('seminar_sessions').update({'is_active': False}).eq('term_id', term_id).execute()
        return jsonify({'status': 'success', 'archived': True})
    except Exception as e:
        app.logger.error(f"seminar_term_delete error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/toggle_active', methods=['POST'])
@login_required(role="admin")
def seminar_session_toggle(session_id):
    try:
        cur = supabase.table('seminar_sessions').select('is_active').eq('id', session_id).single().execute().data
        new_val = not bool(cur.get('is_active'))
        supabase.table('seminar_sessions').update({'is_active': new_val}).eq('id', session_id).execute()
        return jsonify({'status': 'success', 'is_active': new_val})
    except Exception as e:
        app.logger.error(f"seminar_session_toggle error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/admin/seminar_terms/<term_id>')
@login_required(role="admin")
def admin_seminar_term(term_id):
    return redirect(url_for('admin_seminars', term_id=term_id))
    try:
        term = supabase.table('seminar_terms').select('*').eq('id', term_id).single().execute().data
        sessions = supabase.table('seminar_sessions').select('*').eq('term_id', term_id) \
            .order('meeting_date').execute().data or []
        session_ids = [s['id'] for s in sessions]
        votes = []
        if session_ids:
            votes = supabase.table('seminar_votes').select('session_id, member_id, attending') \
                .in_('session_id', session_ids).execute().data or []
        # 회차별 집계 + 참석자 멤버 ID 목록
        agg = {sid: {'yes': 0, 'no': 0, 'attendee_ids': []} for sid in session_ids}
        for v in votes:
            sid = v['session_id']
            if v['attending']:
                agg[sid]['yes'] += 1
                agg[sid]['attendee_ids'].append(v['member_id'])
            else:
                agg[sid]['no'] += 1
        # 멤버 이름 매핑
        member_ids = {mid for a in agg.values() for mid in a['attendee_ids']}
        member_map = {}
        if member_ids:
            mres = supabase.table('members').select('id, name').in_('id', list(member_ids)).execute().data or []
            member_map = {m['id']: m['name'] for m in mres}

        # 진행 기록(history)과 연동: 같은 날짜의 history row 존재 여부 + id
        session_dates = [s['meeting_date'] for s in sessions]
        history_map = {}
        history_session_map = {}
        if session_dates:
            hres = supabase.table('history').select('id, date, book_title, genre, groups, facilitators, seminar_session_id') \
                .in_('date', session_dates).execute().data or []
            for h in hres:
                history_map[h['date']] = h
                if h.get('seminar_session_id'):
                    history_session_map[h['seminar_session_id']] = h

        # 발제문 이벤트/제출 현황 연동
        topic_map = {}
        if session_ids:
            topic_rows = supabase.table('topic_events').select('*') \
                .in_('seminar_session_id', session_ids).execute().data or []
            topic_rows = _auto_close_topic_events(topic_rows)
            topic_ids = [row['id'] for row in topic_rows]
            submission_counts = {}
            if topic_ids:
                submission_rows = supabase.table('topic_submissions').select('event_id') \
                    .in_('event_id', topic_ids).execute().data or []
                for submission in submission_rows:
                    eid = submission['event_id']
                    submission_counts[eid] = submission_counts.get(eid, 0) + 1
            for topic in topic_rows:
                topic['submission_count'] = submission_counts.get(topic['id'], 0)
                topic['share_url'] = f"{request.host_url}shared_topics?token={topic['share_token']}"
                topic_map[topic['seminar_session_id']] = topic

        today_kst = datetime.now(KST).date()
        upcoming_sessions, past_sessions = [], []
        for s in sessions:
            a = agg.get(s['id'], {'yes': 0, 'no': 0, 'attendee_ids': []})
            s['yes_count'] = a['yes']
            s['no_count'] = a['no']
            s['attendees'] = [{'id': mid, 'name': member_map.get(mid, f"id={mid}")} for mid in a['attendee_ids']]
            open_at, close_at = _voting_window_for(s)
            s['voting_open_at'] = open_at.strftime('%Y-%m-%d %H:%M')
            s['voting_close_at'] = close_at.strftime('%Y-%m-%d %H:%M')
            s['voting_open_at_input']  = open_at.strftime('%Y-%m-%dT%H:%M')
            s['voting_close_at_input'] = close_at.strftime('%Y-%m-%dT%H:%M')
            s['voting_custom'] = bool(s.get('vote_open_at') or s.get('vote_close_at'))
            now_kst = datetime.now(KST)
            if now_kst < open_at:
                s['vote_status'] = 'upcoming'
            elif now_kst <= close_at:
                s['vote_status'] = 'open'
            else:
                s['vote_status'] = 'closed'

            # 기록 연동
            h = history_session_map.get(s['id']) or history_map.get(s['meeting_date'])
            s['history_id'] = h['id'] if h else None
            if h and not s.get('book_title') and h.get('book_title'):
                s['book_title'] = h['book_title']
            s['topic_event'] = topic_map.get(s['id'])

            # 과거 회차 분리
            try:
                m_date = date.fromisoformat(s['meeting_date'])
            except Exception:
                m_date = today_kst
            s['is_past'] = m_date < today_kst
            if s['is_past']:
                past_sessions.append(s)
            else:
                upcoming_sessions.append(s)

        # 전체 활성 멤버 (관리자 수동 추가용)
        all_members = supabase.table('members').select('id, name, student_id, department') \
            .eq('is_active', True).order('name').execute().data or []

        share_url = f"{request.host_url}seminar_vote?token={term['share_token']}"
        return render_template('admin_seminar_term.html', term=term, sessions=sessions,
                               upcoming_sessions=upcoming_sessions, past_sessions=past_sessions,
                               share_url=share_url, all_members=all_members)
    except Exception as e:
        app.logger.error(f"admin_seminar_term error: {e}", exc_info=True)
        flash(f"학기 정보를 불러오는 중 오류: {e}", "danger")
        return redirect(url_for('admin_dashboard'))


def _load_weekly_seminar_view(term_id=None):
    terms = supabase.table('seminar_terms').select('*').order('start_date', desc=True).execute().data or []
    if not terms:
        return terms, None, [], []
    term = next((row for row in terms if str(row['id']) == str(term_id)), None)
    if term is None:
        term = next((row for row in terms if row.get('is_active')), terms[0])

    weeks = supabase.table('seminar_weeks').select('*').eq('term_id', term['id']) \
        .order('week_start').execute().data or []
    week_ids = [row['id'] for row in weeks]
    sessions = []
    topics = []
    if week_ids:
        sessions = supabase.table('seminar_sessions').select('*').in_('seminar_week_id', week_ids) \
            .order('meeting_date').execute().data or []
        topics = supabase.table('topic_events').select('*').in_('seminar_week_id', week_ids).execute().data or []
        topics = _auto_close_topic_events(topics)

    session_ids = [row['id'] for row in sessions]
    topic_ids = [row['id'] for row in topics]
    votes = []
    absences = []
    no_shows = []
    histories = []
    review_forms = []
    review_submissions = []
    if session_ids:
        votes = supabase.table('seminar_votes').select('session_id, member_id, attending, added_by_admin') \
            .in_('session_id', session_ids).execute().data or []
        absences = supabase.table('seminar_absences').select('*') \
            .in_('session_id', session_ids).order('created_at').execute().data or []
        no_shows = supabase.table('seminar_no_shows').select('*') \
            .in_('session_id', session_ids).order('created_at').execute().data or []
        histories = supabase.table('history').select('id, date, book_title, groups, seminar_session_id') \
            .in_('seminar_session_id', session_ids).execute().data or []
        review_forms = supabase.table('seminar_review_forms').select('*') \
            .in_('seminar_session_id', session_ids).execute().data or []
        review_form_ids = [row['id'] for row in review_forms]
        if review_form_ids:
            review_submissions = supabase.table('seminar_reviews').select(
                'id, form_id, member_id, memorable_point, discussion_point, free_text, created_at, updated_at'
            ) \
                .in_('form_id', review_form_ids).is_('deleted_at', 'null').execute().data or []
    submissions = []
    if topic_ids:
        submissions = supabase.table('topic_submissions').select('event_id').in_('event_id', topic_ids).execute().data or []

    all_members = supabase.table('members').select('id, name, student_id, department, is_active') \
        .order('name').execute().data or []
    active_members = [row for row in all_members if row.get('is_active')]
    member_by_id = {row['id']: row for row in all_members}
    topic_by_week = {row['seminar_week_id']: row for row in topics if row.get('seminar_week_id')}
    submission_count = defaultdict(int)
    for row in submissions:
        submission_count[row['event_id']] += 1
    history_by_session = {row['seminar_session_id']: row for row in histories if row.get('seminar_session_id')}
    review_count = defaultdict(int)
    reviews_by_form = defaultdict(list)
    for row in review_submissions:
        review_count[row['form_id']] += 1
        row['member'] = member_by_id.get(row.get('member_id'), {})
        row['review_content'] = '\n\n'.join(filter(None, [
            row.get('memorable_point'),
            row.get('discussion_point'),
            row.get('free_text'),
        ]))
        reviews_by_form[row['form_id']].append(row)
    review_by_session = {}
    for row in review_forms:
        row['submission_count'] = review_count[row['id']]
        row['submissions'] = sorted(
            reviews_by_form[row['id']],
            key=lambda item: item.get('updated_at') or item.get('created_at') or '',
            reverse=True,
        )
        row['share_url'] = f"{request.host_url}review/seminar/{row['share_token']}"
        review_by_session[row['seminar_session_id']] = row
    votes_by_session = defaultdict(list)
    for row in votes:
        if row.get('attending'):
            member = member_by_id.get(row['member_id'])
            if member:
                votes_by_session[row['session_id']].append(member)
    absence_audit_by_session = defaultdict(list)
    for row in absences:
        member = member_by_id.get(row['member_id'])
        if member:
            absence_audit_by_session[row['session_id']].append({**row, 'member': member})
    no_show_audit_by_session = defaultdict(list)
    for row in no_shows:
        member = member_by_id.get(row['member_id'])
        if member:
            no_show_audit_by_session[row['session_id']].append({**row, 'member': member})

    sessions_by_week = defaultdict(list)
    today = datetime.now(KST).date()
    for item in sessions:
        item['attendees'] = sorted(votes_by_session[item['id']], key=lambda row: row.get('name') or '')
        item['yes_count'] = len(item['attendees'])
        audit = absence_audit_by_session[item['id']]
        item['absence_audit'] = audit
        item['absences'] = [row for row in audit if not row.get('cancelled_at')]
        item['absent_member_ids'] = [row['member_id'] for row in item['absences']]
        no_show_audit = no_show_audit_by_session[item['id']]
        item['no_show_audit'] = no_show_audit
        item['no_shows'] = [row for row in no_show_audit if not row.get('cancelled_at')]
        item['no_show_member_ids'] = [row['member_id'] for row in item['no_shows']]
        if item.get('participation_mode') == 'opt_in':
            no_show_pool = item['attendees']
        else:
            no_show_pool = [
                member for member in active_members
                if member['id'] not in item['absent_member_ids']
            ]
        item['no_show_candidates'] = [
            member for member in no_show_pool
            if member['id'] not in item['no_show_member_ids']
        ]
        item['expected_count'] = max(0, len(active_members) - len(item['absences'])) \
            if item.get('participation_mode') == 'absence_only' else item['yes_count']
        if item.get('planned_member_ids') is not None:
            planned_set = set(item['planned_member_ids'])
            item['attendees'] = [member for member in all_members if member['id'] in planned_set]
            item['yes_count'] = item['expected_count'] = len(item['attendees'])
            item['no_show_candidates'] = [member for member in item['attendees']
                                          if member['id'] not in item['no_show_member_ids']]
        item['history'] = history_by_session.get(item['id'])
        item['review_form'] = review_by_session.get(item['id'])
        item['is_past'] = date.fromisoformat(item['meeting_date']) < today
        sessions_by_week[item['seminar_week_id']].append(item)

    for week in weeks:
        week['sessions'] = sessions_by_week.get(week['id'], [])
        week['sessions_by_day'] = {row['day_type']: row for row in week['sessions']}
        topic = topic_by_week.get(week['id'])
        if topic:
            topic['submission_count'] = submission_count[topic['id']]
            topic['share_url'] = f"{request.host_url}shared_topics?token={topic['share_token']}"
        week['topic_event'] = topic
        week['is_past'] = bool(week['sessions']) and all(row['is_past'] for row in week['sessions'])
    return terms, term, weeks, active_members


@app.route('/admin/seminars')
@login_required(role="admin")
def admin_seminars():
    try:
        terms, term, weeks, active_members = _load_weekly_seminar_view(request.args.get('term_id'))
        weeks = [row for row in weeks if not row['is_past']] + list(reversed([row for row in weeks if row['is_past']]))
        share_url = f"{request.host_url}seminar_vote?token={term['share_token']}" if term else ''
        return render_template('admin_seminars.html', terms=terms, term=term, weeks=weeks,
                               all_members=active_members, share_url=share_url)
    except Exception as e:
        app.logger.error(f"admin_seminars error: {e}", exc_info=True)
        flash(f"세미나 운영 정보를 불러오는 중 오류가 발생했습니다: {e}", 'danger')
        return redirect(url_for('admin_dashboard'))


@app.route('/seminars')
@login_required(role="ANY")
def seminars():
    try:
        terms, term, weeks, _ = _load_weekly_seminar_view(request.args.get('term_id'))
        weeks = [row for row in weeks if not row['is_past']] + list(reversed([row for row in weeks if row['is_past']]))
        vote_url = f"{request.host_url}seminar_vote?token={term['share_token']}" if term else ''
        return render_template('seminars.html', terms=terms, term=term, weeks=weeks, vote_url=vote_url)
    except Exception as e:
        app.logger.error(f"seminars error: {e}", exc_info=True)
        flash('세미나 일정을 불러오지 못했습니다.', 'danger')
        return redirect(url_for('main_index'))


@app.route('/api/admin/seminar_sessions/<session_id>/update_book', methods=['POST'])
@login_required(role="admin")
def seminar_session_update_book(session_id):
    try:
        data = request.json or {}
        book_title = (data.get('book_title') or '').strip()
        book_author = (data.get('book_author') or '').strip()
        update = {'book_title': book_title or None, 'book_author': book_author or None}
        current = supabase.table('seminar_sessions').select('seminar_week_id') \
            .eq('id', session_id).single().execute().data or {}
        week_id = current.get('seminar_week_id')
        if week_id:
            supabase.table('seminar_weeks').update({
                **update, 'needs_review': False, 'updated_at': datetime.now(timezone.utc).isoformat()
            }).eq('id', week_id).execute()
            supabase.table('seminar_sessions').update(update).eq('seminar_week_id', week_id).execute()
            supabase.table('topic_events').update(update).eq('seminar_week_id', week_id).execute()
            linked_sessions = supabase.table('seminar_sessions').select('id') \
                .eq('seminar_week_id', week_id).execute().data or []
            session_ids = [row['id'] for row in linked_sessions]
            if session_ids:
                supabase.table('history').update({'book_title': book_title or None}) \
                    .in_('seminar_session_id', session_ids).execute()
        else:
            supabase.table('seminar_sessions').update(update).eq('id', session_id).execute()
            supabase.table('topic_events').update(update).eq('seminar_session_id', session_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"seminar_session_update_book error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/update_voting_window', methods=['POST'])
@login_required(role="admin")
def seminar_session_update_voting_window(session_id):
    """관리자: 회차별 투표 오픈/마감 시각을 직접 지정하거나 기본 규칙으로 되돌림.
    body: { vote_open_at: 'YYYY-MM-DDTHH:MM' | '', vote_close_at: 'YYYY-MM-DDTHH:MM' | '',
            reset: bool }
    빈 문자열이거나 reset=True면 NULL로 설정 → 기본 규칙으로 fallback.
    """
    try:
        data = request.json or {}

        def parse_local_kst(val):
            if not val:
                return None
            s = str(val).strip()
            if not s:
                return None
            # datetime-local input: 'YYYY-MM-DDTHH:MM' (초 없음일 수 있음)
            try:
                if len(s) == 16:
                    s += ':00'
                dt = datetime.fromisoformat(s)
            except Exception:
                return None
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=KST)
            return dt.astimezone(timezone.utc).isoformat()

        if data.get('reset'):
            new_open, new_close = None, None
        else:
            new_open = parse_local_kst(data.get('vote_open_at'))
            new_close = parse_local_kst(data.get('vote_close_at'))
            if new_open and new_close and new_open >= new_close:
                return jsonify({'status': 'error', 'message': '오픈 시각은 마감 시각보다 빨라야 합니다.'}), 400

        supabase.table('seminar_sessions').update({
            'vote_open_at': new_open,
            'vote_close_at': new_close,
        }).eq('id', session_id).execute()

        # 새 윈도우를 다시 계산해서 반환 (NULL이면 기본 규칙으로)
        sess = supabase.table('seminar_sessions') \
            .select('meeting_date, vote_open_at, vote_close_at') \
            .eq('id', session_id).single().execute().data or {}
        open_at, close_at = _voting_window_for(sess)
        return jsonify({
            'status': 'success',
            'voting_open_at': open_at.strftime('%Y-%m-%d %H:%M'),
            'voting_close_at': close_at.strftime('%Y-%m-%d %H:%M'),
            'voting_open_at_input':  open_at.strftime('%Y-%m-%dT%H:%M'),
            'voting_close_at_input': close_at.strftime('%Y-%m-%dT%H:%M'),
            'voting_custom': bool(sess.get('vote_open_at') or sess.get('vote_close_at')),
        })
    except Exception as e:
        app.logger.error(f"seminar_session_update_voting_window error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/add_attendee', methods=['POST'])
@login_required(role="admin")
def seminar_session_add_attendee(session_id):
    return jsonify({'status': 'error', 'message': '카카오톡 명단 반영 화면에서 참석 예정 명단을 수정해주세요.'}), 410


@app.route('/api/admin/seminar_sessions/<session_id>/remove_attendee', methods=['POST'])
@login_required(role="admin")
def seminar_session_remove_attendee(session_id):
    return jsonify({'status': 'error', 'message': '카카오톡 명단 반영 화면에서 참석 예정 명단을 수정해주세요.'}), 410


@app.route('/seminar_vote')
def seminar_vote_page():
    return render_template('seminar_vote_retired.html')


@app.route('/api/seminar_vote/verify', methods=['POST'])
def seminar_vote_verify():
    return jsonify({'status': 'error', 'message': '참석 투표는 카카오톡에서 진행합니다. 사이트 신청은 종료되었습니다.'}), 410


@app.route('/api/seminar_vote/counts')
def seminar_vote_counts():
    return jsonify({'status': 'error', 'message': '참석 투표는 카카오톡에서 진행합니다.'}), 410


@app.route('/api/seminar_vote/submit', methods=['POST'])
def seminar_vote_submit():
    return jsonify({'status': 'error', 'message': '참석 투표는 카카오톡에서 진행합니다. 사이트 신청은 종료되었습니다.'}), 410


# ==============================================================================
# --- 6.6 통합 기록 시스템 (회원/세미나/벽돌책/소모임/장르) ---
# ==============================================================================

DEFAULT_GENRES = ['고전문학', '한국문학', '비문학', '시']


def _load_genres():
    try:
        res = supabase.table('genres').select('*') \
            .order('display_order').order('name').execute()
        return res.data or []
    except Exception as e:
        app.logger.warning(f"장르 로드 실패: {e}")
        return []


def _can_view_member_profile(member_id):
    if session.get('user_role') == 'admin':
        return True
    if session.get('user_id') == member_id:
        return True
    return False


def _member_special_events(member_id):
    """해당 멤버가 참여(attendee/organizer/speaker)한 스페셜 이벤트 목록.
    최신 행사 우선으로 정렬해 반환.
    """
    try:
        att_res = supabase.table('special_event_attendees') \
            .select('event_id, role, note') \
            .eq('member_id', member_id).execute()
        att_rows = att_res.data or []
        if not att_rows:
            return []
        ev_ids = list({r['event_id'] for r in att_rows})
        role_by_event = {r['event_id']: (r.get('role') or 'attendee') for r in att_rows}
        ev_res = supabase.table('special_events') \
            .select('id, name, description, event_date, end_date, category, location, is_active') \
            .in_('id', ev_ids).order('event_date', desc=True).execute()
        events = ev_res.data or []
        for e in events:
            e['my_role'] = role_by_event.get(e['id'], 'attendee')
        return events
    except Exception as e:
        app.logger.warning(f"_member_special_events error: {e}")
        return []


def _aggregate_member_activity(member_id, member_name):
    """한 멤버의 세미나/벽돌책/소모임 활동 집계."""
    history_res = supabase.table('history').select('id, date, book_title, genre, groups, facilitators') \
        .order('date', desc=True).execute()
    history = history_res.data or []
    seminar_records, facilitator_count = [], 0
    for row in history:
        groups = row.get('groups') or []
        present = set()
        for g in groups:
            if isinstance(g, list):
                present.update(g)
            elif isinstance(g, str):
                present.add(g)
        if member_name in present:
            is_fac = member_name in (row.get('facilitators') or [])
            if is_fac:
                facilitator_count += 1
            seminar_records.append({
                'history_id': row['id'],
                'date': row.get('date'),
                'book_title': row.get('book_title') or '',
                'genre': row.get('genre') or '',
                'is_facilitator': is_fac,
            })

    # 벽돌책 세션 참여
    try:
        bb_parts = supabase.table('brick_session_members').select(
            'session_id, brick_book_sessions(id, meeting_date, notes, brick_book_id, brick_books(id, title))'
        ).eq('member_id', member_id).execute().data or []
    except Exception:
        bb_parts = []
    brick_sessions = []
    for p in bb_parts:
        s = p.get('brick_book_sessions') or {}
        b = s.get('brick_books') or {}
        brick_sessions.append({
            'book_id': b.get('id'), 'book_title': b.get('title'),
            'session_id': s.get('id'), 'meeting_date': s.get('meeting_date'),
        })
    brick_sessions.sort(key=lambda x: x.get('meeting_date') or '', reverse=True)

    # 소모임 세션 참여
    try:
        sg_parts = supabase.table('study_session_members').select(
            'session_id, study_group_sessions(id, meeting_date, notes, study_group_id, study_groups(id, name))'
        ).eq('member_id', member_id).execute().data or []
    except Exception:
        sg_parts = []
    study_sessions = []
    for p in sg_parts:
        s = p.get('study_group_sessions') or {}
        g = s.get('study_groups') or {}
        study_sessions.append({
            'group_id': g.get('id'), 'group_name': g.get('name'),
            'session_id': s.get('id'), 'meeting_date': s.get('meeting_date'),
        })
    study_sessions.sort(key=lambda x: x.get('meeting_date') or '', reverse=True)

    return {
        'seminar_records': seminar_records,
        'seminar_count': len(seminar_records),
        'facilitator_count': facilitator_count,
        'brick_sessions': brick_sessions,
        'study_sessions': study_sessions,
    }


# --- 장르 API ---
@app.route('/api/genres')
@login_required(role="ANY")
def list_genres():
    return jsonify({'genres': _load_genres()})


@app.route('/api/admin/genres/create', methods=['POST'])
@login_required(role="admin")
def create_genre():
    try:
        name = ((request.json or {}).get('name') or '').strip()
        if not name:
            return jsonify({'status': 'error', 'message': '장르명을 입력하세요.'}), 400
        existing = supabase.table('genres').select('id').eq('name', name).execute().data or []
        if existing:
            return jsonify({'status': 'error', 'message': '이미 존재하는 장르입니다.'}), 400
        res = supabase.table('genres').insert({'name': name, 'is_default': False}).execute()
        return jsonify({'status': 'success', 'genre': res.data[0]})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/genres/<genre_id>/delete', methods=['POST'])
@login_required(role="admin")
def delete_genre(genre_id):
    try:
        cur = supabase.table('genres').select('is_default').eq('id', genre_id).single().execute().data
        if cur and cur.get('is_default'):
            return jsonify({'status': 'error', 'message': '기본 장르는 삭제할 수 없습니다.'}), 400
        supabase.table('genres').delete().eq('id', genre_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================================================
# 스페셜 이벤트 (MT, 워크숍, 강연 등 1회성 행사)
# ============================================================================

@app.route('/admin/special_events')
@login_required(role="admin")
def admin_special_events():
    """관리자: 스페셜 이벤트 목록 + 생성/관리 페이지"""
    try:
        events = supabase.table('special_events').select('*') \
            .order('event_date', desc=True).execute().data or []
        event_ids = [evt['id'] for evt in events]
        attendee_counts = {}
        if event_ids:
            attendee_rows = supabase.table('special_event_attendees') \
                .select('event_id').in_('event_id', event_ids).execute().data or []
            for attendee in attendee_rows:
                event_id = attendee.get('event_id')
                attendee_counts[event_id] = attendee_counts.get(event_id, 0) + 1
        for evt in events:
            evt['attendee_count'] = attendee_counts.get(evt['id'], 0)
        terms = supabase.table('seminar_terms').select('id, name, start_date, end_date') \
            .order('start_date', desc=True).execute().data or []
        return render_template('admin_special_events.html', events=events, terms=terms)
    except Exception as e:
        app.logger.error(f"admin_special_events error: {e}")
        flash("스페셜 이벤트 페이지 로드 중 오류.", "danger")
        return redirect(url_for('admin_dashboard'))


@app.route('/admin/special_events/<event_id>')
@login_required(role="admin")
def admin_special_event_detail(event_id):
    """관리자: 단일 이벤트 상세 + 참석자 관리"""
    try:
        evt = supabase.table('special_events').select('*').eq('id', event_id).single().execute().data
        if not evt:
            flash("이벤트를 찾을 수 없습니다.", "danger")
            return redirect(url_for('admin_special_events'))
        attendees = supabase.table('special_event_attendees') \
            .select('id, member_id, role, note, members(id, name, student_id, department)') \
            .eq('event_id', event_id).order('created_at').execute().data or []
        all_members = supabase.table('members').select('id, name, student_id, department') \
            .eq('is_active', True).order('name').execute().data or []
        attendee_ids = {a['member_id'] for a in attendees}
        candidates = [m for m in all_members if m['id'] not in attendee_ids]
        return render_template('admin_special_event_detail.html',
                               event=evt, attendees=attendees, candidates=candidates)
    except Exception as e:
        app.logger.error(f"admin_special_event_detail error: {e}")
        flash("이벤트 상세 로드 중 오류.", "danger")
        return redirect(url_for('admin_special_events'))


@app.route('/api/admin/special_events/create', methods=['POST'])
@login_required(role="admin")
def create_special_event():
    try:
        data = request.json or {}
        name = (data.get('name') or '').strip()
        event_date = data.get('event_date')
        if not name or not event_date:
            return jsonify({'error': '이름과 날짜는 필수입니다.'}), 400
        payload = {
            'name': name,
            'description': (data.get('description') or '').strip() or None,
            'event_date': event_date,
            'end_date': data.get('end_date') or None,
            'category': (data.get('category') or 'event'),
            'term_id': data.get('term_id') or None,
            'created_by': session.get('user_id'),
        }
        res = supabase.table('special_events').insert(payload).execute()
        return jsonify({'status': 'success', 'event': res.data[0]})
    except Exception as e:
        app.logger.error(f"create_special_event error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/special_events/<event_id>/update', methods=['POST'])
@login_required(role="admin")
def update_special_event(event_id):
    try:
        data = request.json or {}
        update = {}
        for k in ('name', 'description', 'event_date', 'end_date', 'category', 'term_id'):
            if k in data:
                update[k] = data[k] if data[k] != '' else None
        if 'is_active' in data:
            update['is_active'] = bool(data['is_active'])
        if not update:
            return jsonify({'error': '수정할 내용이 없습니다.'}), 400
        supabase.table('special_events').update(update).eq('id', event_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/special_events/<event_id>/toggle_active', methods=['POST'])
@login_required(role="admin")
def toggle_special_event(event_id):
    try:
        cur = supabase.table('special_events').select('is_active').eq('id', event_id).single().execute().data
        if not cur:
            return jsonify({'error': '이벤트를 찾을 수 없습니다.'}), 404
        new_state = not bool(cur.get('is_active'))
        supabase.table('special_events').update({'is_active': new_state}).eq('id', event_id).execute()
        return jsonify({'status': 'success', 'is_active': new_state})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/special_events/<event_id>/delete', methods=['POST'])
@login_required(role="admin")
def delete_special_event(event_id):
    try:
        # cascade로 참석자도 함께 삭제됨
        supabase.table('special_events').delete().eq('id', event_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/special_events/<event_id>/attendees/add', methods=['POST'])
@login_required(role="admin")
def add_special_event_attendee(event_id):
    try:
        data = request.json or {}
        member_ids = data.get('member_ids') or []
        if isinstance(member_ids, (int, str)):
            member_ids = [member_ids]
        if not member_ids:
            return jsonify({'error': '추가할 회원을 선택하세요.'}), 400
        rows = [{'event_id': event_id, 'member_id': int(mid),
                 'role': data.get('role', 'attendee'),
                 'note': (data.get('note') or '').strip() or None}
                for mid in member_ids]
        # ON CONFLICT DO NOTHING 동작 — 이미 등록된 멤버는 스킵
        supabase.table('special_event_attendees').upsert(rows, on_conflict='event_id,member_id').execute()
        return jsonify({'status': 'success', 'added': len(rows)})
    except Exception as e:
        app.logger.error(f"add_special_event_attendee error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/special_events/<event_id>/attendees/<int:member_id>/remove', methods=['POST'])
@login_required(role="admin")
def remove_special_event_attendee(event_id, member_id):
    try:
        supabase.table('special_event_attendees').delete() \
            .eq('event_id', event_id).eq('member_id', member_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============================================================================
# 출석 매트릭스 (관리자 도구)
# 가로축: 확정된 세미나 회차 날짜, 세로축: 회원, 셀: O/X
# 데이터 소스: history (확정된 세미나 진행 기록의 present 명단)
# ============================================================================

def _build_attendance_matrix(start_date=None, end_date=None):
    """날짜별 실제 출석표. 학기 인정 횟수는 별도 화면에서 주차 단위로 센다."""
    from attendance_workflow import build_term_attendance
    members = supabase.table('members').select('id, name, student_id, department, is_active') \
        .order('name').execute().data or []
    start, end = start_date or '0001-01-01', end_date or '9999-12-31'
    histories = supabase.table('history').select('id,date,present,book_title,seminar_session_id,attendance_confirmed_at,actual_member_ids') \
        .gte('date', start).lte('date', end).execute().data or []
    seminar_sessions = supabase.table('seminar_sessions').select('*') \
        .gte('meeting_date', start).lte('meeting_date', end).execute().data or []
    session_ids = [s['id'] for s in seminar_sessions]
    no_shows = supabase.table('seminar_no_shows').select('session_id,member_id,cancelled_at') \
        .in_('session_id', session_ids).is_('cancelled_at', 'null').execute().data or [] if session_ids else []
    report = build_term_attendance(members, histories, seminar_sessions, [], [], [], [], start, end,
                                   today=datetime.now(KST).date(), no_shows=no_shows)
    weekday_labels = ('월', '화', '수', '목', '금', '토', '일')
    for column in report['columns']:
        day = date.fromisoformat(column['date'])
        column['label'] = f"{day.month}/{day.day}({weekday_labels[day.weekday()]})"
        column['sort'] = day
    return members, report['columns'], report['matrix']


@app.route('/admin/attendance_matrix')
@login_required(role="admin")
def admin_attendance_matrix():
    start_date = request.args.get('start_date') or None
    end_date = request.args.get('end_date') or None
    term_id = request.args.get('term_id') or None

    # 학기 선택 시 학기 기간으로 자동 채움
    if term_id and not (start_date and end_date):
        try:
            t = supabase.table('seminar_terms').select('start_date, end_date') \
                .eq('id', term_id).single().execute().data
            if t:
                start_date = start_date or t['start_date']
                end_date = end_date or t['end_date']
        except Exception:
            pass

    try:
        members, weeks, matrix = _build_attendance_matrix(start_date, end_date)
        terms = supabase.table('seminar_terms').select('id, name, start_date, end_date') \
            .order('start_date', desc=True).execute().data or []

        # 회원별 출석 횟수 집계
        member_counts = {m['id']: sum(1 for w in weeks if matrix.get(m['id'], {}).get(w['key'])) for m in members}

        return render_template('admin_attendance_matrix.html',
                               members=members, weeks=weeks, matrix=matrix,
                               member_counts=member_counts,
                               terms=terms, term_id=term_id,
                               start_date=start_date or '', end_date=end_date or '')
    except Exception as e:
        app.logger.error(f"admin_attendance_matrix error: {e}")
        flash("매트릭스 로드 중 오류.", "danger")
        return redirect(url_for('admin_dashboard'))


@app.route('/admin/attendance_matrix/export')
@login_required(role="admin")
def admin_attendance_matrix_export():
    """엑셀(.xlsx)로 출석 매트릭스 내보내기"""
    start_date = request.args.get('start_date') or None
    end_date = request.args.get('end_date') or None
    term_id = request.args.get('term_id') or None
    if term_id and not (start_date and end_date):
        try:
            t = supabase.table('seminar_terms').select('start_date, end_date, name') \
                .eq('id', term_id).single().execute().data
            if t:
                start_date = start_date or t['start_date']
                end_date = end_date or t['end_date']
        except Exception:
            pass

    members, weeks, matrix = _build_attendance_matrix(start_date, end_date)

    # DataFrame 구성: 행=회원, 열=세미나 회차
    columns = ['이름', '학번', '소속'] + [f"{w['label']} {w['title']}" for w in weeks] + ['출석횟수']
    rows = []
    for m in members:
        row = [m.get('name', ''), m.get('student_id', '') or '', m.get('department', '') or '']
        cnt = 0
        for w in weeks:
            attended = bool(matrix.get(m['id'], {}).get(w['key']))
            row.append('O' if attended else 'X')
            if attended:
                cnt += 1
        row.append(cnt)
        rows.append(row)

    df = pd.DataFrame(rows, columns=columns)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='출석매트릭스', index=False)
    buf.seek(0)

    fname_parts = ['attendance_matrix']
    if start_date: fname_parts.append(start_date)
    if end_date: fname_parts.append(end_date)
    fname = '_'.join(fname_parts) + '.xlsx'

    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=fname)


# ==============================================================================
# --- 세미나실 예약 현황 (외부 게시판 캐시 기반) ---
# ==============================================================================
@app.route('/admin/seminar_rooms')
@login_required(role="admin")
def admin_seminar_rooms():
    """경북대 총동연 세미나실 예약 현황을 달력으로 보여준다.

    데이터는 seminar_room_posts 테이블에서 가져오며,
    실제 외부 게시판 크롤은 /api/admin/seminar_rooms/refresh 로 분리되어 있다.
    """
    today = datetime.now(timezone(timedelta(hours=9))).date()
    month_str = request.args.get('month', '')

    try:
        if month_str:
            year, month = map(int, month_str.split('-'))
            _ = date(year, month, 1)
        else:
            year, month = today.year, today.month
    except Exception:
        year, month = today.year, today.month

    first_day = date(year, month, 1)
    next_first = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)
    last_day = next_first - timedelta(days=1)

    from seminar_rooms import (
        compute_available_dates, extract_club_name, get_room_from_title,
        load_settings, parse_dates_from_title, WRITE_URL,
    )

    try:
        posts_res = supabase.table('seminar_room_posts').select('*').execute()
        posts = posts_res.data or []
    except Exception as e:
        app.logger.error(f"seminar_room_posts 조회 실패: {e}")
        posts = []
        flash(f"세미나실 캐시 조회 오류: {e}", "danger")

    # 예전 파서로 캐시된 글도 화면에서는 최신 규칙으로 즉시 보정한다.
    # 연도 없는 제목은 기존 캐시 날짜의 연도를 우선 사용해 해가 바뀌어도 안전하다.
    for post in posts:
        cached_dates = post.get('dates') or []
        parse_year = today.year
        if cached_dates:
            try:
                parse_year = date.fromisoformat(str(cached_dates[0])[:10]).year
            except (TypeError, ValueError):
                pass

        reparsed_dates = parse_dates_from_title(post.get('title') or '', parse_year)
        if reparsed_dates:
            post['dates'] = [item.isoformat() for item in reparsed_dates]
        reparsed_room = get_room_from_title(post.get('title') or '')
        if reparsed_room:
            post['room'] = reparsed_room
        reparsed_club = extract_club_name(post.get('title') or '')
        if reparsed_club:
            post['club_name'] = reparsed_club

    # 날짜 → 예약 목록 매핑 (월 한정 + 전체)
    by_date = defaultdict(list)        # 달력에 표시할 월 내 데이터
    all_by_date = defaultdict(list)    # 예약 가능 날짜 계산용 (전체)
    for p in posts:
        for d_str in (p.get('dates') or []):
            try:
                d = date.fromisoformat(d_str[:10])
            except Exception:
                continue
            all_by_date[d.isoformat()].append(p)
            if first_day <= d <= last_day:
                by_date[d.isoformat()].append(p)

    # 달력 그리드 (월요일 시작)
    grid_start = first_day - timedelta(days=first_day.weekday())
    grid_end = last_day + timedelta(days=(6 - last_day.weekday()))

    weeks = []
    cur = grid_start
    while cur <= grid_end:
        week = []
        for i in range(7):
            d = cur + timedelta(days=i)
            week.append({
                'date': d,
                'iso': d.isoformat(),
                'in_month': d.month == month,
                'is_today': d == today,
                'is_target': d.weekday() in (0, 3),  # 월/목 강조
                'posts': by_date.get(d.isoformat(), []),
            })
        weeks.append(week)
        cur += timedelta(days=7)

    prev_month = (first_day - timedelta(days=1)).strftime('%Y-%m')
    next_month = next_first.strftime('%Y-%m')

    last_check = None
    if posts:
        checked_values = [p.get('last_checked_at') for p in posts if p.get('last_checked_at')]
        if checked_values:
            last_check = max(checked_values)

    # 상태별 카운트(현재 월 한정)
    month_posts = []
    seen_ids = set()
    for d_iso, lst in by_date.items():
        for p in lst:
            if p['wr_id'] not in seen_ids:
                seen_ids.add(p['wr_id'])
                month_posts.append(p)
    counts = {'approved': 0, 'pending': 0, 'rejected': 0}
    for p in month_posts:
        counts[p.get('status', 'pending')] = counts.get(p.get('status', 'pending'), 0) + 1

    # 설정 로드 (DB) + 예약 가능 날짜 계산
    settings = load_settings(supabase)

    def _to_date(s):
        if not s:
            return None
        try:
            return date.fromisoformat(str(s)[:10])
        except Exception:
            return None

    sem_start = _to_date(settings.get('semester_start'))
    sem_end = _to_date(settings.get('semester_end'))

    try:
        d_min = int(settings.get('days_ahead_min') or 7)
        d_max = int(settings.get('days_ahead_max') or 28)
    except (TypeError, ValueError):
        d_min, d_max = 7, 28

    available_dates = compute_available_dates(
        all_by_date,
        today=today,
        days_ahead_min=d_min,
        days_ahead_max=d_max,
        semester_start=sem_start,
        semester_last=sem_end,
    )

    return render_template(
        'admin_seminar_rooms.html',
        weeks=weeks,
        year=year,
        month=month,
        prev_month=prev_month,
        next_month=next_month,
        today=today,
        last_check=last_check,
        total_cached=len(posts),
        month_counts=counts,
        available_dates=available_dates,
        settings=settings,
        write_url=WRITE_URL,
    )


@app.route('/api/admin/seminar_rooms/settings', methods=['POST'])
@login_required(role="admin")
def admin_seminar_rooms_settings_save():
    """세미나실 도우미 설정값(동아리명/전화/학기일자 등) 저장."""
    from seminar_rooms import save_settings
    payload = request.get_json(silent=True) or {}
    try:
        updated = save_settings(supabase, payload)
        return jsonify({'status': 'success', 'settings': updated})
    except Exception as e:
        app.logger.error(f"세미나실 설정 저장 실패: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_rooms/refresh', methods=['POST'])
@login_required(role="admin")
def admin_seminar_rooms_refresh():
    """외부 게시판을 크롤링하여 캐시를 갱신한다."""
    from seminar_rooms import crawl
    payload = request.get_json(silent=True) or {}
    pages = int(payload.get('pages') or 3)
    pages = max(1, min(pages, 10))
    try:
        result = crawl(supabase, max_pages=pages, recheck_pending=True,
                       logger=app.logger)
        return jsonify({'status': 'success', **result})
    except Exception as e:
        app.logger.error(f"세미나실 크롤 실패: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


# --- 학기 필터 헬퍼 ---
def _get_terms_for_filter():
    """기록 페이지의 학기 필터 드롭다운에 사용할 학기 목록 (최신순)."""
    try:
        return supabase.table('seminar_terms').select('id, name, start_date, end_date') \
            .order('start_date', desc=True).execute().data or []
    except Exception:
        return []


def _get_term_range(term_id):
    """선택된 term_id의 (start_date, end_date, term_dict) 반환. 없으면 (None, None, None)."""
    if not term_id:
        return None, None, None
    try:
        term = supabase.table('seminar_terms').select('*').eq('id', term_id).single().execute().data
        if not term:
            return None, None, None
        return term['start_date'], term['end_date'], term
    except Exception:
        return None, None, None


def _date_in_range(d, start, end):
    """문자열 또는 date를 받아 [start, end] 범위에 있는지 확인."""
    if not d or not start or not end:
        return False
    s = str(d)[:10]
    return str(start)[:10] <= s <= str(end)[:10]


# --- 도움말 페이지 ---
@app.route('/help/admin')
@login_required(role="admin")
def help_admin():
    return render_template('help_admin.html')


@app.route('/help/member')
@login_required(role="ANY")
def help_member():
    return render_template('help_member.html')


# --- 버그·개선 제보 ---
@app.route('/api/bug-reports', methods=['POST'])
@login_required(role="ANY")
def submit_bug_report():
    data = request.get_json(silent=True) or {}
    category = str(data.get('category') or 'bug').strip()
    title = str(data.get('title') or '').strip()
    description = str(data.get('description') or '').strip()
    source_page = str(data.get('source_page') or 'unknown').strip()[:160]
    if category not in ('bug', 'suggestion'):
        return jsonify({'status': 'error', 'message': '올바른 제보 종류를 선택해주세요.'}), 400
    if not 2 <= len(title) <= 120:
        return jsonify({'status': 'error', 'message': '제목은 2~120자로 적어주세요.'}), 400
    if not 10 <= len(description) <= 3000:
        return jsonify({'status': 'error', 'message': '내용은 10~3000자로 적어주세요.'}), 400
    try:
        inserted = supabase.table('bug_reports').insert({
            'reporter_id': session.get('user_id'),
            'reporter_name': session.get('user_name') or '이름 미상',
            'category': category,
            'title': title,
            'description': description,
            'source_page': source_page or 'unknown',
            'user_agent': (request.headers.get('User-Agent') or '')[:500] or None,
        }).execute().data or []
        return jsonify({'status': 'success', 'id': inserted[0]['id'] if inserted else None})
    except Exception as e:
        app.logger.error(f"submit_bug_report error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': '제보 저장 중 오류가 발생했습니다.'}), 500


@app.route('/admin/bug-reports')
@login_required(role="admin")
def admin_bug_reports():
    try:
        reports = supabase.table('bug_reports').select(
            'id, reporter_name, category, title, description, source_page, status, admin_note, created_at'
        ).order('created_at', desc=True).limit(200).execute().data or []
        counts = {'new': 0, 'reviewing': 0, 'resolved': 0}
        for report in reports:
            status = report.get('status')
            if status in counts:
                counts[status] += 1
        return render_template('admin_bug_reports.html', reports=reports, counts=counts)
    except Exception as e:
        app.logger.error(f"admin_bug_reports error: {e}", exc_info=True)
        flash('버그 제보 목록을 불러오지 못했습니다.', 'danger')
        return render_template('admin_bug_reports.html', reports=[], counts={
            'new': 0, 'reviewing': 0, 'resolved': 0
        })


@app.route('/api/admin/bug-reports/<report_id>', methods=['PATCH'])
@login_required(role="admin")
def update_bug_report(report_id):
    data = request.get_json(silent=True) or {}
    status = str(data.get('status') or '').strip()
    admin_note = str(data.get('admin_note') or '').strip()
    if status not in ('new', 'reviewing', 'resolved'):
        return jsonify({'status': 'error', 'message': '올바른 처리 상태를 선택해주세요.'}), 400
    if len(admin_note) > 1000:
        return jsonify({'status': 'error', 'message': '운영진 메모는 1000자 이내로 적어주세요.'}), 400
    now = datetime.now(timezone.utc).isoformat()
    try:
        supabase.table('bug_reports').update({
            'status': status,
            'admin_note': admin_note or None,
            'updated_at': now,
            'resolved_at': now if status == 'resolved' else None,
        }).eq('id', report_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"update_bug_report error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': '처리 상태 저장 중 오류가 발생했습니다.'}), 500


# --- 기록 허브 ---
@app.route('/records')
@login_required(role="admin")
def records_hub():
    try:
        seminar_count = supabase.table('history').select('id', count='exact').execute().count or 0
        bb_count = supabase.table('brick_books').select('id', count='exact').execute().count or 0
        sg_count = supabase.table('study_groups').select('id', count='exact').execute().count or 0
        member_count = supabase.table('members').select('id', count='exact').eq('is_active', True).execute().count or 0
    except Exception as e:
        app.logger.warning(f"records_hub count error: {e}")
        seminar_count = bb_count = sg_count = member_count = 0
    return render_template('records_hub.html',
                           seminar_count=seminar_count, brick_book_count=bb_count,
                           study_group_count=sg_count, member_count=member_count)


# --- 회원 관리 ---
@app.route('/admin/members')
@login_required(role="admin")
def records_members():
    try:
        members = supabase.table('members').select(
            'id, name, gender, department, student_id, recruiting_class, member_status, role, email'
        ).order('name').execute().data or []
        duplicate_student_ids = duplicate_student_id_values(members)
        for member in members:
            member['student_id_duplicate'] = (
                normalize_member_student_id(member.get('student_id')) in duplicate_student_ids
            )
        # 회원명부 클릭 시 미리 활동 요약을 보여주기 위해 각 멤버의 활동 카운트 집계
        if members:
            # 1) 세미나/발제 카운트 (history.groups 이름 매칭)
            history_rows = supabase.table('history').select('groups, facilitators').execute().data or []
            seminar_cnt, fac_cnt = {}, {}
            for r in history_rows:
                names = set()
                for g in (r.get('groups') or []):
                    if isinstance(g, list): names.update(g)
                    elif isinstance(g, str): names.add(g)
                facs = set(r.get('facilitators') or [])
                for n in names: seminar_cnt[n] = seminar_cnt.get(n, 0) + 1
                for n in facs: fac_cnt[n] = fac_cnt.get(n, 0) + 1
            # 2) 벽돌책 + 소모임 세션 카운트 (member_id 기반)
            try:
                bb_rows = supabase.table('brick_session_members').select('member_id').execute().data or []
            except Exception: bb_rows = []
            try:
                sg_rows = supabase.table('study_session_members').select('member_id').execute().data or []
            except Exception: sg_rows = []
            bb_cnt, sg_cnt = {}, {}
            for r in bb_rows: bb_cnt[r['member_id']] = bb_cnt.get(r['member_id'], 0) + 1
            for r in sg_rows: sg_cnt[r['member_id']] = sg_cnt.get(r['member_id'], 0) + 1
            name_counts = {}
            for member in members:
                member_name = member.get('name')
                name_counts[member_name] = name_counts.get(member_name, 0) + 1
            ambiguous_names = {name for name, count in name_counts.items() if name and count > 1}
            for m in members:
                is_ambiguous = m.get('name') in ambiguous_names
                m['activity_count_ambiguous'] = is_ambiguous
                m['seminar_count'] = None if is_ambiguous else seminar_cnt.get(m['name'], 0)
                m['facilitator_count'] = None if is_ambiguous else fac_cnt.get(m['name'], 0)
                m['brick_count'] = bb_cnt.get(m['id'], 0)
                m['study_count'] = sg_cnt.get(m['id'], 0)
    except Exception as e:
        app.logger.error(f"records_members error: {e}", exc_info=True)
        flash(f"멤버 로딩 오류: {e}", 'danger')
        members = []
        duplicate_student_ids = set()
    return render_template(
        'records_members.html',
        members=members,
        duplicate_student_id_count=len(duplicate_student_ids),
        can_manage_roles=_current_user_is_primary_admin(),
    )


@app.route('/records/members')
@login_required(role="admin")
def records_members_legacy():
    return redirect(url_for('records_members'), code=301)


@app.route('/records/members/<int:member_id>')
@login_required(role="ANY")
def records_member_profile(member_id):
    if not _can_view_member_profile(member_id):
        flash("접근 권한이 없습니다.", "danger")
        return redirect(url_for('main_index'))
    try:
        member = supabase.table('members').select('*').eq('id', member_id).single().execute().data
        if not member:
            flash("멤버를 찾을 수 없습니다.", "danger")
            return redirect(url_for('records_members') if session.get('user_role') == 'admin' else url_for('my_page'))
        activity = _aggregate_member_activity(member_id, member.get('name', ''))
        special_events = _member_special_events(member_id)
        return render_template('records_member_profile.html', member=member,
                               special_events=special_events, **activity)
    except Exception as e:
        app.logger.error(f"records_member_profile error: {e}", exc_info=True)
        flash(f"프로필 로딩 오류: {e}", 'danger')
        return redirect(url_for('records_hub'))


# --- 세미나 ---
@app.route('/records/seminars')
@login_required(role="admin")
def records_seminars():
    term_id = request.args.get('term_id') or ''
    page = max(1, request.args.get('page', 1, type=int))
    page_size = 100
    try:
        q = supabase.table('history').select('*', count='exact').order('date', desc=True)
        start, end, _ = _get_term_range(term_id)
        if start and end:
            q = q.gte('date', start).lte('date', end)
        history_response = q.range((page - 1) * page_size, page * page_size - 1).execute()
        history = history_response.data or []
        total_history = history_response.count if history_response.count is not None else len(history)
        total_pages = max(1, math.ceil(total_history / page_size))
        for row in history:
            groups = row.get('groups') or []
            present = []
            for g in groups:
                if isinstance(g, list):
                    present.extend(g)
                elif isinstance(g, str):
                    present.append(g)
            row['present_count'] = len(present)
            row['group_count'] = len(groups) if groups and isinstance(groups[0], list) else 0
            row['facilitator_count'] = len(row.get('facilitators') or [])

        # 6개월 단위 버킷 그룹화 (예: "2025년 상반기", "2025년 하반기")
        buckets = []  # [{'key': str, 'label': str, 'rows': [..]}]
        seen = {}
        for row in history:
            try:
                d = date.fromisoformat(row['date'])
                half = '상반기' if d.month <= 6 else '하반기'
                key = f"{d.year}-{1 if d.month <= 6 else 2}"
                label = f"{d.year}년 {half}"
            except Exception:
                key, label = 'unknown', '날짜 미상'
            if key not in seen:
                seen[key] = {'key': key, 'label': label, 'rows': []}
                buckets.append(seen[key])
            seen[key]['rows'].append(row)
        # history는 이미 date desc 정렬이므로 buckets도 자연스럽게 최신순
        genres = _load_genres()
        terms = _get_terms_for_filter()
    except Exception as e:
        app.logger.error(f"records_seminars error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        history, genres, terms, buckets = [], [], [], []
        total_history, total_pages = 0, 1
    return render_template('records_seminars.html', history=history, genres=genres,
                           terms=terms, selected_term_id=term_id, buckets=buckets,
                           page=page, total_pages=total_pages, total_history=total_history)


@app.route('/records/seminars/<history_id>')
@login_required(role="admin")
def records_seminar_detail(history_id):
    try:
        row = supabase.table('history').select('*').eq('id', history_id).single().execute().data
        groups = row.get('groups') or []
        all_present = []
        for g in groups:
            if isinstance(g, list):
                all_present.extend(g)
            elif isinstance(g, str):
                all_present.append(g)
        member_map = {}
        if all_present:
            mres = supabase.table('members').select('id, name, gender').in_('name', list(set(all_present))).execute().data or []
            member_map = {m['name']: m['id'] for m in mres}
        else:
            mres = []
        member_genders = {m['name']: normalize_gender(m.get('gender')) for m in mres}
        genres = _load_genres()
        all_members = supabase.table('members').select('id, name, gender') \
            .eq('is_active', True).order('name').execute().data or []
        for member in all_members:
            member['gender_code'] = normalize_gender(member.get('gender'))
            member_genders.setdefault(member['name'], member['gender_code'])
        meeting_history = _meeting_details_from_history(
            _effective_group_history_rows(), before_date=row.get('date'), exclude_history_id=history_id
        )
        return render_template('records_seminar_detail.html',
                               row=row, groups=groups, member_map=member_map,
                               genres=genres, total_present=len(all_present),
                               all_members=all_members, member_genders=member_genders,
                               meeting_history=meeting_history)
    except Exception as e:
        app.logger.error(f"records_seminar_detail error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        return redirect(url_for('records_seminars'))


@app.route('/api/admin/history/<history_id>/update_meta', methods=['POST'])
@login_required(role="admin")
def update_history_meta(history_id):
    """세미나 기록 메타 + 본문(날짜/발제자/조) 수정.
    허용 필드: book_title, genre, date, facilitators(list[str]), groups(list[list[str]])
    """
    try:
        data = request.json or {}
        update = {}
        if 'book_title' in data:
            update['book_title'] = (data.get('book_title') or '').strip() or None
        if 'genre' in data:
            update['genre'] = (data.get('genre') or '').strip() or None
        linked_session_id = None
        if 'book_title' in data or 'date' in data:
            link_rows = supabase.table('history').select('seminar_session_id') \
                .eq('id', history_id).execute().data or []
            linked_session_id = (link_rows[0] if link_rows else {}).get('seminar_session_id')
        if 'date' in data:
            d = (data.get('date') or '').strip()
            if not d:
                return jsonify({'status': 'error', 'message': '날짜는 비울 수 없습니다.'}), 400
            if linked_session_id:
                linked_session = supabase.table('seminar_sessions').select('meeting_date') \
                    .eq('id', linked_session_id).single().execute().data
                if linked_session and d != linked_session['meeting_date']:
                    return jsonify({
                        'status': 'error',
                        'message': '세미나 회차에 연결된 기록의 날짜는 변경할 수 없습니다.'
                    }), 400
            update['date'] = d
        if 'facilitators' in data:
            facs = data.get('facilitators') or []
            if not isinstance(facs, list):
                return jsonify({'status': 'error', 'message': 'facilitators는 배열이어야 합니다.'}), 400
            update['facilitators'] = [str(x).strip() for x in facs if str(x).strip()]
        if 'groups' in data:
            try:
                groups, editor_state = normalize_group_editor_payload(data.get('groups'), data.get('group_editor_state'))
                previous_rows = supabase.table('history').select('group_editor_state') \
                    .eq('id', history_id).execute().data or []
                previous_state = (previous_rows[0] if previous_rows else {}).get('group_editor_state') or {}
                if data.get('group_editor_state') is None:
                    if previous_state:
                        assigned_names = [name for group in groups for name in group]
                        exclusions = [item for item in previous_state.get('excluded', []) if item.get('name') not in assigned_names]
                        retained_state = {
                            'participants': list(dict.fromkeys([*previous_state.get('participants', []), *assigned_names])),
                            'excluded': exclusions,
                            'group_names': previous_state.get('group_names') if len(previous_state.get('group_names', [])) == len(groups) else editor_state['group_names'],
                        }
                        # Older clients may move members, but cannot silently discard
                        # tracked people or their exclusion reasons.
                        groups, editor_state = normalize_group_editor_payload(groups, retained_state)
                # A supplied snapshot is not authority to erase previously tracked
                # participants. They must remain assigned or explicitly excluded.
                if not set(previous_state.get('participants', [])).issubset(set(editor_state['participants'])):
                    raise ValueError('기존 참여자를 명단에서 삭제할 수 없습니다. 미배정으로 옮기거나 사유를 입력해 편성에서 제외해주세요.')
                _validate_group_member_names(groups)
            except ValueError as exc:
                return jsonify({'status': 'error', 'message': str(exc)}), 400
            if not _validate_groups_against_restrictions(groups):
                return jsonify({'status': 'error', 'message': '관리자가 지정한 비공개 편성 제한 조건과 충돌합니다. 명단을 다시 편성해주세요.'}), 409
            update['groups'] = groups
            update['present'] = [name for group in groups for name in group]
            update['group_editor_state'] = editor_state
            if 'facilitators' in update:
                update['facilitators'] = list(dict.fromkeys(name for name in update['facilitators'] if name in update['present']))
        if not update and data.get('confirm_attendance') is not True:
            return jsonify({'status': 'error', 'message': '변경할 필드가 없습니다.'}), 400

        # groups/date 변경 시 co_matrix 재계산
        groups_changed = 'groups' in update
        date_changed = 'date' in update
        old_row = None
        if groups_changed or date_changed or data.get('confirm_attendance') is True:
            old_res = supabase.table('history').select('*').eq('id', history_id).execute()
            old_row = (old_res.data or [None])[0]

        if old_row and old_row.get('attendance_confirmed_at') and old_row.get('actual_member_ids') is None:
            # Freeze legacy actual attendance BEFORE replacing the planned groups.
            legacy_present = old_row.get('present')
            if legacy_present is None:
                legacy_present = [name for group in old_row.get('groups') or [] for name in group if isinstance(name, str)]
            try:
                update['actual_member_ids'] = _actual_member_ids_from_names(legacy_present)
            except ValueError as exc:
                return jsonify({'status': 'error', 'message': str(exc)}), 400

        if data.get('confirm_attendance') is True:
            if not old_row:
                return jsonify({'status': 'error', 'message': '기록을 찾을 수 없습니다.'}), 404
            if old_row.get('seminar_session_id'):
                return jsonify({'status': 'error', 'message': '연결된 회차의 실제 출석은 세미나 운영 화면에서 확정해주세요.'}), 400
            actual_date = update.get('date', old_row.get('date'))
            try:
                if not actual_date or datetime.fromisoformat(actual_date).date() > datetime.now(KST).date():
                    raise ValueError('미래 세미나는 실제 참석으로 확정할 수 없습니다.')
                actual_names = update.get('present', old_row.get('present') or [])
                update['actual_member_ids'] = _actual_member_ids_from_names(actual_names)
            except (ValueError, TypeError) as exc:
                return jsonify({'status': 'error', 'message': str(exc)}), 400
            update['attendance_confirmed_at'] = datetime.now(timezone.utc).isoformat()

        supabase.table('history').update(update).eq('id', history_id).execute()

        # 세미나 운영 화면에서 생성된 기록은 도서명이 주차·회차·발제문에도 공유된다.
        # 기록 화면에서 고쳐도 원본 일정과 짝 회차가 즉시 같은 제목을 보도록 역동기화한다.
        if 'book_title' in update and linked_session_id:
            linked_session = supabase.table('seminar_sessions').select('seminar_week_id') \
                .eq('id', linked_session_id).single().execute().data or {}
            week_id = linked_session.get('seminar_week_id')
            title_update = {'book_title': update['book_title']}
            if week_id:
                supabase.table('seminar_weeks').update({
                    **title_update,
                    'needs_review': False,
                    'updated_at': datetime.now(timezone.utc).isoformat(),
                }).eq('id', week_id).execute()
                linked_sessions = supabase.table('seminar_sessions').select('id') \
                    .eq('seminar_week_id', week_id).execute().data or []
                linked_session_ids = [row['id'] for row in linked_sessions]
                supabase.table('seminar_sessions').update(title_update) \
                    .eq('seminar_week_id', week_id).execute()
                supabase.table('topic_events').update(title_update) \
                    .eq('seminar_week_id', week_id).execute()
                if linked_session_ids:
                    supabase.table('history').update(title_update) \
                        .in_('seminar_session_id', linked_session_ids).execute()
            else:
                supabase.table('seminar_sessions').update(title_update) \
                    .eq('id', linked_session_id).execute()
                supabase.table('topic_events').update(title_update) \
                    .eq('seminar_session_id', linked_session_id).execute()

        if old_row is not None and (groups_changed or date_changed or data.get('confirm_attendance') is True):
            old_groups = old_row.get('groups') or []
            old_date = old_row.get('date') or ''
            new_groups = update.get('groups', old_groups)
            new_date = update.get('date', old_date)
            affected_keys = _pair_keys_from_groups(old_groups) | _pair_keys_from_groups(new_groups)
            if affected_keys:
                rebuild_co_matrix(affected_keys)

        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"update_history_meta error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/history/<history_id>/delete', methods=['POST'])
@login_required(role="admin")
def records_history_delete(history_id):
    """세미나 기록 삭제 + co_matrix 차감."""
    try:
        old_res = supabase.table('history').select('groups, date').eq('id', history_id).execute()
        old_row = (old_res.data or [None])[0]
        supabase.table('history').delete().eq('id', history_id).execute()
        if old_row:
            affected_keys = _pair_keys_from_groups(old_row.get('groups') or [])
            if affected_keys:
                rebuild_co_matrix(affected_keys)
        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"records_history_delete error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/capacity', methods=['PATCH'])
@login_required(role="admin")
def seminar_session_capacity(session_id):
    return jsonify({'status': 'error', 'message': '월요일 정원·선착순 제한을 사용하지 않습니다.'}), 410


@app.route('/api/admin/seminar_sessions/<session_id>/absences', methods=['POST'])
@login_required(role="admin")
def seminar_session_add_absence(session_id):
    try:
        data = request.json or {}
        member_ids = normalize_member_ids(data)
        note = (data.get('note') or '').strip()[:500]
        target = supabase.table('seminar_sessions').select('participation_mode,planned_member_ids') \
            .eq('id', session_id).single().execute().data or {}
        if target.get('planned_member_ids') is not None:
            return jsonify({'status': 'error', 'message': '카카오톡 명단 반영 화면에서 불참 명단을 수정해주세요.'}), 409
        if target.get('participation_mode') != 'absence_only':
            return jsonify({'status': 'error', 'message': '목요일 불참 입력 회차가 아닙니다.'}), 400
        members = supabase.table('members').select('id').in_('id', member_ids) \
            .eq('is_active', True).execute().data or []
        valid_member_ids = {row['id'] for row in members}
        invalid_member_ids = [member_id for member_id in member_ids if member_id not in valid_member_ids]
        if invalid_member_ids:
            return jsonify({'status': 'error', 'message': '비활성 또는 존재하지 않는 회원이 포함되어 있습니다.'}), 400

        current = supabase.table('seminar_absences').select('member_id') \
            .eq('session_id', session_id).in_('member_id', member_ids) \
            .is_('cancelled_at', 'null').execute().data or []
        current_member_ids = {row['member_id'] for row in current}
        new_member_ids = [member_id for member_id in member_ids if member_id not in current_member_ids]
        stamp = datetime.now(timezone.utc).isoformat()
        if new_member_ids:
            supabase.table('seminar_absences').insert([
                {
                    'session_id': session_id, 'member_id': member_id, 'note': note or None,
                    'recorded_by': session.get('user_id'), 'updated_at': stamp,
                }
                for member_id in new_member_ids
            ]).execute()
        return jsonify({
            'status': 'success',
            'added_count': len(new_member_ids),
            'already_registered_count': len(current_member_ids),
        })
    except ValueError as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400
    except Exception as e:
        app.logger.error(f"seminar_session_add_absence error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/absences/<int:member_id>', methods=['DELETE'])
@login_required(role="admin")
def seminar_session_cancel_absence(session_id, member_id):
    try:
        target = supabase.table('seminar_sessions').select('planned_member_ids') \
            .eq('id', session_id).single().execute().data or {}
        if target.get('planned_member_ids') is not None:
            return jsonify({'status': 'error', 'message': '카카오톡 명단 반영 화면에서 불참 명단을 수정해주세요.'}), 409
        stamp = datetime.now(timezone.utc).isoformat()
        supabase.table('seminar_absences').update({
            'cancelled_at': stamp, 'cancelled_by': session.get('user_id'), 'updated_at': stamp,
        }).eq('session_id', session_id).eq('member_id', member_id).is_('cancelled_at', 'null').execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"seminar_session_cancel_absence error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/no_shows', methods=['POST'])
@login_required(role="admin")
def seminar_session_add_no_shows(session_id):
    try:
        data = request.json or {}
        member_ids = normalize_member_ids(data)
        note = (data.get('note') or '').strip()[:500]
        target = supabase.table('seminar_sessions').select('participation_mode,planned_member_ids,actual_member_ids,attendance_confirmed_at') \
            .eq('id', session_id).single().execute().data or {}
        if not target:
            return jsonify({'status': 'error', 'message': '세미나 회차를 찾을 수 없습니다.'}), 404

        members = supabase.table('members').select('id').in_('id', member_ids) \
            .eq('is_active', True).execute().data or []
        eligible_ids = {row['id'] for row in members}
        if target.get('planned_member_ids') is not None:
            eligible_ids = set(target['planned_member_ids'])
        elif target.get('participation_mode') == 'opt_in':
            votes = supabase.table('seminar_votes').select('member_id') \
                .eq('session_id', session_id).eq('attending', True).execute().data or []
            eligible_ids &= {row['member_id'] for row in votes}
        elif target.get('participation_mode') == 'absence_only':
            notices = supabase.table('seminar_absences').select('member_id') \
                .eq('session_id', session_id).is_('cancelled_at', 'null').execute().data or []
            eligible_ids -= {row['member_id'] for row in notices}

        invalid_member_ids = [member_id for member_id in member_ids if member_id not in eligible_ids]
        if invalid_member_ids:
            return jsonify({
                'status': 'error',
                'message': '이 회차의 참석 예정자가 아니거나 이미 사전 불참 연락을 남긴 회원이 포함되어 있습니다.',
            }), 400

        current = supabase.table('seminar_no_shows').select('member_id') \
            .eq('session_id', session_id).in_('member_id', member_ids) \
            .is_('cancelled_at', 'null').execute().data or []
        current_member_ids = {row['member_id'] for row in current}
        new_member_ids = [member_id for member_id in member_ids if member_id not in current_member_ids]
        stamp = datetime.now(timezone.utc).isoformat()
        if new_member_ids:
            supabase.table('seminar_no_shows').insert([
                {
                    'session_id': session_id,
                    'member_id': member_id,
                    'note': note or None,
                    'recorded_by': session.get('user_id'),
                    'updated_at': stamp,
                }
                for member_id in new_member_ids
            ]).execute()
        if target.get('attendance_confirmed_at') and target.get('actual_member_ids') is not None:
            supabase.table('seminar_sessions').update({
                'actual_member_ids': [mid for mid in target['actual_member_ids'] if mid not in member_ids],
            }).eq('id', session_id).execute()
        # 이미 저장된 계획표가 있더라도 미연락 불참자는 실제 만남으로 세지 않는다.
        _rebuild_matrix_for_session(session_id)
        return jsonify({
            'status': 'success',
            'added_count': len(new_member_ids),
            'already_registered_count': len(current_member_ids),
        })
    except ValueError as e:
        return jsonify({'status': 'error', 'message': str(e)}), 400
    except Exception as e:
        app.logger.error(f"seminar_session_add_no_shows error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_sessions/<session_id>/no_shows/<int:member_id>', methods=['DELETE'])
@login_required(role="admin")
def seminar_session_cancel_no_show(session_id, member_id):
    try:
        stamp = datetime.now(timezone.utc).isoformat()
        supabase.table('seminar_no_shows').update({
            'cancelled_at': stamp,
            'cancelled_by': session.get('user_id'),
            'updated_at': stamp,
        }).eq('session_id', session_id).eq('member_id', member_id) \
            .is_('cancelled_at', 'null').execute()
        # 취소된 경우에는 해당 회원의 만남을 원래 계획표 기준으로 다시 복구한다.
        _rebuild_matrix_for_session(session_id)
        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"seminar_session_cancel_no_show error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/seminar_weeks/<week_id>', methods=['PATCH'])
@login_required(role="admin")
def seminar_week_update(week_id):
    try:
        data = request.json or {}
        sessions = supabase.table('seminar_sessions').select('id') \
            .eq('seminar_week_id', week_id).order('meeting_date').execute().data or []
        if not sessions:
            return jsonify({'status': 'error', 'message': '주차를 찾을 수 없습니다.'}), 404
        book_title = (data.get('book_title') or '').strip()
        book_author = (data.get('book_author') or '').strip()
        note = (data.get('note') or '').strip()
        update = {'book_title': book_title or None, 'book_author': book_author or None}
        supabase.table('seminar_weeks').update({
            **update, 'note': note or None, 'needs_review': False,
            'updated_at': datetime.now(timezone.utc).isoformat(),
        }).eq('id', week_id).execute()
        supabase.table('seminar_sessions').update(update).eq('seminar_week_id', week_id).execute()
        supabase.table('topic_events').update(update).eq('seminar_week_id', week_id).execute()
        session_ids = [row['id'] for row in sessions]
        supabase.table('history').update({'book_title': book_title or None}) \
            .in_('seminar_session_id', session_ids).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        app.logger.error(f"seminar_week_update error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/bookclub/rebuild_matrix', methods=['POST'])
@login_required(role="admin")
def admin_rebuild_bookclub_matrix():
    """history 전체를 기준으로 만남 매트릭스를 복구한다."""
    try:
        result = rebuild_co_matrix()
        return jsonify({'status': 'success', **result})
    except Exception as e:
        app.logger.error(f"admin_rebuild_bookclub_matrix error: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500


# --- 벽돌책 ---
@app.route('/records/brick_books')
@login_required(role="admin")
def records_brick_books():
    term_id = request.args.get('term_id') or ''
    try:
        books = supabase.table('brick_books').select('*').order('created_at', desc=True).execute().data or []
        start, end, _ = _get_term_range(term_id)
        if books:
            ids = [b['id'] for b in books]
            sq = supabase.table('brick_book_sessions').select('brick_book_id, meeting_date').in_('brick_book_id', ids)
            if start and end:
                sq = sq.gte('meeting_date', start).lte('meeting_date', end)
            sess_res = sq.execute().data or []
            cnt = {}
            for s in sess_res:
                cnt[s['brick_book_id']] = cnt.get(s['brick_book_id'], 0) + 1
            for b in books:
                b['session_count'] = cnt.get(b['id'], 0)
            # 학기 필터링 시: 해당 학기 내 세션이 있는 책만 표시
            if start and end:
                books = [b for b in books if b['session_count'] > 0]
        terms = _get_terms_for_filter()
    except Exception as e:
        app.logger.error(f"records_brick_books error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        books, terms = [], []
    return render_template('records_brick_books.html', books=books,
                           terms=terms, selected_term_id=term_id)


@app.route('/records/brick_books/<book_id>')
@login_required(role="admin")
def records_brick_book_detail(book_id):
    try:
        book = supabase.table('brick_books').select('*').eq('id', book_id).single().execute().data
        sessions_res = supabase.table('brick_book_sessions').select('*') \
            .eq('brick_book_id', book_id).order('meeting_date', desc=True).execute()
        sessions = sessions_res.data or []
        sess_ids = [s['id'] for s in sessions]
        member_map_by_sess = {}
        if sess_ids:
            parts = supabase.table('brick_session_members').select('session_id, member_id, members(id, name)') \
                .in_('session_id', sess_ids).execute().data or []
            for p in parts:
                sid = p['session_id']
                member_map_by_sess.setdefault(sid, []).append(p.get('members') or {})
        for s in sessions:
            s['members'] = member_map_by_sess.get(s['id'], [])
        all_members = supabase.table('members').select('id, name').eq('is_active', True).order('name').execute().data or []
        return render_template('records_brick_book_detail.html',
                               book=book, sessions=sessions, all_members=all_members)
    except Exception as e:
        app.logger.error(f"records_brick_book_detail error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        return redirect(url_for('records_brick_books'))


@app.route('/api/admin/brick_books/create', methods=['POST'])
@login_required(role="admin")
def brick_book_create():
    try:
        data = request.json or {}
        title = (data.get('title') or '').strip()
        if not title:
            return jsonify({'status': 'error', 'message': '제목은 필수입니다.'}), 400
        res = supabase.table('brick_books').insert({
            'title': title,
            'notes': (data.get('notes') or '').strip() or None,
        }).execute()
        return jsonify({'status': 'success', 'book': res.data[0]})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/brick_books/<book_id>/delete', methods=['POST'])
@login_required(role="admin")
def brick_book_delete(book_id):
    try:
        supabase.table('brick_books').delete().eq('id', book_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/brick_books/<book_id>/sessions/add', methods=['POST'])
@login_required(role="admin")
def brick_session_add(book_id):
    try:
        data = request.json or {}
        meeting_date = data.get('meeting_date')
        if not meeting_date:
            return jsonify({'status': 'error', 'message': '날짜가 필요합니다.'}), 400
        sess = supabase.table('brick_book_sessions').insert({
            'brick_book_id': book_id,
            'meeting_date': meeting_date,
            'notes': (data.get('notes') or '').strip() or None,
        }).execute().data[0]
        member_ids = [int(m) for m in (data.get('member_ids') or []) if m]
        if member_ids:
            supabase.table('brick_session_members').insert([
                {'session_id': sess['id'], 'member_id': mid} for mid in member_ids
            ]).execute()
        return jsonify({'status': 'success', 'session': sess})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/brick_sessions/<session_id>/delete', methods=['POST'])
@login_required(role="admin")
def brick_session_delete(session_id):
    try:
        supabase.table('brick_book_sessions').delete().eq('id', session_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# --- 소모임 ---
@app.route('/records/study_groups')
@login_required(role="admin")
def records_study_groups():
    term_id = request.args.get('term_id') or ''
    try:
        groups = supabase.table('study_groups').select('*').order('created_at', desc=True).execute().data or []
        start, end, _ = _get_term_range(term_id)
        if groups:
            ids = [g['id'] for g in groups]
            sq = supabase.table('study_group_sessions').select('study_group_id, meeting_date').in_('study_group_id', ids)
            if start and end:
                sq = sq.gte('meeting_date', start).lte('meeting_date', end)
            sess_res = sq.execute().data or []
            cnt = {}
            for s in sess_res:
                cnt[s['study_group_id']] = cnt.get(s['study_group_id'], 0) + 1
            for g in groups:
                g['session_count'] = cnt.get(g['id'], 0)
            if start and end:
                groups = [g for g in groups if g['session_count'] > 0]
        terms = _get_terms_for_filter()
    except Exception as e:
        app.logger.error(f"records_study_groups error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        groups, terms = [], []
    return render_template('records_study_groups.html', groups=groups,
                           terms=terms, selected_term_id=term_id)


@app.route('/records/study_groups/<group_id>')
@login_required(role="admin")
def records_study_group_detail(group_id):
    try:
        group = supabase.table('study_groups').select('*').eq('id', group_id).single().execute().data
        sessions_res = supabase.table('study_group_sessions').select('*') \
            .eq('study_group_id', group_id).order('meeting_date', desc=True).execute()
        sessions = sessions_res.data or []
        sess_ids = [s['id'] for s in sessions]
        member_map_by_sess = {}
        if sess_ids:
            parts = supabase.table('study_session_members').select('session_id, member_id, members(id, name)') \
                .in_('session_id', sess_ids).execute().data or []
            for p in parts:
                sid = p['session_id']
                member_map_by_sess.setdefault(sid, []).append(p.get('members') or {})
        for s in sessions:
            s['members'] = member_map_by_sess.get(s['id'], [])
        all_members = supabase.table('members').select('id, name').eq('is_active', True).order('name').execute().data or []
        return render_template('records_study_group_detail.html',
                               group=group, sessions=sessions, all_members=all_members)
    except Exception as e:
        app.logger.error(f"records_study_group_detail error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        return redirect(url_for('records_study_groups'))


@app.route('/api/admin/study_groups/create', methods=['POST'])
@login_required(role="admin")
def study_group_create():
    try:
        data = request.json or {}
        name = (data.get('name') or '').strip()
        if not name:
            return jsonify({'status': 'error', 'message': '이름은 필수입니다.'}), 400
        res = supabase.table('study_groups').insert({
            'name': name,
            'notes': (data.get('notes') or '').strip() or None,
        }).execute()
        return jsonify({'status': 'success', 'group': res.data[0]})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/study_groups/<group_id>/delete', methods=['POST'])
@login_required(role="admin")
def study_group_delete(group_id):
    try:
        supabase.table('study_groups').delete().eq('id', group_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/study_groups/<group_id>/sessions/add', methods=['POST'])
@login_required(role="admin")
def study_session_add(group_id):
    try:
        data = request.json or {}
        meeting_date = data.get('meeting_date')
        if not meeting_date:
            return jsonify({'status': 'error', 'message': '날짜가 필요합니다.'}), 400
        sess = supabase.table('study_group_sessions').insert({
            'study_group_id': group_id,
            'meeting_date': meeting_date,
            'notes': (data.get('notes') or '').strip() or None,
        }).execute().data[0]
        member_ids = [int(m) for m in (data.get('member_ids') or []) if m]
        if member_ids:
            supabase.table('study_session_members').insert([
                {'session_id': sess['id'], 'member_id': mid} for mid in member_ids
            ]).execute()
        return jsonify({'status': 'success', 'session': sess})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/admin/study_sessions/<session_id>/delete', methods=['POST'])
@login_required(role="admin")
def study_session_delete(session_id):
    try:
        supabase.table('study_group_sessions').delete().eq('id', session_id).execute()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# --- 통계 ---
@app.route('/records/analytics')
@login_required(role="admin")
def records_analytics():
    term_id = request.args.get('term_id') or ''
    try:
        start, end, _ = _get_term_range(term_id)
        hq = supabase.table('history').select('id, date, genre, present, book_title')
        if start and end:
            hq = hq.gte('date', start).lte('date', end)
        history = hq.order('date').execute().data or []
        genre_counts, monthly_counts = {}, {}
        for row in history:
            g = row.get('genre') or '미분류'
            genre_counts[g] = genre_counts.get(g, 0) + 1
            d = row.get('date')
            if d:
                ym = str(d)[:7]
                monthly_counts[ym] = monthly_counts.get(ym, 0) + 1

        attendance_members, attendance_sessions, attendance_matrix = _build_attendance_matrix(start, end)
        attendance_counts = {
            member['id']: sum(
                1 for session_item in attendance_sessions
                if attendance_matrix.get(member['id'], {}).get(session_item['key'])
            )
            for member in attendance_members
        }
        visible_attendance_members = [
            member for member in attendance_members
            if member.get('is_active') or attendance_counts.get(member['id'], 0) > 0
        ]
        # Ranking also uses confirmed actual attendance, never the saved group plan.
        top_attendees = sorted(
            [(member['name'], attendance_counts.get(member['id'], 0)) for member in attendance_members
             if attendance_counts.get(member['id'], 0) > 0],
            key=lambda item: (-item[1], item[0]),
        )[:15]

        # 벽돌책/소모임 월별 세션 카운트 (학기 필터 적용)
        bb_monthly, sg_monthly = {}, {}
        try:
            bbq = supabase.table('brick_book_sessions').select('meeting_date')
            if start and end:
                bbq = bbq.gte('meeting_date', start).lte('meeting_date', end)
            bb_sess = bbq.execute().data or []
            for s in bb_sess:
                ym = str(s.get('meeting_date') or '')[:7]
                if ym: bb_monthly[ym] = bb_monthly.get(ym, 0) + 1
        except Exception: pass
        try:
            sgq = supabase.table('study_group_sessions').select('meeting_date')
            if start and end:
                sgq = sgq.gte('meeting_date', start).lte('meeting_date', end)
            sg_sess = sgq.execute().data or []
            for s in sg_sess:
                ym = str(s.get('meeting_date') or '')[:7]
                if ym: sg_monthly[ym] = sg_monthly.get(ym, 0) + 1
        except Exception: pass

        # 벽돌책/소모임 카운트: 학기 내 세션이 있는 것만
        if start and end:
            bb_count = len({s.get('brick_book_id') for s in (supabase.table('brick_book_sessions')
                .select('brick_book_id, meeting_date').gte('meeting_date', start).lte('meeting_date', end)
                .execute().data or [])})
            sg_count = len({s.get('study_group_id') for s in (supabase.table('study_group_sessions')
                .select('study_group_id, meeting_date').gte('meeting_date', start).lte('meeting_date', end)
                .execute().data or [])})
        else:
            bb_count = supabase.table('brick_books').select('id', count='exact').execute().count or 0
            sg_count = supabase.table('study_groups').select('id', count='exact').execute().count or 0

        terms = _get_terms_for_filter()
        return render_template('records_analytics.html',
                               total_seminars=len(history),
                               total_brick_books=bb_count,
                               total_study_groups=sg_count,
                               genre_counts=genre_counts,
                               monthly_counts=dict(sorted(monthly_counts.items())),
                               bb_monthly=dict(sorted(bb_monthly.items())),
                               sg_monthly=dict(sorted(sg_monthly.items())),
                               top_attendees=top_attendees,
                               attendance_members=visible_attendance_members,
                               attendance_sessions=attendance_sessions,
                               attendance_matrix=attendance_matrix,
                               attendance_counts=attendance_counts,
                               attendance_start_date=start or '',
                               attendance_end_date=end or '',
                               terms=terms, selected_term_id=term_id)
    except Exception as e:
        app.logger.error(f"records_analytics error: {e}", exc_info=True)
        flash(f"오류: {e}", 'danger')
        return redirect(url_for('records_hub'))




# 게시판은 별도 모듈에 두되 기존 Flask 세션/Supabase 클라이언트를 그대로 공유한다.
from boards import init_board_routes
init_board_routes(app, supabase, login_required)

# 공개 링크 중심 참여 흐름(후기·책 추천·벽돌책)을 별도 모듈로 관리한다.
from engagement import init_engagement_routes
init_engagement_routes(app, supabase, login_required, _voting_window_for,
                       topic_event_lifecycle=_auto_close_topic_events)

from attendance_routes import init_attendance_routes
init_attendance_routes(app, lambda: supabase, login_required, _rebuild_matrix_for_session)


# ==============================================================================
# --- 7. 서버 실행 ---
# ==============================================================================
if __name__ == '__main__':
    # 로컬 개발 전용. 프로덕션(Render)에서는 Dockerfile의 gunicorn으로 구동됨.
    app.run(host='0.0.0.0', port=5000, debug=True)
